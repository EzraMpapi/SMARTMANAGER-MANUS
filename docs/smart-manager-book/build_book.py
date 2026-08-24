from __future__ import annotations

import os
import sys
from datetime import date
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ModuleNotFoundError:
    Document = None

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
FIGURES = ASSETS / "figures"
OUT = ROOT / "deliverables"
TYPST = ROOT / "typst"
MANAGED_ASSET_SOURCE_ROOT = Path(
    os.environ.get(
        "SMART_MANAGER_BOOK_ASSET_ROOT",
        "/home/ubuntu/webdev-static-assets/businesssphere-doc-assets",
    )
).expanduser()
OUT.mkdir(exist_ok=True)
TYPST.mkdir(exist_ok=True)

OWNER = MANAGED_ASSET_SOURCE_ROOT / "ezra-mpapi-owner.png"
LOGO = MANAGED_ASSET_SOURCE_ROOT / "smart-manager-logo.png"
TODAY = "23 August 2026"

modules = [
    ("Public Brand & Marketing Entry", "Ukurasa wa chapa na masoko", "Accueil de marque et marketing", "IMPLEMENTED", "Home.tsx and the public entry experience establish the product proposition and sign-in path."),
    ("Authentication & Secure Onboarding", "Uthibitishaji na uanzishaji salama", "Authentification et intégration sécurisée", "IMPLEMENTED", "LoginModuleEcosystem, account verification, company creation, and protected session boundaries are present."),
    ("Master Application Shell & Navigation", "Ganda kuu la programu na urambazaji", "Coquille applicative et navigation", "IMPLEMENTED", "DashboardLayout and BusinessSphereDashboard provide the authenticated workspace shell."),
    ("Profile Identity Center", "Kituo cha utambulisho wa wasifu", "Centre d’identité du profil", "IMPLEMENTED", "Profile identity and preference services are backed by authenticated self-service operations."),
    ("Executive Dashboard", "Dashibodi ya uongozi", "Tableau de bord exécutif", "IMPLEMENTED", "Executive and command-center surfaces summarize connected business information."),
    ("Daily Business Briefing", "Muhtasari wa kila siku wa biashara", "Briefing commercial quotidien", "PARTIALLY IMPLEMENTED", "Briefing and scheduled report services exist; availability depends on configured data and delivery services."),
    ("CRM & Customer Pipeline", "CRM na mfuatano wa wateja", "CRM et pipeline clients", "IMPLEMENTED", "Customer and lead persistence boundaries, interactions, and sales-facing workflows are represented."),
    ("Sales & Billing", "Mauzo na malipo", "Ventes et facturation", "IMPLEMENTED", "Sales document contracts, invoices, payments, and billing controls are represented in source and database history."),
    ("Point of Sale (POS)", "Mauzo ya moja kwa moja (POS)", "Point de vente (POS)", "IMPLEMENTED", "Transaction, return, customer credit, register control, loyalty, reconciliation, and audit paths are present."),
    ("Inventory & Warehouse Management", "Usimamizi wa hesabu na ghala", "Stocks et entrepôts", "IMPLEMENTED", "Inventory movement and warehouse command-center surfaces connect operational quantities to sales and procurement."),
    ("Procurement & Vendor Management", "Ununuzi na wasambazaji", "Achats et fournisseurs", "IMPLEMENTED", "Procurement persistence boundaries and vendor workflows connect purchasing to stock and finance."),
    ("Finance & Accounting", "Fedha na uhasibu", "Finance et comptabilité", "IMPLEMENTED", "Finance foundations, journal core, reconciliation, and financial command centers are migration-backed."),
    ("Reports & Scheduled Reporting", "Ripoti na ripoti zilizopangwa", "Rapports et rapports planifiés", "IMPLEMENTED", "Report delivery, schedules, exports, and management views are represented by tested server services."),
    ("Human Resources & Payroll", "Rasilimali watu na mishahara", "Ressources humaines et paie", "IMPLEMENTED", "Employee, workforce authorization, Tanzania payroll, leave decisions, and role approvals are present."),
    ("Manufacturing & Work Orders", "Uzalishaji na maagizo ya kazi", "Fabrication et ordres de travail", "PARTIALLY IMPLEMENTED", "The project inventory recognizes the surface and persistence boundaries; operational depth must be verified per deployment."),
    ("Supply Chain & Fleet", "Mnyororo wa ugavi na magari", "Chaîne logistique et flotte", "IMPLEMENTED", "Fleet management source and controls cover operational fleet records and scheduled controls."),
    ("Marketing Campaigns", "Kampeni za masoko", "Campagnes marketing", "UI ONLY", "The surface exists in the module registry; campaign execution depth is not presented as fully verified."),
    ("E-Commerce Storefront", "Duka la kielektroniki", "Boutique e-commerce", "PLANNED", "The project inventory names the surface, but a production commerce contract is not asserted by this book."),
    ("Documents & Secure Files", "Nyaraka na mafaili salama", "Documents et fichiers sécurisés", "IMPLEMENTED", "Document and notebook persistence boundaries are tested; storage configuration remains environment-dependent."),
    ("Projects & Task Management", "Miradi na usimamizi wa kazi", "Projets et tâches", "PARTIALLY IMPLEMENTED", "Project persistence boundaries are represented; complete project operations should be validated with customer data."),
    ("Customer Support & Helpdesk", "Huduma kwa wateja na helpdesk", "Support client et helpdesk", "IMPLEMENTED", "Support workflow, ticket persistence, authenticated policy, metrics, and inbox surfaces are present."),
    ("Enterprise Analytics & BI", "Uchambuzi wa biashara na BI", "Analytique d’entreprise et BI", "PARTIALLY IMPLEMENTED", "Analytics command centers and predictive surfaces exist; results depend on connected records and configuration."),
    ("Notifications & Alerting", "Arifa na tahadhari", "Notifications et alertes", "IMPLEMENTED", "Notification history, reminders, scheduled jobs, and subscription notifications are represented."),
    ("Activity Stream & Audit Evidence", "Mtiririko wa shughuli na ushahidi wa ukaguzi", "Flux d’activité et preuves d’audit", "IMPLEMENTED", "Audit logs, evidence export, compliance views, and billing audits provide traceability."),
    ("Integration Hub", "Kitovu cha miunganisho", "Hub d’intégrations", "CONFIGURATION REQUIRED", "Webhook, provider, email, mobile-money, and data integrations are available only when their services are configured."),
    ("Workflow Studio & Marketplace", "Studio ya mtiririko wa kazi na soko", "Studio des flux et marketplace", "UI ONLY", "The project inventory recognizes the surface; a complete marketplace execution contract is not asserted."),
    ("Collaboration Hub", "Kitovu cha ushirikiano", "Hub de collaboration", "PARTIALLY IMPLEMENTED", "Collaboration and email-link checks exist, subject to configured delivery channels."),
    ("TRA VFD Fiscalization Portal", "Portal ya TRA VFD", "Portail de fiscalisation TRA VFD", "EXTERNAL SERVICE REQUIRED", "TRA fiscal and Z-report services are present but require external credentials and provider availability."),
    ("AI Assistant & Smart Intelligence", "Msaidizi wa AI na akili ya biashara", "Assistant IA et intelligence", "EXTERNAL SERVICE REQUIRED", "AI approvals, summaries, and assistant workflows are protected and require configured model/service access."),
    ("WhatsApp Web Integration", "Muunganisho wa WhatsApp Web", "Intégration WhatsApp Web", "CONFIGURATION REQUIRED", "The UI supports linked-device interaction patterns; external WhatsApp/provider configuration is required."),
    ("Microfinance", "Mikopo midogo", "Microfinance", "IMPLEMENTED", "Production microfinance workflows, credit scoring, collections, and escalation services are represented."),
    ("Money Agent", "Wakala wa fedha", "Agent de transfert d’argent", "IMPLEMENTED", "Money Agent core operations, fee/commission controls, ledger repair, and reconciliation are present."),
    ("VICOBA / SACCOS & Community Groups", "VICOBA / SACCOS na vikundi vya jamii", "VICOBA / SACCOS et groupes communautaires", "IMPLEMENTED", "Community group schema, relationship guards, documents, approvals, and security hardening are migration-backed."),
    ("Healthcare / Clinic", "Afya na kliniki", "Santé et clinique", "IMPLEMENTED", "Healthcare operations, claims, reminders, portals, interoperability, and internal-role boundaries are represented."),
    ("School Management", "Usimamizi wa shule", "Gestion scolaire", "IMPLEMENTED", "School management and linked learner portal surfaces are present with fee, attendance, assignment, and report-card paths."),
    ("Pharmacy Management", "Usimamizi wa famasi", "Gestion de pharmacie", "IMPLEMENTED", "Pharmacy operations and tenant-safe persistence boundaries are represented."),
    ("Hotel & Hospitality", "Hoteli na ukarimu", "Hôtellerie et hospitalité", "IMPLEMENTED", "Hospitality core, POS/services, guest engagement, finance reconciliation, and access remediation are present."),
    ("Restaurant & F&B", "Mgahawa na chakula/vinywaji", "Restaurant et restauration", "IMPLEMENTED", "Restaurant F&B core, operations extension, Tanzania fiscal configuration, and helper hardening are present."),
    ("Fleet Management", "Usimamizi wa magari", "Gestion de flotte", "IMPLEMENTED", "Fleet management core and scheduled control surfaces are represented."),
    ("Banking & MFI", "Benki na MFI", "Banque et IMF", "IMPLEMENTED", "Bank/MFI core, security hardening, credit ledger, disbursement, and workflow completion are present."),
    ("Employee Portal", "Portal ya mfanyakazi", "Portail employé", "IMPLEMENTED", "Employee portal, self-service boundaries, payroll, leave, and authorization controls are present."),
    ("Property Management", "Usimamizi wa mali isiyohamishika", "Gestion immobilière", "IMPLEMENTED", "Property portfolio controls, tenant-safe data boundaries, and scheduled property controls are represented."),
    ("Subscription & Billing", "Usajili na malipo", "Abonnements et facturation", "IMPLEMENTED", "FREE_15 and six paid bonus packages, monthly-only payment contracts, idempotency, and provider verification are live in the applied schema."),
    ("Global Admin Control Center", "Kituo cha udhibiti wa msimamizi wa mfumo", "Centre de contrôle administrateur global", "IMPLEMENTED", "Platform-wide company, user, subscription, audit, integration, security, and health controls are represented."),
    ("Enterprise Settings & Security Control Center", "Mipangilio na udhibiti wa usalama", "Paramètres et contrôle de sécurité", "IMPLEMENTED", "Settings, permissions, passkeys, session recovery, RLS evidence, and security hardening surfaces are present."),
    ("Predictive Analytics", "Uchambuzi wa utabiri", "Analytique prédictive", "PARTIALLY IMPLEMENTED", "Predictive analytics surfaces exist; model quality and business outcomes depend on data quality and configured services."),
]

chapters = [
    {
        "title": "The Story Behind SMART MANAGER",
        "sw": "Hadithi ya SMART MANAGER",
        "fr": "L’histoire de SMART MANAGER",
        "en": "SMART MANAGER is designed around a practical observation: a growing organization loses time and control when its operations are scattered across notebooks, spreadsheets, messaging channels, separate finance tools, disconnected stock records, and manual reports. The product therefore brings the verified business surfaces of this repository into one authenticated workspace, where records, roles, workflows, and management views can be connected rather than repeatedly reconstructed.",
        "sw_text": "SMART MANAGER imejengwa juu ya uelewa wa vitendo: taasisi inayokua hupoteza muda na udhibiti pale shughuli zinapotawanyika kwenye daftari, lahajedwali, ujumbe, zana tofauti za fedha, rekodi za bidhaa na ripoti za mikono. Bidhaa hii huleta sehemu za biashara zilizothibitishwa na mradi huu kwenye nafasi moja salama ya kazi, ambako rekodi, majukumu, mtiririko wa kazi na maoni ya uongozi vinaweza kuunganishwa.",
        "fr_text": "SMART MANAGER part d’un constat pratique : une organisation en croissance perd du temps et du contrôle lorsque ses opérations sont dispersées entre carnets, tableurs, messageries, outils financiers séparés, stocks isolés et rapports manuels. Le produit réunit donc les surfaces métier vérifiées du dépôt dans un espace de travail authentifié, afin de relier les données, les rôles, les processus et les vues de pilotage.",
        "points": [
            ("Fragmented records", "Rekodi zilizotawanyika", "Données fragmentées", "Centralized access reduces duplicate re-entry and improves the path from an operational action to a management view."),
            ("Limited visibility", "Mwonekano mdogo", "Visibilité limitée", "Dashboards, reports, audit evidence, and command centers make the state of connected operations easier to inspect."),
            ("Weak accountability", "Uwajibikaji dhaifu", "Responsabilité limitée", "Authentication, roles, company scope, audit events, and protected procedures establish a clearer control boundary."),
            ("Scaling pressure", "Shinikizo la ukuaji", "Pression de croissance", "A modular architecture allows an organization to adopt relevant capabilities without claiming that every surface is equally complete."),
        ],
    },
    {
        "title": "Why SMART MANAGER?",
        "sw": "Kwa nini SMART MANAGER?",
        "fr": "Pourquoi SMART MANAGER ?",
        "en": "The product name expresses a management promise rather than a claim of artificial intelligence everywhere. SMART means connected information, automation where implemented, intelligent analysis where configured, real-time visibility where data is available, and support for decisions. MANAGER means control, planning, coordination, accountability, monitoring, execution, and growth. The intended movement is from uncertainty to visibility, from visibility to control, from control to intelligence, and from intelligence to action.",
        "sw_text": "Jina la bidhaa linaeleza ahadi ya usimamizi, si dai kwamba kila sehemu ina akili bandia. SMART inamaanisha taarifa zilizounganishwa, otomatiki pale ilipojengwa, uchambuzi wa akili pale umewekwa, mwonekano wa karibu na wakati halisi pale data ipo, na msaada wa maamuzi. MANAGER inamaanisha udhibiti, mipango, uratibu, uwajibikaji, ufuatiliaji, utekelezaji na ukuaji.",
        "fr_text": "Le nom du produit exprime une promesse de gestion, et non l’affirmation que chaque surface utilise l’intelligence artificielle. SMART signifie informations connectées, automatisation lorsqu’elle est implémentée, analyse intelligente lorsqu’elle est configurée, visibilité proche du temps réel lorsque les données sont disponibles et aide à la décision. MANAGER signifie contrôle, planification, coordination, responsabilité, suivi, exécution et croissance.",
        "points": [
            ("Uncertainty → visibility", "Kutokuwa na uhakika → mwonekano", "Incertitude → visibilité", "Management begins with a trustworthy view of what is recorded and what remains unconfigured."),
            ("Visibility → control", "Mwonekano → udhibiti", "Visibilité → contrôle", "Roles, permissions, tenant boundaries, and auditability turn information into controlled action."),
            ("Control → intelligence", "Udhibiti → uelewa", "Contrôle → intelligence", "Reports, analytics, AI services, and reconciliations add interpretation when their prerequisites exist."),
            ("Intelligence → growth", "Uelewa → ukuaji", "Intelligence → croissance", "A disciplined operational record becomes a foundation for measured expansion."),
        ],
    },
    {
        "title": "SMART MANAGER for Every Business Size",
        "sw": "SMART MANAGER kwa kila ukubwa wa biashara",
        "fr": "SMART MANAGER pour chaque taille d’entreprise",
        "en": "Small businesses can use the platform to establish orderly customer, sales, stock, payment, and accountability records early. Growing businesses can connect more people, branches, departments, and transaction streams. Larger organizations can use the company-scoped architecture, role boundaries, audit evidence, dashboards, integrations, and global administration surfaces where those capabilities are enabled and configured. The product should be adopted according to verified operational need, not according to a generic promise that every module is complete for every organization.",
        "sw_text": "Biashara ndogo zinaweza kutumia mfumo kuanzisha mapema rekodi zilizopangwa za wateja, mauzo, bidhaa, malipo na uwajibikaji. Biashara zinazokua zinaweza kuunganisha watu, matawi, idara na miamala zaidi. Mashirika makubwa yanaweza kutumia usanifu wa kampuni, mipaka ya majukumu, ushahidi wa ukaguzi, dashibodi, miunganisho na usimamizi wa jumla pale uwezo huo umewezeshwa na kusanidiwa.",
        "fr_text": "Les petites entreprises peuvent utiliser la plateforme pour structurer tôt les données clients, ventes, stocks, paiements et responsabilités. Les entreprises en croissance peuvent relier davantage de personnes, de succursales, de départements et de flux de transactions. Les grandes organisations peuvent exploiter l’architecture par entreprise, les rôles, les preuves d’audit, les tableaux de bord, les intégrations et l’administration globale lorsque ces capacités sont activées et configurées.",
        "points": [
            ("Small business", "Biashara ndogo", "Petite entreprise", "Start with the core surfaces that create control: customers, sales, stock, payments, finance, users, and reporting."),
            ("Growing business", "Biashara inayokua", "Entreprise en croissance", "Add procurement, POS, workforce, support, property, industry workspaces, and scheduled controls as operational complexity increases."),
            ("Large business", "Biashara kubwa", "Grande entreprise", "Use company scope, role-aware access, audit, command centers, integrations, and platform administration with governance."),
        ],
    },
    {
        "title": "The Product Tour and Operating Model",
        "sw": "Ziara ya bidhaa na namna ya uendeshaji",
        "fr": "Visite du produit et modèle opératoire",
        "en": "The verified user journey begins at the public brand entry, continues through authentication and account verification, then establishes or joins a company workspace. The authenticated shell presents role-aware navigation and connects the user to the modules that are available for the workspace. An operational action moves through UI validation, a protected server boundary, database constraints and tenant scope, an auditable result, and—where configured—a notification, report, integration, or scheduled follow-up.",
        "sw_text": "Safari ya mtumiaji iliyo kwenye mradi huanza kwenye ukurasa wa chapa, inaendelea kupitia uthibitishaji wa akaunti, kisha huunda au kujiunga na nafasi ya kampuni. Ganda salama la programu huonyesha urambazaji unaozingatia jukumu na kuunganisha mtumiaji na moduli zinazopatikana. Kitendo cha uendeshaji hupitia uthibitishaji wa UI, mpaka wa seva, masharti ya database na kampuni, matokeo yanayoweza kukaguliwa, na—pale kumesanidiwa—arifa, ripoti, muunganisho au ufuatiliaji uliopangwa.",
        "fr_text": "Le parcours vérifié commence sur l’entrée publique de la marque, se poursuit par l’authentification et la vérification du compte, puis crée ou rejoint un espace d’entreprise. La coquille authentifiée présente une navigation adaptée au rôle et relie l’utilisateur aux modules disponibles. Une action opérationnelle traverse la validation de l’interface, une frontière serveur protégée, les contraintes de base de données et le périmètre de l’entreprise, un résultat auditable et, lorsque configuré, une notification, un rapport, une intégration ou un suivi planifié.",
        "points": [
            ("Discover → register", "Gundua → sajili", "Découvrir → s’inscrire", "Public entry and secure onboarding establish the first relationship with the platform."),
            ("Company → workspace", "Kampuni → nafasi ya kazi", "Entreprise → espace de travail", "Company identity becomes the boundary for data, roles, and operational context."),
            ("Action → evidence", "Kitendo → ushahidi", "Action → preuve", "Protected procedures, database constraints, and audit records provide the control path."),
            ("Result → management", "Matokeo → uongozi", "Résultat → pilotage", "Reports and command centers turn recorded work into a management conversation."),
        ],
    },
    {
        "title": "The SMART MANAGER Module Encyclopedia",
        "sw": "Ensaiklopidia ya moduli za SMART MANAGER",
        "fr": "Encyclopédie des modules SMART MANAGER",
        "en": "The following inventory is grounded in the project’s verified module registry, source files, migration names, and tested persistence boundaries. Status labels are deliberate: IMPLEMENTED means a connected source and/or database workflow is present; PARTIALLY IMPLEMENTED means only part of the end-to-end contract is verified; UI ONLY means a surface exists without a complete persistence claim; CONFIGURATION REQUIRED and EXTERNAL SERVICE REQUIRED identify prerequisites; PLANNED and NOT VERIFIED are not product promises.",
        "sw_text": "Orodha ifuatayo imetokana na rejista ya moduli, mafaili ya chanzo, majina ya migrations na mipaka ya uhifadhi iliyojaribiwa. Lebo za hali zimetumika kwa makusudi: IMEJENGWA inamaanisha chanzo na/au mtiririko wa database umeonekana; IMEJENGWA KWA SEHEMU inamaanisha mkataba kamili bado haujathibitishwa; UI TU inamaanisha mwonekano upo bila dai la uhifadhi kamili; INAHITAJI USANIDI na INAHITAJI HUDUMA YA NJE zinaonyesha masharti; ILIYOPANGWA si ahadi ya sasa.",
        "fr_text": "L’inventaire suivant est fondé sur le registre des modules, les fichiers source, les noms de migrations et les frontières de persistance testées. Les statuts sont intentionnels : IMPLÉMENTÉ signifie qu’un flux source et/ou base de données est présent ; PARTIELLEMENT IMPLÉMENTÉ signifie que le contrat de bout en bout n’est pas entièrement vérifié ; INTERFACE SEULE indique une surface sans promesse de persistance complète ; CONFIGURATION REQUISE et SERVICE EXTERNE REQUIS signalent des prérequis ; PLANIFIÉ n’est pas une promesse actuelle.",
        "points": [],
    },
    {
        "title": "Connected Workflows",
        "sw": "Mtiririko wa kazi uliounganishwa",
        "fr": "Flux de travail connectés",
        "en": "SMART MANAGER should be understood as a set of connected control paths rather than isolated screens. The most important connections are sales to inventory and finance; procurement to vendors, stock, and payables; customers to CRM, sales, payments, and support; users to roles, permissions, modules, actions, and audit; and subscriptions to company access, payment verification, invoices, notifications, and reporting.",
        "sw_text": "SMART MANAGER inapaswa kueleweka kama seti ya njia za udhibiti zilizounganishwa, si skrini zilizotengwa. Miunganisho muhimu ni mauzo na hesabu pamoja na fedha; ununuzi na wasambazaji, bidhaa na madeni; wateja na CRM, mauzo, malipo na huduma; watumiaji na majukumu, ruhusa, moduli, vitendo na ukaguzi; na usajili na ruhusa ya kampuni, uthibitishaji wa malipo, ankara, arifa na ripoti.",
        "fr_text": "SMART MANAGER doit être compris comme un ensemble de chemins de contrôle connectés et non comme des écrans isolés. Les connexions majeures sont ventes–stocks–finance ; achats–fournisseurs–stocks–dettes ; clients–CRM–ventes–paiements–support ; utilisateurs–rôles–permissions–modules–actions–audit ; et abonnements–accès entreprise–vérification des paiements–factures–notifications–rapports.",
        "points": [
            ("Sales → Inventory → Finance → Reports", "Mauzo → Hesabu → Fedha → Ripoti", "Ventes → Stocks → Finance → Rapports", "A recorded sale can become stock movement, financial evidence, and management information where the relevant contracts are enabled."),
            ("Procurement → Supplier → Payables → Finance", "Ununuzi → Msambazaji → Madeni → Fedha", "Achats → Fournisseur → Dettes → Finance", "Purchasing controls are designed to preserve the relationship between what was ordered, received, owed, and reported."),
            ("User → Role → Permission → Action → Audit", "Mtumiaji → Jukumu → Ruhusa → Kitendo → Ukaguzi", "Utilisateur → Rôle → Permission → Action → Audit", "Security is a workflow: identity and role boundaries precede access, mutation, and evidence."),
            ("Subscription → Access → Payment → Invoice", "Usajili → Ruhusa → Malipo → Ankara", "Abonnement → Accès → Paiement → Facture", "The billing model uses server and database checks so browser state cannot grant an entitlement."),
        ],
    },
    {
        "title": "Subscriptions and Customer Value",
        "sw": "Usajili na thamani kwa mteja",
        "fr": "Abonnements et valeur client",
        "en": "The live subscription model has one free package and six paid packages. FREE_15 provides 15 days at TZS 0 and requires no payment. Each paid package is monthly-only and represents one paid calendar month plus one promotional bonus calendar month, for two total calendar months. The application and database verify the selected package, amount, tenant, provider order, payment state, and idempotency key. Expired Free access transitions to RequiresPlan without an automatic charge; paid expiry uses calendar-month arithmetic.",
        "sw_text": "Mfumo wa sasa wa usajili una kifurushi kimoja cha bure na vifurushi sita vya kulipia. FREE_15 hutoa siku 15 kwa TZS 0 bila malipo. Kila kifurushi cha kulipia ni cha kila mwezi na kinawakilisha mwezi mmoja uliolipiwa pamoja na mwezi mmoja wa ziada wa ofa, jumla ya miezi miwili ya kalenda. Programu na database huthibitisha kifurushi, kiasi, kampuni, oda ya mtoa huduma, hali ya malipo na ufunguo wa kutorudia.",
        "fr_text": "Le modèle actuel comprend une offre gratuite et six offres payantes. FREE_15 offre 15 jours à 0 TZS sans paiement. Chaque offre payante est mensuelle et représente un mois civil payé plus un mois civil promotionnel, soit deux mois au total. L’application et la base vérifient l’offre, le montant, l’entreprise, la commande du prestataire, l’état du paiement et la clé d’idempotence. L’expiration du gratuit passe à RequiresPlan sans prélèvement automatique ; l’expiration payante utilise des mois calendaires.",
        "points": [
            ("FREE_15", "FREE_15", "FREE_15", "TZS 0; 15 days; no payment; no automatic charge; expiry leads to a paid-plan requirement while data is retained."),
            ("TWIGA / TEMBO / SIMBA", "TWIGA / TEMBO / SIMBA", "TWIGA / TEMBO / SIMBA", "TZS 5,000 / 10,000 / 15,000 per month; one paid month plus one bonus month; two calendar months total."),
            ("SIMBA SC / YANGA SC / AZAM FC", "SIMBA SC / YANGA SC / AZAM FC", "SIMBA SC / YANGA SC / AZAM FC", "TZS 4,500 / 9,000 / 7,000 per month; one paid month plus one bonus month; two calendar months total."),
        ],
    },
    {
        "title": "Operating Routines and Learning Paths",
        "sw": "Ratiba za uendeshaji na njia za kujifunza",
        "fr": "Routines opératoires et parcours d’apprentissage",
        "en": "A disciplined customer should establish a first-day baseline, a weekly review habit, and a monthly management rhythm. Beginners learn identity, navigation, company context, and safe record creation. Operational users learn their module’s inputs, validations, outputs, and exception paths. Managers learn dashboards, approvals, reconciliations, reports, and accountability. Administrators learn tenancy, permissions, integrations, scheduled services, database migrations, security evidence, and recovery boundaries.",
        "sw_text": "Mteja mwenye nidhamu aanzishe msingi wa siku ya kwanza, utaratibu wa mapitio ya kila wiki na mzunguko wa uongozi wa kila mwezi. Mwanzo ajifunze utambulisho, urambazaji, kampuni na uundaji salama wa rekodi. Mtumiaji wa uendeshaji ajifunze pembejeo, uthibitishaji, matokeo na njia za kasoro za moduli yake. Meneja ajifunze dashibodi, vibali, upatanisho, ripoti na uwajibikaji. Msimamizi ajifunze kampuni, ruhusa, miunganisho, huduma zilizopangwa, migrations, ushahidi wa usalama na mipaka ya urejeshaji.",
        "fr_text": "Un client discipliné établit une base le premier jour, une revue hebdomadaire et un rythme mensuel de pilotage. Le débutant apprend l’identité, la navigation, le contexte de l’entreprise et la création sûre de données. L’utilisateur opérationnel apprend les entrées, validations, sorties et exceptions de son module. Le manager apprend les tableaux de bord, validations, rapprochements, rapports et responsabilités. L’administrateur apprend la multi-tenance, les permissions, les intégrations, les services planifiés, les migrations, les preuves de sécurité et les limites de reprise.",
        "points": [
            ("Beginner", "Mwanzo", "Débutant", "Learn the workspace, role, module purpose, safe input, and where to find help."),
            ("Operational user", "Mtumiaji wa uendeshaji", "Utilisateur opérationnel", "Execute repeatable workflows, check validations, and resolve normal exceptions."),
            ("Manager", "Meneja", "Manager", "Review movement, cash, customers, people, exceptions, reports, and approvals."),
            ("Administrator / expert", "Msimamizi / mtaalamu", "Administrateur / expert", "Govern access, configuration, integrations, scheduled services, audit, and recovery."),
        ],
    },
    {
        "title": "Security, Administration, and Trust",
        "sw": "Usalama, usimamizi na uaminifu",
        "fr": "Sécurité, administration et confiance",
        "en": "The project’s security model is layered. Authentication establishes the user session. Authorization establishes role and manager boundaries. Company scope and Row Level Security help separate tenant data. Server-side API handlers keep provider credentials away from the browser. Database functions use explicit search paths, role grants, constraints, and audit records. Global administration is a separate control plane and must not be confused with ordinary company operations. Customers should treat configuration, provider credentials, and recovery procedures as part of the security program, not as afterthoughts.",
        "sw_text": "Mfumo wa usalama una tabaka kadhaa. Uthibitishaji huanzisha kikao cha mtumiaji. Uidhinishaji huweka mipaka ya jukumu na meneja. Kampuni na Row Level Security husaidia kutenganisha data. Handlers za seva huweka siri za watoa huduma mbali na browser. Functions za database hutumia search path wazi, grants, masharti na rekodi za ukaguzi. Usimamizi wa jumla ni eneo tofauti na shughuli za kampuni za kawaida.",
        "fr_text": "Le modèle de sécurité est en couches. L’authentification établit la session. L’autorisation définit les limites de rôle et de gestion. Le périmètre entreprise et la sécurité au niveau des lignes contribuent à séparer les données. Les handlers serveur gardent les secrets des prestataires hors du navigateur. Les fonctions de base utilisent des chemins explicites, des droits, des contraintes et des audits. L’administration globale est un plan de contrôle distinct des opérations ordinaires.",
        "points": [
            ("Authentication", "Uthibitishaji", "Authentification", "Use verified sessions and do not share credentials or bypass the supported entry path."),
            ("Authorization and RLS", "Uidhinishaji na RLS", "Autorisation et RLS", "Use role, company, and policy boundaries; never treat a browser-visible control as proof of permission."),
            ("Secrets and integrations", "Siri na miunganisho", "Secrets et intégrations", "Keep Supabase service credentials, provider keys, and webhook secrets server-side."),
            ("Audit and recovery", "Ukaguzi na urejeshaji", "Audit et reprise", "Keep audit evidence, backups, migration records, and incident procedures reviewable."),
        ],
    },
    {
        "title": "Current Limitations and the Future",
        "sw": "Mipaka ya sasa na mustakabali",
        "fr": "Limites actuelles et avenir",
        "en": "A trustworthy product book must distinguish what is implemented from what is planned. Some surfaces depend on external providers, credentials, scheduled execution, or connected business records. Some registry entries are UI-only or partially verified. The roadmap should therefore prioritize deployment alignment, complete provider configuration, stronger endpoint-by-endpoint review of privileged routines, richer customer-facing screenshots, and measured validation of each industry workflow. Future vision should extend the verified foundation rather than obscure current limitations.",
        "sw_text": "Kitabu cha bidhaa kinachoaminika lazima kitenganishe kilichojengwa na kilichopangwa. Baadhi ya sehemu zinahitaji watoa huduma wa nje, siri, uendeshaji uliopangwa au rekodi za biashara zilizounganishwa. Baadhi ya entries za rejista ni UI tu au zimehakikiwa kwa sehemu. Kwa hiyo roadmap ianze na ulinganifu wa deployment, usanidi wa huduma, mapitio ya routines zenye nguvu, screenshots zaidi za wateja na uthibitishaji wa kila workflow ya sekta.",
        "fr_text": "Un livre produit fiable doit distinguer ce qui est implémenté de ce qui est planifié. Certaines surfaces dépendent de prestataires externes, de secrets, d’exécution planifiée ou de données métier connectées. Certaines entrées du registre sont uniquement des interfaces ou partiellement vérifiées. La feuille de route doit donc commencer par l’alignement du déploiement, la configuration des services, la revue des routines privilégiées, davantage de captures client et la validation mesurée de chaque workflow sectoriel.",
        "points": [
            ("Implemented", "Imejengwa", "Implémenté", "Connected source, database, or tested workflow evidence exists."),
            ("Configuration required", "Inahitaji usanidi", "Configuration requise", "The contract exists but depends on environment variables, roles, storage, provider, or schedule configuration."),
            ("External service required", "Inahitaji huduma ya nje", "Service externe requis", "TRA, AI, WhatsApp, email, mobile money, or other providers must be configured and available."),
            ("Planned / not verified", "Iliyopangwa / haijathibitishwa", "Planifié / non vérifié", "Do not market this classification as a current guaranteed capability."),
        ],
    },
]


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "1F2937"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(8.3)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, title: str, sw: str, fr: str, level: int = 1):
    p = doc.add_heading(title, level=level)
    p.style.font.name = "Aptos Display"
    p.add_run(f"\n{sw}\n{fr}").italic = True
    return p


def add_trilingual_block(doc: Document, en: str, sw: str, fr: str):
    for label, text, color in [("ENGLISH", en, "0B5D3B"), ("KISWAHILI", sw, "1D4ED8"), ("FRANÇAIS", fr, "7C2D12")]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(color)
        r.font.size = Pt(8.5)
        r2 = p.add_run(text)
        r2.font.name = "Aptos"
        r2.font.size = Pt(10.5)


def add_bilingual_table(doc: Document, rows: Iterable[tuple[str, str, str, str]], headers: tuple[str, str, str, str]):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "0B5D3B")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.3):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].italic = True
    cp.runs[0].font.size = Pt(8.5)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("SMART MANAGER ERP  •  ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def build_docx():
    if Document is None:
        raise RuntimeError("The python-docx package is required to rebuild the editable Word document.")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in [("Title", 30, "0B5D3B"), ("Heading 1", 20, "0B5D3B"), ("Heading 2", 15, "1D4ED8"), ("Heading 3", 12, "7C2D12")]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
        styles[name].font.bold = True
    header = section.header.paragraphs[0]
    header.text = "SMART MANAGER ERP  |  Official Enterprise Manual"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(36)
    cover.add_run().add_picture(str(LOGO), width=Inches(6.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("SMART MANAGER ERP")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(11, 93, 59)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THE OFFICIAL ENTERPRISE MANUAL\nAND BUSINESS OPERATIONS GUIDE")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(31, 41, 55)
    for line in ["Your Business. Connected. Controlled. Intelligent.", "Prepared by Ezra Mpapi", "Owner & Creator of SMART MANAGER ERP", "Dar es Salaam, Tanzania", "English • Kiswahili • Français", TODAY]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10.5)
    doc.add_page_break()

    # Creator page
    add_heading(doc, "About the Creator", "Kuhusu Muumba", "À propos du créateur", 1)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(OWNER), width=Inches(3.2))
    add_trilingual_block(doc,
        "Ezra Mpapi is identified in the supplied materials as the Owner & Creator of SMART MANAGER ERP. This book intentionally does not add an unverified biography, qualification, award, or employment history.",
        "Ezra Mpapi ametajwa kwenye nyenzo zilizotolewa kama Mmiliki na Muumba wa SMART MANAGER ERP. Kitabu hiki hakiongezi wasifu, sifa, tuzo au historia ya ajira ambayo haijathibitishwa.",
        "Les documents fournis identifient Ezra Mpapi comme propriétaire et créateur de SMART MANAGER ERP. Ce livre n’ajoute volontairement aucune biographie, qualification, récompense ou expérience non vérifiée.")
    doc.add_paragraph("Headquarters / Makao makuu / Siège : Dar es Salaam, Tanzania")
    doc.add_page_break()

    # Language and contents
    add_heading(doc, "How to Use This Book", "Jinsi ya kutumia kitabu hiki", "Comment utiliser ce livre", 1)
    add_trilingual_block(doc,
        "This is a trilingual product book, user guide, business operations handbook, administrator reference, subscription guide, and architecture overview. English is the primary reference language; every major chapter is accompanied by Kiswahili and French editorial guidance. Technical names, API routes, table names, and package codes remain in their verified project form.",
        "Hiki ni kitabu cha bidhaa, mwongozo wa mtumiaji, kitabu cha uendeshaji wa biashara, rejea ya msimamizi, mwongozo wa usajili na muhtasari wa usanifu kwa lugha tatu. Kiingereza ni lugha kuu ya rejea; kila sura kuu ina maelezo ya Kiswahili na Kifaransa. Majina ya kiufundi, routes, table na codes za vifurushi zimeachwa kama zilivyothibitishwa kwenye mradi.",
        "Il s’agit d’un livre produit trilingue, d’un guide utilisateur, d’un manuel des opérations, d’une référence administrateur, d’un guide des abonnements et d’un aperçu d’architecture. L’anglais est la langue de référence principale ; chaque chapitre majeur est accompagné d’un contenu éditorial en kiswahili et en français. Les noms techniques, routes, tables et codes restent ceux vérifiés dans le projet.")
    add_heading(doc, "Contents", "Yaliyomo", "Sommaire", 1)
    for i, chapter in enumerate(chapters, 1):
        doc.add_paragraph(f"{i}. {chapter['title']} / {chapter['sw']} / {chapter['fr']}")
    doc.add_paragraph("Reference sections: module encyclopedia, subscription table, workflows, limitations, glossary, FAQ, troubleshooting, best practices, and indexes.")
    doc.add_page_break()

    # Main chapters
    for i, chapter in enumerate(chapters, 1):
        add_heading(doc, f"Part {i}: {chapter['title']}", chapter["sw"], chapter["fr"], 1)
        add_trilingual_block(doc, chapter["en"], chapter["sw_text"], chapter["fr_text"])
        if chapter["points"]:
            rows = [(a, b, c, d) for a, b, c, d in chapter["points"]]
            add_bilingual_table(doc, rows, ("English focus", "Kiswahili", "Français", "Verified guidance"))
        if chapter["title"] == "The Story Behind SMART MANAGER":
            add_figure(doc, FIGURES / "01-live-homepage.png", "Figure 1 — Existing SMART MANAGER public homepage evidence.")
        elif chapter["title"] == "The Product Tour and Operating Model":
            add_figure(doc, FIGURES / "03-auth-workflow.png", "Figure 2 — Authentication and secure onboarding workflow evidence.")
            add_figure(doc, FIGURES / "02-live-module-map.png", "Figure 3 — Live module map evidence from the project architecture review.")
        elif chapter["title"] == "Connected Workflows":
            add_figure(doc, FIGURES / "04-business-workflow.png", "Figure 4 — Connected business workflow evidence.")
            add_figure(doc, FIGURES / "09-sales-to-receipt.png", "Figure 5 — Sales-to-receipt workflow evidence.")
            add_figure(doc, FIGURES / "10-inventory-movement.png", "Figure 6 — Inventory movement workflow evidence.")
        elif chapter["title"] == "Subscriptions and Customer Value":
            add_figure(doc, FIGURES / "08-webhook-state-machine.png", "Figure 7 — Provider/webhook state-machine evidence for verified payment handling.")
            add_bilingual_table(doc, [
                ("FREE_15", "FREE", "FREE", "TZS 0; 15 days; 0 paid / 0 bonus / 0 total months; no payment."),
                ("TWIGA", "TWIGA", "TWIGA", "TZS 5,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("TEMBO", "TEMBO", "TEMBO", "TZS 10,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("SIMBA", "SIMBA", "SIMBA", "TZS 15,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("SIMBA_SC", "SIMBA SC", "SIMBA SC", "TZS 4,500/month; 1 paid + 1 bonus; 2 calendar months."),
                ("YANGA_SC", "YANGA SC", "YANGA SC", "TZS 9,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("AZAM_FC", "AZAM FC", "AZAM FC", "TZS 7,000/month; 1 paid + 1 bonus; 2 calendar months."),
            ], ("Code", "Name", "Nom", "Verified commercial contract"))
        elif chapter["title"] == "Security, Administration, and Trust":
            add_figure(doc, FIGURES / "06-auth-tenancy.png", "Figure 8 — Authentication, tenancy, and protected access evidence.")
            add_figure(doc, FIGURES / "07-financial-control.png", "Figure 9 — Financial control-flow evidence.")
        elif chapter["title"] == "Current Limitations and the Future":
            add_figure(doc, FIGURES / "05-deployment-observation.png", "Figure 10 — Deployment observation evidence; deployment configuration remains an operational prerequisite.")
        doc.add_page_break()

    add_heading(doc, "Complete Module Encyclopedia", "Ensaiklopidia kamili ya moduli", "Encyclopédie complète des modules", 1)
    doc.add_paragraph("The module table is intentionally concise so that customers can see the verified purpose, localization, status, and evidence boundary. A status is not a marketing embellishment; it is a delivery control.")
    add_bilingual_table(doc, [(a, b, c, f"{d}: {e}") for a, b, c, d, e in modules], ("Module / English", "Kiswahili", "Français", "Status and verified evidence"))
    doc.add_page_break()

    add_heading(doc, "Workflow Reference", "Rejea ya mitiririko ya kazi", "Référence des workflows", 1)
    workflows = [
        ("Authentication", "Public entry → session verification → profile/company context → protected workspace → audit where applicable."),
        ("Company onboarding", "Register → verify account → create or join company → activate FREE_15 through authenticated server/database control → enter workspace."),
        ("Sales", "Customer → quote/order/document → validation → inventory/finance path where enabled → payment/receipt → report and audit."),
        ("Procurement", "Supplier → purchase request/order → receipt → stock movement → payable/finance evidence → management report."),
        ("POS", "Cashier/register → item/customer → server validation → transaction → receipt/return/credit path → reconciliation and audit."),
        ("Property", "Property portfolio → company-scoped records → scheduled controls → tenant-safe access → reporting and maintenance pathways where enabled."),
        ("Subscription", "Plan catalog → authenticated access → server-verified payment intent → provider order → verified status/webhook → payment/subscription/invoice/audit."),
        ("Administration", "Platform administrator → scoped control plane → company/user/subscription/security/integration action → audit evidence → support or recovery path."),
    ]
    add_bilingual_table(doc, [(a, a, a, b) for a, b in workflows], ("Workflow", "Kiswahili", "Français", "Verified path"))
    doc.add_paragraph("The workflow descriptions are deliberately qualified: a path is only fully operational when its external provider, environment, tenant data, role, and scheduled-service prerequisites are available.")

    add_heading(doc, "Business Problem → SMART MANAGER Solution", "Changamoto ya biashara → Suluhisho", "Problème → Solution SMART MANAGER", 1)
    matrix = [
        ("Fragmented records", "Rekodi zilizotawanyika", "Données fragmentées", "Connected workspace, module boundaries, and shared management views."),
        ("Poor inventory visibility", "Mwonekano mdogo wa bidhaa", "Visibilité limitée des stocks", "Inventory movements connected to sales, procurement, POS, and reporting where enabled."),
        ("Uncontrolled access", "Ufikiaji usiodhibitiwa", "Accès non contrôlé", "Authenticated sessions, roles, company scope, RLS, policies, and audit evidence."),
        ("Payment uncertainty", "Kutokuwa na uhakika wa malipo", "Incertitude de paiement", "Server-side amount verification, provider order checks, idempotency, invoices, and audit."),
        ("Manual reporting", "Ripoti za mikono", "Rapports manuels", "Dashboards, report services, scheduled delivery, and exports where configured."),
    ]
    add_bilingual_table(doc, matrix, ("Business struggle", "Kiswahili", "Français", "Verified response"))

    add_heading(doc, "Frequently Asked Questions", "Maswali yanayoulizwa mara kwa mara", "Questions fréquentes", 1)
    faqs = [
        ("What is SMART MANAGER?", "SMART MANAGER ni nini?", "Qu’est-ce que SMART MANAGER ?", "It is an integrated, authenticated business operating platform whose verified repository surfaces span core ERP, industry workspaces, controls, and administration."),
        ("Can a small business use it?", "Biashara ndogo inaweza kuitumia?", "Une petite entreprise peut-elle l’utiliser ?", "Yes, beginning with the relevant verified core surfaces and adopting more controls as the organization grows."),
        ("How does subscription work?", "Usajili unafanyaje kazi?", "Comment fonctionne l’abonnement ?", "FREE_15 gives 15 days at TZS 0; paid packages are monthly-only with one paid month plus one bonus month."),
        ("Does the browser decide the price?", "Browser huamua bei?", "Le navigateur décide-t-il du prix ?", "No. Server and database logic verify the package and amount; client-supplied values are not authoritative."),
        ("How do permissions work?", "Ruhusa zinafanyaje kazi?", "Comment fonctionnent les permissions ?", "Through authenticated sessions, roles, company scope, protected handlers, database grants, constraints, and RLS policies."),
        ("Is every listed module fully complete?", "Kila moduli imekamilika?", "Chaque module est-il totalement complet ?", "No. The encyclopedia labels partial, UI-only, configuration-required, external-service, and planned surfaces explicitly."),
        ("How do I contact support?", "Ninawasilianaje na huduma?", "Comment contacter le support ?", "Use the configured support/helpdesk path available to the workspace; delivery channels and contacts must be configured by the operating organization."),
    ]
    add_bilingual_table(doc, faqs, ("Question", "Kiswahili", "Français", "Verified answer"))

    add_heading(doc, "Troubleshooting and Best Practices", "Utatuzi na mbinu bora", "Dépannage et bonnes pratiques", 1)
    troubleshooting = [
        ("Catalog unavailable", "Katalogi haipatikani", "Catalogue indisponible", "Check server-side Supabase billing configuration, deployment alignment, route registration, and provider/database health. Do not expose service credentials."),
        ("Free activation rejected", "Uanzishaji wa Free umekataliwa", "Activation Free refusée", "Confirm a verified session, active company profile, manager boundary, and active FREE_15 catalog row. Do not retry with a paid code."),
        ("Checkout rejected", "Checkout imekataliwa", "Paiement refusé", "Confirm authenticated workspace, valid Tanzanian mobile number, active paid plan, Monthly cycle, and no pending payment. The server decides the amount."),
        ("External integration unavailable", "Muunganisho wa nje haupatikani", "Intégration externe indisponible", "Check environment configuration, credentials, provider status, webhook registration, and scheduled execution. Keep the issue classified as configuration/external-service dependent."),
        ("Unexpected access state", "Hali ya ruhusa haijatarajiwa", "État d’accès inattendu", "Use the server access snapshot, inspect role/company scope, review audit evidence, and fail closed rather than trusting local browser state."),
    ]
    add_bilingual_table(doc, troubleshooting, ("Problem", "Kiswahili", "Français", "Diagnosis and prevention"))
    doc.add_paragraph("Best practice is to preserve one source of truth: capture the operational fact once, validate it at the server boundary, enforce it in the database, and expose the resulting evidence through the appropriate role-aware view.")

    add_heading(doc, "Glossary and Reference Index", "Kamusi na faharasa", "Glossaire et index", 1)
    glossary = [
        ("Company scope", "Mipaka ya kampuni", "Périmètre entreprise", "The current organization boundary used to separate tenant operations."),
        ("RLS", "Usalama wa kiwango cha mstari", "Sécurité au niveau des lignes", "Database policy enforcement that restricts which rows a role can read or change."),
        ("Idempotency", "Kutokurudia madhara", "Idempotence", "A repeated request with the same key returns the existing operation rather than creating a duplicate."),
        ("FREE_15", "Kifurushi cha siku 15", "Offre de 15 jours", "Free access for 15 days at TZS 0 with no payment."),
        ("Calendar month", "Mwezi wa kalenda", "Mois calendaire", "Expiry calculated using calendar-month arithmetic rather than a fixed number of days."),
        ("SECURITY DEFINER", "Function yenye mamlaka ya mmiliki", "Fonction SECURITY DEFINER", "A database function that executes with owner privileges and therefore requires deliberate grants and body review."),
        ("RequiresPlan", "Inahitaji kifurushi", "Forfait requis", "The access state used after Free access ends or when an active package is not confirmed."),
    ]
    add_bilingual_table(doc, glossary, ("Term", "Kiswahili", "Français", "Meaning"))
    add_heading(doc, "Source and Evidence Notes", "Vyanzo na ushahidi", "Sources et preuves", 1)
    add_trilingual_block(doc,
        "This book was prepared from the supplied master prompt, the current repository source, the verified module inventory and architecture assets, the live Supabase migration ledger, and the subscription verification report. Historical migration files are preserved as immutable history; they are not treated as active runtime behavior. Where evidence was incomplete, the text uses an explicit limitation label.",
        "Kitabu hiki kimeandaliwa kutoka kwenye prompt kuu iliyotolewa, chanzo cha sasa cha repository, inventory ya moduli na assets za usanifu, migration ledger ya Supabase na ripoti ya uthibitishaji wa usajili. Migrations za zamani zimehifadhiwa kama historia isiyobadilishwa; hazichukuliwi kama runtime hai. Pale ushahidi haujakamilika, maandishi yametumia lebo ya mpaka wazi.",
        "Ce livre a été préparé à partir du prompt principal fourni, du dépôt actuel, de l’inventaire vérifié des modules et des éléments d’architecture, du registre des migrations Supabase et du rapport de vérification des abonnements. Les anciennes migrations sont conservées comme historique immuable et ne sont pas considérées comme comportement runtime actif. Lorsque les preuves étaient incomplètes, une étiquette de limite explicite est utilisée.")
    refs = [
        "[1] client/src/BusinessSphereDashboard.jsx — current authenticated workspace and module composition.",
        "[2] client/src/components/BrandLogo.tsx — official logo reference used by the application.",
        "[3] server/subscriptionBilling.ts — protected catalog, Free activation, payment, provider, and invoice handlers.",
        "[4] server/propertyManagement.ts — scheduled property control boundary.",
        "[5] supabase/migrations/20260823_062_subscription_free_plan_model.sql — FREE_15 and paid bonus-package model.",
        "[6] supabase/migrations/20260823_063_subscription_monthly_constraint_correction.sql — Monthly-only constraint correction.",
        "[7] verification/SUBSCRIPTION-API-SECURITY-MIGRATION-REVIEW-2026-08-23.md — API, RLS, Security Advisor, and migration review evidence.",
        "[8] client/public/Smart_Manager_ERP_Executive_Presentation_Inventory.pdf — verified 40-surface module inventory.",
    ]
    for ref in refs:
        doc.add_paragraph(ref)
    add_heading(doc, "SMART MANAGER — From Business Operations to Business Intelligence", "SMART MANAGER — Kutoka uendeshaji hadi akili ya biashara", "SMART MANAGER — Des opérations à l’intelligence d’affaires", 1)
    add_trilingual_block(doc,
        "SMART MANAGER is strongest when it makes work measurable, responsibility visible, and decisions easier to explain. Its long-term value is not a claim that every problem is solved automatically. It is the disciplined connection of people, records, workflows, controls, integrations, and insight so that an organization can move with greater confidence.",
        "SMART MANAGER huwa na nguvu zaidi inapofanya kazi ipimwe, uwajibikaji uonekane na maamuzi yaelezeke kwa urahisi. Thamani yake ya muda mrefu si dai kwamba kila tatizo linatatuliwa moja kwa moja. Ni kuunganisha watu, rekodi, mitiririko, udhibiti, miunganisho na uelewa ili shirika liende kwa kujiamini zaidi.",
        "SMART MANAGER est le plus utile lorsqu’il rend le travail mesurable, les responsabilités visibles et les décisions explicables. Sa valeur à long terme n’est pas de prétendre résoudre automatiquement tous les problèmes. Elle réside dans la connexion disciplinée des personnes, données, processus, contrôles, intégrations et analyses afin de progresser avec davantage de confiance.")
    doc.add_paragraph("End of official manual • English / Kiswahili / Français • Dar es Salaam, Tanzania")
    doc.save(OUT / "SMART_MANAGER_ERP_OFFICIAL_MANUAL.docx")


def typst_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("#", "\\#").replace("[", "\\[").replace("]", "\\]")


def typst_paragraph(text: str) -> str:
    return typst_escape(text)


def typst_table(rows: list[tuple[str, str, str, str]], headers: tuple[str, str, str, str]) -> str:
    cells = [f"[*{typst_escape(h)}*]" for h in headers]
    for row in rows:
        cells.extend(f"[{typst_escape(v)}]" for v in row)
    return "#table(columns: (2.4cm, 3.6cm, 3.3cm, 7.0cm), inset: 5pt, stroke: 0.35pt + luma(190), table.header(" + ", ".join(cells[:4]) + "), " + ", ".join(cells[4:]) + ")\n"


def build_typst():
    lines: list[str] = []
    lines.append('#import "report-theme.typ": report-accent, report-theme')
    lines.append('#show: report-theme.with(title: "SMART MANAGER ERP", author: "Ezra Mpapi", rhythm: "longform", running-header: true)')
    lines.append('#set text(font: ("Libertinus Serif", "Noto Serif CJK SC"), lang: "en")')
    lines.append('#set par(justify: true)')
    lines.append('')
    lines.append('#page(fill: rgb("#061A13"), margin: (top: 1.5cm, bottom: 1.5cm, x: 2cm), numbering: none, header: none)[#align(center)[#v(1cm)#image("assets/smart-manager-logo.png", width: 15cm)#v(0.6cm)#text(fill: white, size: 25pt, weight: "bold")[SMART MANAGER ERP]#v(0.3cm)#text(fill: rgb("#D6EADF"), size: 14pt)[The Official Enterprise Manual \\ and Business Operations Guide]#v(1cm)#line(length: 70%, stroke: 1pt + rgb("#D6B36A"))#v(0.8cm)#text(fill: white, size: 11pt)[Your Business. Connected. Controlled. Intelligent.]#v(1.3cm)#text(fill: white, size: 11pt)[Prepared by Ezra Mpapi \\ Owner & Creator \\ Dar es Salaam, Tanzania \\ English • Kiswahili • Français \\ 23 August 2026]]]')
    lines.append('#page(numbering: none, header: none)[#align(center)[#text(size: 20pt, weight: "bold", fill: report-accent)[About the Creator / Kuhusu Muumba / À propos du créateur]#v(0.8cm)#image("assets/ezra-mpapi-owner.png", width: 7cm)#v(0.5cm)#text(size: 11pt)[Ezra Mpapi is identified in the supplied materials as the Owner & Creator of SMART MANAGER ERP. This book does not add an unverified biography, qualification, award, or employment history. \\ \\ Ezra Mpapi ametajwa kwenye nyenzo zilizotolewa kama Mmiliki na Muumba wa SMART MANAGER ERP. Kitabu hiki hakiongezi wasifu au sifa ambazo hazijathibitishwa. \\ \\ Les documents fournis identifient Ezra Mpapi comme propriétaire et créateur de SMART MANAGER ERP. Aucune biographie non vérifiée n’est ajoutée. \\ \\ Headquarters / Makao makuu / Siège: Dar es Salaam, Tanzania]]]')
    lines.append('#page(numbering: none, header: none)[#outline(title: [Contents / Yaliyomo / Sommaire], indent: 1.5em)]')
    lines.append('#counter(page).update(1)')
    lines.append('= Editorial and Evidence Note')
    lines.append('This trilingual manual is grounded in the supplied master prompt, the current repository, verified architecture and workflow assets, the live Supabase migration ledger, and the subscription verification report. English is the primary reference language; Kiswahili and French accompany every major chapter. Technical routes, table names, package codes, and status labels remain in their verified project form.')
    lines.append('*Editorial note / Dokezo la uhariri / Note éditoriale:* A status such as PLANNED, UI ONLY, PARTIALLY IMPLEMENTED, CONFIGURATION REQUIRED, or EXTERNAL SERVICE REQUIRED is an evidence boundary, not a negative judgment and not a current feature promise.')
    lines.append('')
    for i, chapter in enumerate(chapters, 1):
        lines.append(f'= Part {i}: {typst_escape(chapter["title"])}')
        lines.append(f'*{typst_escape(chapter["sw"])}*  \\  *{typst_escape(chapter["fr"])}*')
        lines.append(typst_paragraph(chapter["en"]))
        lines.append(typst_paragraph(chapter["sw_text"]))
        lines.append(typst_paragraph(chapter["fr_text"]))
        if chapter["points"]:
            lines.append(typst_table(chapter["points"], ("English focus", "Kiswahili", "Français", "Verified guidance")))
        title = chapter["title"]
        figure_map = {
            "The Story Behind SMART MANAGER": [("01-live-homepage.png", "Figure 1 — Existing SMART MANAGER public homepage evidence.")],
            "The Product Tour and Operating Model": [("03-auth-workflow.png", "Figure 2 — Authentication and secure onboarding workflow evidence."), ("02-live-module-map.png", "Figure 3 — Live module map evidence.")],
            "Connected Workflows": [("04-business-workflow.png", "Figure 4 — Connected business workflow evidence."), ("09-sales-to-receipt.png", "Figure 5 — Sales-to-receipt workflow evidence."), ("10-inventory-movement.png", "Figure 6 — Inventory movement workflow evidence.")],
            "Subscriptions and Customer Value": [("08-webhook-state-machine.png", "Figure 7 — Verified payment/webhook state-machine evidence.")],
            "Security, Administration, and Trust": [("06-auth-tenancy.png", "Figure 8 — Authentication and tenancy evidence."), ("07-financial-control.png", "Figure 9 — Financial control-flow evidence.")],
            "Current Limitations and the Future": [("05-deployment-observation.png", "Figure 10 — Deployment observation evidence.")],
        }
        for filename, caption in figure_map.get(title, []):
            lines.append(f'#figure(image("assets/figures/{filename}", width: 15cm), caption: [{typst_escape(caption)}])')
        if title == "Subscriptions and Customer Value":
            lines.append(typst_table([
                ("FREE_15", "FREE", "FREE", "TZS 0; 15 days; no payment."),
                ("TWIGA", "TWIGA", "TWIGA", "TZS 5,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("TEMBO", "TEMBO", "TEMBO", "TZS 10,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("SIMBA", "SIMBA", "SIMBA", "TZS 15,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("SIMBA_SC", "SIMBA SC", "SIMBA SC", "TZS 4,500/month; 1 paid + 1 bonus; 2 calendar months."),
                ("YANGA_SC", "YANGA SC", "YANGA SC", "TZS 9,000/month; 1 paid + 1 bonus; 2 calendar months."),
                ("AZAM_FC", "AZAM FC", "AZAM FC", "TZS 7,000/month; 1 paid + 1 bonus; 2 calendar months."),
            ], ("Code", "Name", "Nom", "Verified commercial contract")))
    lines.append('= Complete Module Encyclopedia')
    lines.append('The inventory below is grounded in the project module registry, source files, migration names, and tested persistence boundaries. Status labels distinguish evidence from aspiration.')
    lines.append(typst_table([(a, b, c, f"{d}: {e}") for a, b, c, d, e in modules], ("Module / English", "Kiswahili", "Français", "Status and verified evidence")))
    lines.append('= Workflow Reference')
    workflow_rows = [
        ("Authentication", "Uthibitishaji", "Authentification", "Public entry → verified session → company/profile context → protected workspace."),
        ("Company onboarding", "Uanzishaji wa kampuni", "Intégration entreprise", "Register → verify → create/join company → authenticated FREE_15 activation → workspace."),
        ("Sales", "Mauzo", "Ventes", "Customer → document → validation → stock/finance → payment/receipt → report/audit."),
        ("Procurement", "Ununuzi", "Achats", "Supplier → order → receipt → stock → payable/finance → report."),
        ("POS", "POS", "PDV", "Register → item/customer → validation → transaction → receipt/return/credit → reconciliation."),
        ("Property", "Mali isiyohamishika", "Immobilier", "Portfolio → company-scoped records → scheduled controls → tenant-safe reporting."),
        ("Subscription", "Usajili", "Abonnement", "Catalog → access → verified payment intent → provider status → payment/subscription/invoice/audit."),
        ("Administration", "Usimamizi", "Administration", "Platform control → scoped action → audit evidence → support/recovery."),
    ]
    lines.append(typst_table(workflow_rows, ("Workflow", "Kiswahili", "Français", "Verified path")))
    lines.append('= Business Problem to Solution')
    lines.append(typst_table([
        ("Fragmented records", "Rekodi zilizotawanyika", "Données fragmentées", "Connected workspace and shared management views."),
        ("Poor inventory visibility", "Mwonekano mdogo wa bidhaa", "Stocks peu visibles", "Inventory movement linked to sales, procurement, POS, and reports where enabled."),
        ("Uncontrolled access", "Ufikiaji usiodhibitiwa", "Accès non contrôlé", "Sessions, roles, company scope, RLS, policies, and audit."),
        ("Payment uncertainty", "Kutokuwa na uhakika wa malipo", "Paiement incertain", "Server amount verification, provider checks, idempotency, invoice, and audit."),
        ("Manual reporting", "Ripoti za mikono", "Rapports manuels", "Dashboards, report services, schedules, and exports where configured."),
    ], ("Business struggle", "Kiswahili", "Français", "Verified response")))
    lines.append('= FAQ, Troubleshooting, and Best Practices')
    lines.append('*Can a small business use SMART MANAGER? / Je, biashara ndogo inaweza kuitumia? / Une petite entreprise peut-elle l’utiliser ?* Yes, by starting with relevant verified core surfaces and adopting additional controls as the business grows.')
    lines.append('*How does subscription work? / Usajili unafanyaje kazi? / Comment fonctionne l’abonnement ?* FREE_15 provides 15 days at TZS 0. Paid packages are monthly-only with one paid month and one bonus calendar month. Server and database logic, not browser values, decide price and entitlement.')
    lines.append('*What is the safest troubleshooting rule? / Kanuni salama ya utatuzi ni ipi? / Quelle est la règle de dépannage la plus sûre ?* Confirm the verified session, company scope, role, configuration, provider health, and audit evidence; fail closed rather than trusting a local browser state.')
    lines.append('= Glossary and Reference Notes')
    lines.append(typst_table([
        ("Company scope", "Mipaka ya kampuni", "Périmètre entreprise", "The organization boundary used to separate tenant operations."),
        ("RLS", "Usalama wa kiwango cha mstari", "Sécurité au niveau des lignes", "Database policy enforcement over rows."),
        ("Idempotency", "Kutokurudia madhara", "Idempotence", "Repeated request keys do not create duplicate operations."),
        ("FREE_15", "Kifurushi cha siku 15", "Offre de 15 jours", "Free access for 15 days at TZS 0."),
        ("RequiresPlan", "Inahitaji kifurushi", "Forfait requis", "Access state requiring a confirmed package."),
    ], ("Term", "Kiswahili", "Français", "Meaning")))
    lines.append('= SMART MANAGER — From Business Operations to Business Intelligence')
    lines.append('SMART MANAGER is strongest when it makes work measurable, responsibility visible, and decisions easier to explain. Its value is the disciplined connection of people, records, workflows, controls, integrations, and insight—not a claim that every problem is solved automatically.')
    lines.append('SMART MANAGER huwa na nguvu zaidi inapofanya kazi ipimwe, uwajibikaji uonekane na maamuzi yaelezeke. Thamani yake ni kuunganisha watu, rekodi, mitiririko, udhibiti, miunganisho na uelewa kwa nidhamu—si dai kwamba kila tatizo linatatuliwa kiotomatiki.')
    lines.append('SMART MANAGER est le plus utile lorsqu’il rend le travail mesurable, les responsabilités visibles et les décisions explicables. Sa valeur est la connexion disciplinée des personnes, données, processus, contrôles, intégrations et analyses—et non la promesse de résoudre automatiquement chaque problème.')
    lines.append('')
    document = '\n\n'.join(lines).rstrip() + '\n'
    document = document.replace('image("assets/smart-manager-logo.png"', f'image("{LOGO.as_posix()}"')
    document = document.replace('image("assets/ezra-mpapi-owner.png"', f'image("{OWNER.as_posix()}"')
    TYPST.joinpath('main.typ').write_text(document, encoding='utf-8')


if __name__ == "__main__":
    if "--typst-only" in sys.argv:
        build_typst()
        print(f"Created {TYPST / 'main.typ'}")
    else:
        build_docx()
        build_typst()
        print(f"Created {OUT / 'SMART_MANAGER_ERP_OFFICIAL_MANUAL.docx'}")
        print(f"Created {TYPST / 'main.typ'}")
