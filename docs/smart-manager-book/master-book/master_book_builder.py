from __future__ import annotations

import json
import os
import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
REPO = ROOT.parents[2]
TYPST_PROJECT = ROOT / "typst-project"
ASSETS = TYPST_PROJECT / "assets"
DELIVERABLES = ROOT / "deliverables"
SOURCE_MD = ROOT / "SMART_MANAGER_MASTER_BOOK_EN_SW.md"
DOCX_PATH = DELIVERABLES / "SMART_MANAGER_MASTER_BOOK_EN_SW.docx"
TYPOUT = TYPST_PROJECT / "main.typ"
EVIDENCE = ROOT / "evidence"
TODAY = "24 August 2026"
MANAGED_BRANDING_ASSET_ROOT = Path(
    os.environ.get(
        "SMART_MANAGER_MASTER_BOOK_BRANDING_ROOT",
        "/home/ubuntu/webdev-static-assets/smart-manager-master-book/branding",
    )
).expanduser()
MASTER_BOOK_LOGO = MANAGED_BRANDING_ASSET_ROOT / "smart-manager-logo.png"
MASTER_BOOK_LOGO_STORAGE = "/manus-storage/smart-manager-logo_b1db8065.png"
MASTER_BOOK_OWNER = MANAGED_BRANDING_ASSET_ROOT / "ezra-mpapi-owner.png"

LIVE_TABLES = EVIDENCE / "live_supabase_tables_2026-08-24.json"
LIVE_MIGRATIONS = EVIDENCE / "live_supabase_migrations_2026-08-24.json"
SECURITY_ADVISORS = EVIDENCE / "live_supabase_security_advisors_2026-08-24.json"
PERF_ADVISORS = EVIDENCE / "live_supabase_performance_advisors_2026-08-24.json"

MODULES: list[dict[str, Any]] = [
    {"name":"Public Brand and Marketing Entry","sw":"Ukurasa wa Umma wa Chapa na Masoko","status":"IMPLEMENTED","ui":"client/src/pages/Home.tsx","server":"server/_core/apiApp.ts","prefixes":["companies"],"roles":"Public visitor; prospective customer","focus":"The public landing page communicates the Smart Manager proposition, links to the application, exposes language and theme controls, and provides a passkey entry point without treating the marketing preview as operational data.","workflow":"Discover product → review capabilities → choose launch or authentication → enter the secure workspace.","limit":"Marketing delivery and conversion analytics are not the same as a signed-in company workflow; external campaign channels require configuration.","sw_focus":"Ukurasa wa mwanzo unaeleza thamani ya Smart Manager, unaunganisha mgeni na kuingia kwenye mfumo, na unaonyesha kwa uaminifu kwamba mwonekano wa masoko si rekodi ya biashara ya kampuni.","image":"screenshots/live-homepage.png"},
    {"name":"Authentication and Secure Onboarding","sw":"Uthibitishaji na Uanzishaji Salama","status":"IMPLEMENTED","ui":"client/src/components/PublicAuthGateway.jsx; client/src/App.tsx","server":"server/_core/apiApp.ts; server/_core/oauth.ts; server/authHeaders.ts","prefixes":["profiles","companies","company_memberships"],"roles":"Unauthenticated visitor; verified user; company owner","focus":"The public gateway supports password sign-in, password recovery, email confirmation, approved OAuth providers, passkeys, session persistence, and the transition into a company-aware workspace.","workflow":"Load public Supabase configuration → authenticate or recover → persist token/session → resolve verified profile → enter /app.","limit":"Actual OAuth delivery, email delivery, and passkey availability remain dependent on Supabase/provider configuration and browser capability.","sw_focus":"Gateway ya umma inasaidia kuingia kwa nenosiri, kurejesha nenosiri, kuthibitisha barua pepe, OAuth iliyokubaliwa, passkeys na uhifadhi wa kikao kabla ya kuingia kwenye nafasi ya kampuni.","image":"screenshots/live-workspace-entry.png"},
    {"name":"Master Application Shell and Navigation","sw":"Ganda Kuu la Programu na Urambazaji","status":"IMPLEMENTED","ui":"client/src/App.tsx; client/src/BusinessSphereDashboard.jsx","server":"server/routers.ts; server/_core/apiApp.ts","prefixes":["company_modules","workspaces"],"roles":"All authenticated users, filtered by role and entitlement","focus":"The application shell provides route handling, error boundaries, theme and language providers, dashboard preferences, lazy workspace loading, command navigation, and fail-closed subscription access checks.","workflow":"Route match → auth decision → lazy workspace load → role/entitlement-aware navigation → module action.","limit":"The dashboard remains a large legacy single-file boundary and produces a non-fatal bundle-size warning; safe decomposition is a future engineering workstream.","sw_focus":"Ganda la programu linaunganisha routes, uthibitishaji, lugha, mandhari, mapendeleo ya dashibodi, lazy loading na urambazaji unaozingatia jukumu pamoja na entitlement.","image":"diagrams/live-module-map.png"},
    {"name":"Profile Identity Center","sw":"Kituo cha Utambulisho wa Wasifu","status":"IMPLEMENTED","ui":"client/src/components/ProfileIdentityCenter.jsx","server":"server/profileIdentity.ts; server/routers.ts","prefixes":["profiles","user_table_preferences"],"roles":"Authenticated user; workspace administrator for linked context","focus":"The profile center separates personal identity, work identity, security, preferences, notifications, and activity while keeping protected identity fields server-authoritative.","workflow":"Open profile → read verified profile → update permitted self-service fields → upload or remove avatar through controlled storage → refresh confirmed state.","limit":"Device/session listing and workspace switching are explicitly unavailable unless supported by a verified backend contract.","sw_focus":"Kituo cha wasifu hutenganisha utambulisho binafsi, kazi, usalama, mapendeleo, arifa na shughuli. Mtumiaji hubadilisha tu sehemu anazoruhusiwa na server ndiyo mamlaka ya mwisho.","image":"diagrams/live-auth-workflow.png"},
    {"name":"Executive Dashboard","sw":"Dashibodi ya Uongozi","status":"IMPLEMENTED","ui":"client/src/components/ExecutiveCommandCenter.jsx; client/src/BusinessSphereDashboard.jsx","server":"server/dashboardReports.ts; server/routers.ts","prefixes":["custom_kpis","scheduled_reports","financial_benchmarks"],"roles":"CEO; COO; CFO; organization owner; manager roles","focus":"The executive surface consolidates operational signals from connected modules. It is intended for review and prioritization, not as an independent source of financial truth.","workflow":"Select company context → load confirmed metrics → inspect trends/exceptions → navigate to source module → approve or act through the protected boundary.","limit":"A dashboard is only as complete as the connected records, schedules, and integrations that feed it; empty or unavailable states must remain visible.","sw_focus":"Dashibodi ya uongozi hukusanya ishara za uendeshaji kutoka moduli zilizounganishwa. Ni ya mapitio na kupanga kipaumbele, si chanzo tofauti cha ukweli wa fedha.","image":"figures/01-live-homepage.png"},
    {"name":"CRM and Customer Pipeline","sw":"CRM na Mfuatano wa Wateja","status":"TESTING","ui":"client/src/components/CommercialCommandCenters.jsx; client/src/BusinessSphereDashboard.jsx","server":"server/routers.ts; server/salesInteractions.ts","prefixes":["crm","customer"],"roles":"Sales Manager; sales user; account owner; support user","focus":"CRM organizes contacts, leads, interactions, customer feedback, and pipeline signals so commercial activity can be related to sales and support records.","workflow":"Create or import lead → qualify → record interaction → convert or follow up → connect to quotation/invoice/support.","limit":"The repository contains strong persistence and truthfulness contracts, but broader authenticated CRUD walkthroughs remain required for every legacy screen.","sw_focus":"CRM huweka pamoja mawasiliano, leads, mwingiliano na mrejesho wa wateja. Lengo ni kuunganisha mazungumzo ya kibiashara na mauzo pamoja na huduma kwa mteja.","image":"figures/04-business-workflow.png"},
    {"name":"Sales and Billing","sw":"Mauzo na Utozaji","status":"TESTING","ui":"client/src/components/SalesDetailWorkspace.jsx; client/src/components/CommercialCommandCenters.jsx","server":"server/routers.ts; server/subscriptionBilling.ts","prefixes":["sales","ecommerce_orders","ecommerce_products"],"roles":"Sales Manager; finance manager; sales user; billing manager","focus":"Sales surfaces cover quotations, invoices, subscriptions, payments, returns, and customer interactions. The customer-facing amount and status must come from confirmed server/database results.","workflow":"Customer → quotation/order → invoice → payment or collection → receipt/ledger/report → follow-up.","limit":"Legacy invoice payment posting still requires an atomic database RPC and durable idempotency constraint before concurrent-safe payment posting can be claimed.","sw_focus":"Sehemu ya mauzo inahusu quotations, ankara, subscriptions, malipo, returns na mwingiliano wa wateja. Bei na hali ya mwisho vinatoka server na database, si thamani ya browser.","image":"figures/09-sales-to-receipt.png"},
    {"name":"Point of Sale","sw":"Mauzo ya Moja kwa Moja (POS)","status":"PASSED","ui":"client/src/components/OperationsCommandCenters.jsx; client/src/BusinessSphereDashboard.jsx","server":"server/posWorkforceRpcAdapters.ts; server/posTransactionEngine.ts","prefixes":["pos"],"roles":"Cashier; finance manager; warehouse manager; administrator","focus":"POS supports registers, terminals, shifts, sale headers and lines, tenders, returns, tax, promotions, loyalty, pending offline transport, reconciliation, and audit paths.","workflow":"Open register/shift → scan or select item → validate price and stock → collect tender → confirm sale → receipt → reconcile or return.","limit":"Controlled staging credentials and real device/provider behavior remain operational acceptance requirements; client pending queues are not treated as durable server saves.","sw_focus":"POS inasimamia register, shift, bidhaa, malipo, returns, kodi, promotions, loyalty, reconciliation na audit. Mstari wa usalama ni kuthibitisha muamala kabla ya kuuita umehifadhiwa.","image":"diagrams/financial-control-flow.png"},
    {"name":"Inventory and Warehouse Management","sw":"Usimamizi wa Hesabu na Maghala","status":"TESTING","ui":"client/src/components/OperationsCommandCenters.jsx","server":"server/supabasePersistence.ts; server/routers.ts","prefixes":["inventory","stock","warehouse"],"roles":"Warehouse Manager; procurement officer; sales user; administrator","focus":"Inventory records cover items, batches, stock movement, suppliers, warehouses, transfers, audits, and the operational connection from stock to sales and procurement.","workflow":"Create item → receive or adjust stock → transfer or sell → record movement → inspect balance/audit → replenish.","limit":"Full authenticated import/export and concurrency walkthroughs remain part of the broader acceptance backlog.","sw_focus":"Hesabu na maghala huunganisha items, batches, movement, suppliers, warehouses, transfers na audits. Kila mabadiliko ya quantity inapaswa kuonekana kama movement yenye ushahidi.","image":"figures/10-inventory-movement.png"},
    {"name":"Procurement and Vendor Management","sw":"Ununuzi na Wasambazaji","status":"TESTING","ui":"client/src/components/OperationsCommandCenters.jsx","server":"server/routers.ts; server/procurementPersistence.ts","prefixes":["procurement","purchase_order"],"roles":"Procurement Officer; warehouse manager; finance manager; administrator","focus":"Procurement relates suppliers, purchase orders, contracts, receiving, stock, approvals, and payables evidence. The workflow must preserve what was requested, approved, received, and owed.","workflow":"Supplier → request/order → approval → receipt → stock update → payable/finance evidence → report.","limit":"Downstream integrations and external supplier channels are configuration-dependent; recorded integration evidence must not be replaced with fabricated acknowledgements.","sw_focus":"Ununuzi unaunganisha wasambazaji, purchase orders, contracts, mapokezi, stock, approvals na madeni. Mfumo unapaswa kutenganisha kilichoombwa, kilichoidhinishwa, kilichopokelewa na kinachodaiwa.","image":"figures/04-business-workflow.png"},
    {"name":"Finance and Accounting","sw":"Fedha na Uhasibu","status":"TESTING","ui":"client/src/components/FinanceCommandCenters.jsx","server":"server/financeCommandCenters.ts; server/financePersistence.ts","prefixes":["finance","fin","journal","expense","period"],"roles":"CFO; Finance Manager; accountant; organization owner","focus":"Finance includes accounts, journals, income and expenses, budgets, tax/VAT, period closes, bank reconciliation, assets, and management reports. It is the control plane for monetary evidence rather than a decorative dashboard.","workflow":"Capture transaction → validate company scope → post or approve → reconcile → close period → report.","limit":"Production reporting depends on intended schema alignment and real company data; tax filing itself is not implied by a local summary.","sw_focus":"Fedha na uhasibu vina accounts, journal, mapato na matumizi, bajeti, kodi/VAT, kufunga vipindi, reconciliation ya benki, assets na ripoti. Data halisi ndiyo msingi wa ripoti.","image":"figures/07-financial-control.png"},
    {"name":"Reports and Scheduled Reporting","sw":"Ripoti na Ripoti Zilizopangwa","status":"TESTING","ui":"client/src/components/FinanceCommandCenters.jsx; client/src/BusinessSphereDashboard.jsx","server":"server/dashboardReports.ts; server/reportSchedules.ts","prefixes":["scheduled_reports","report"],"roles":"Executive roles; manager; administrator","focus":"Reporting provides trend views, exports, schedules, delivery states, and compliance-oriented summaries. A generated report is evidence of recorded data and configured delivery, not a guarantee that an external message was delivered.","workflow":"Choose report → apply company/time scope → compute from source records → export or schedule → inspect delivery history.","limit":"Email delivery remains provider-configured; CSV is the dependency-safe export path when spreadsheet dependencies are intentionally absent.","sw_focus":"Ripoti huonyesha mwenendo, exports, schedules na hali ya delivery. Ripoti inaonyesha data iliyorekodiwa; haimaanishi moja kwa moja kwamba email ya nje imefika.","image":"figures/05-deployment-observation.png"},
    {"name":"Human Resources and Payroll","sw":"Rasilimali Watu na Mishahara","status":"TESTING","ui":"client/src/components/PeopleCommandCenters.jsx; client/src/components/EmployeePortalWorkspace.jsx","server":"server/teamWorkforce.ts; server/tanzaniaPayroll.ts; server/routers.ts","prefixes":["hr","workforce","employee"],"roles":"HR Manager; manager; employee; payroll approver; administrator","focus":"HR covers employees, attendance, benefits, leave, goals, approvals, payroll, announcements, and secure invitations. Tenant-scoped persistence replaces browser-only identity and invite paths.","workflow":"Invite employee → accept securely → assign role → record attendance/leave → approve → calculate payroll → report.","limit":"Payroll and employee records need a full authenticated tenant-data walkthrough; external salary/payment rails are outside the repository’s verified contract.","sw_focus":"HR inahusu wafanyakazi, mahudhurio, benefits, leave, malengo, approvals, payroll, announcements na invitations salama. Data ya mfanyakazi inapaswa kubaki ndani ya kampuni yake.","image":"diagrams/live-business-workflow.png"},
    {"name":"Manufacturing and Work Orders","sw":"Uzalishaji na Maagizo ya Kazi","status":"PARTIALLY IMPLEMENTED","ui":"client/src/BusinessSphereDashboard.jsx; client/src/components/OperationsCommandCenters.jsx","server":"server/manufacturingPersistence.ts","prefixes":["manufacturing"],"roles":"Production manager; warehouse manager; project manager","focus":"The repository contains manufacturing and work-order persistence boundaries and navigation signals. The master book treats operational depth as partial unless a complete deployment-backed flow is demonstrated.","workflow":"Define work order → allocate or consume material → update status → record output → review cost and exceptions.","limit":"The visible surface is broader than the independently verified end-to-end contract for every manufacturing scenario.","sw_focus":"Repository ina mipaka ya manufacturing na work orders, lakini kitabu kinatenganisha uwepo wa surface na uthibitisho wa workflow kamili wa uzalishaji.","image":"figures/04-business-workflow.png"},
    {"name":"Supply Chain and Fleet","sw":"Mnyororo wa Ugavi na Magari","status":"TESTING","ui":"client/src/components/OperationsCommandCenters.jsx; client/src/components/FleetWorkspace.jsx","server":"server/fleetManagement.ts","prefixes":["scm","fleet","flt"],"roles":"Fleet manager; logistics user; administrator","focus":"Supply-chain and fleet surfaces connect shipments, vehicles, trips, drivers, fuel, maintenance, incidents, routes, telematics events, and alerts.","workflow":"Plan route/trip → assign driver/vehicle → record movement or fuel → ingest event → inspect alert → close maintenance or incident.","limit":"Telematics and external vehicle providers require approved webhook configuration and secrets; no external movement is claimed from a local screen alone.","sw_focus":"Ugavi na fleet huunganisha shipments, magari, trips, madereva, mafuta, matengenezo, incidents, routes, telematics na alerts.","image":"figures/04-business-workflow.png"},
    {"name":"Marketing Campaigns","sw":"Kampeni za Masoko","status":"TESTING","ui":"client/src/components/CommercialCommandCenters.jsx","server":"server/routers.ts; server/emailTemplateWorkflow.ts","prefixes":["marketing","emails","crm"],"roles":"CMO; marketing manager; sales manager","focus":"Marketing provides campaign-oriented surfaces, customer segmentation signals, message templates, and analytics boundaries. Delivery is kept separate from the platform’s verified persistence contract.","workflow":"Define audience → draft message → validate template/link → obtain approval → dispatch through configured channel → measure response.","limit":"External email, SMS, WhatsApp, and campaign providers remain configuration-dependent; the UI must not claim dispatch when only a draft exists.","sw_focus":"Masoko huandaa segments, templates, ujumbe na analytics. Ujumbe wa draft unatenganishwa na dispatch halisi ya huduma ya nje.","image":"figures/01-live-homepage.png"},
    {"name":"E-Commerce Storefront","sw":"Duka la Kielektroniki","status":"PARTIALLY IMPLEMENTED","ui":"client/src/components/CommercialCommandCenters.jsx","server":"server/routers.ts","prefixes":["ecommerce"],"roles":"Store manager; sales manager; customer","focus":"The repository names e-commerce products and orders and includes storefront-oriented surfaces. The evidence boundary does not invent a live Shopify or third-party commerce activation.","workflow":"Publish catalog → receive order → validate customer/order → connect fulfilment and finance → report.","limit":"External commerce provider credentials and production storefront activation are not verified by this book.","sw_focus":"Mfumo una surfaces za bidhaa na oda za e-commerce, lakini haujadai kuwa Shopify au mtoa huduma mwingine amewezeshwa bila ushahidi wa configuration.","image":"figures/04-business-workflow.png"},
    {"name":"Documents and Secure Files","sw":"Nyaraka na Mafaili Salama","status":"TESTING","ui":"client/src/components/PeopleCommandCenters.jsx; client/src/components/ProfileIdentityCenter.jsx","server":"server/storage.ts; server/documentDownloadBoundaries.ts","prefixes":["documents","digital_signatures","signatures","approval_signatures"],"roles":"All authenticated users within permitted company scope; approver","focus":"Document records, storage references, signatures, exports, and download boundaries keep file bytes separate from transactional metadata and preserve tenant scope.","workflow":"Create metadata → upload through storage boundary → verify reference → share or sign under role control → download/export evidence.","limit":"S3/storage configuration is required for live file bytes; metadata persistence alone is not evidence that a file was uploaded successfully.","sw_focus":"Nyaraka huhifadhi metadata na reference ya storage, si file bytes ndani ya business table. Upload, kusaini na download vinahitaji mipaka ya ruhusa.","image":"figures/06-auth-tenancy.png"},
    {"name":"Projects and Task Management","sw":"Miradi na Usimamizi wa Kazi","status":"PARTIALLY IMPLEMENTED","ui":"client/src/components/PeopleCommandCenters.jsx","server":"server/routers.ts; server/projectsPersistence.ts","prefixes":["projects","project"],"roles":"Project Manager; team member; executive","focus":"Projects provide tasks, milestones, expenses, and progress-oriented persistence boundaries so work can be coordinated with financial and operational context.","workflow":"Create project → add milestone/task → assign owner → update progress → record expense → review completion.","limit":"The repository contains the persistence boundary, but complete project operations still depend on authenticated data and role-specific acceptance.","sw_focus":"Miradi huunganisha tasks, milestones, expenses na maendeleo. Uthibitisho wa kina wa kila operation unahitaji data halisi ya kampuni na majukumu sahihi.","image":"figures/04-business-workflow.png"},
    {"name":"Customer Support and Helpdesk","sw":"Huduma kwa Wateja na Helpdesk","status":"TESTING","ui":"client/src/components/IntelligenceCommandCenters.jsx; client/src/components/PeopleCommandCenters.jsx","server":"server/supportOperations.ts; server/supportMetrics.ts","prefixes":["support"],"roles":"Support agent; manager; customer; administrator","focus":"Support includes tickets, timelines, internal notes, inbox/chat evidence, calls, SLA policies, workflow rules, search, and support metrics.","workflow":"Receive issue → create ticket → classify/SLA → assign → communicate → resolve → audit and measure.","limit":"Automated external support execution is intentionally not claimed without configured channels; internal ticket persistence remains the source of record.","sw_focus":"Helpdesk ina tickets, timelines, internal notes, chat/call evidence, SLA, assignment, resolution na metrics. Dispatch ya nje inategemea channel iliyosanidiwa.","image":"figures/04-business-workflow.png"},
    {"name":"Enterprise Analytics and BI","sw":"Uchambuzi wa Biashara na BI","status":"PARTIALLY IMPLEMENTED","ui":"client/src/components/IntelligenceCommandCenters.jsx; client/src/components/PredictiveAnalyticsWorkspace.jsx","server":"server/dashboardReports.ts; server/marketIntelligence.ts","prefixes":["custom_kpis","financial_benchmarks","competitors"],"roles":"Executive; analyst; manager","focus":"Analytics combines KPI, benchmark, market, trend, and command-center views. It is a decision-support layer over recorded operations, not a replacement for source ledgers.","workflow":"Select scope → read source metrics → compare or trend → investigate source record → document action.","limit":"Analytical quality depends on complete, correctly scoped input data and configured market services.","sw_focus":"BI huleta KPI, benchmarks, market na trend pamoja. Ubora wa uchambuzi unategemea data iliyokamilika na iliyowekwa ndani ya kampuni sahihi.","image":"figures/01-live-homepage.png"},
    {"name":"Notifications and Alerting","sw":"Arifa na Tahadhari","status":"TESTING","ui":"client/src/BusinessSphereDashboard.jsx; module workspaces","server":"server/notificationHistory.ts; server/transactionalEmail.ts","prefixes":["notification","sms","emails"],"roles":"Authenticated users; managers; administrators","focus":"Notifications capture reminders, delivery history, unread/read state, scheduled signals, and module-specific alerts while distinguishing a persisted notification from successful external delivery.","workflow":"Detect event → create notification record → deliver through configured channel → mark read or retry → audit result.","limit":"SMS, email, push, and WhatsApp delivery depend on external provider configuration and should show explicit unavailable states.","sw_focus":"Arifa zinarekodi tukio, delivery history, hali ya kusoma na retries. Rekodi ya notification si uthibitisho wa delivery ya mtoa huduma wa nje.","image":"figures/05-deployment-observation.png"},
    {"name":"Activity Stream and Audit Evidence","sw":"Mtiririko wa Shughuli na Ushahidi wa Ukaguzi","status":"TESTING","ui":"client/src/components/ComplianceAuditLogView.tsx; dashboard views","server":"server/auditLogs.ts; server/tenantAuditViewer.ts","prefixes":["audit_log","billing_plan_audit_log","community_group_audit_log","hospitality_audit_log","bank_audit_events"],"roles":"Auditor; administrator; company manager with scope","focus":"Audit evidence records who performed an action, where applicable, in which company scope, and with what result. Sensitive global administration records have a separate control boundary.","workflow":"Authenticated action → protected mutation → audit record → scoped search/export → review or incident follow-up.","limit":"An audit record increases traceability but does not correct an unsafe write; mutation confirmation and database constraints remain primary controls.","sw_focus":"Ushahidi wa audit unaonyesha nani alifanya kitendo, kwenye kampuni ipi na matokeo yake pale taarifa hiyo ipo. Rekodi za global admin zinalindwa tofauti.","image":"figures/06-auth-tenancy.png"},
    {"name":"Integration Hub","sw":"Kitovu cha Miunganisho","status":"CONFIGURATION REQUIRED","ui":"client/src/components/FinanceCommandCenters.jsx; client/src/components/PeopleCommandCenters.jsx","server":"server/webhooks.ts; server/transactionalEmail.ts; server/integration connections","prefixes":["integration","webhooks","notification_channels"],"roles":"Administrator; integration manager; finance manager","focus":"The hub represents integration metadata, webhook configuration, provider readiness, delivery history, and safe test boundaries.","workflow":"Define connection → validate configuration → test safely → enable dispatch → monitor delivery → retry or disable.","limit":"No provider credential, endpoint, or delivery success is assumed when configuration is missing; controls must remain disabled or clearly unavailable.","sw_focus":"Kitovu cha miunganisho huonyesha metadata, webhooks, readiness na delivery history. Bila configuration halali, mfumo haupaswi kuonyesha kana kwamba channel inafanya kazi.","image":"diagrams/integration-topology.png"},
    {"name":"Workflow Studio and Marketplace","sw":"Studio ya Mitiririko na Soko","status":"PARTIALLY IMPLEMENTED","ui":"client/src/components/PeopleCommandCenters.jsx; dashboard registry","server":"server/workflows.ts; server/aiApprovals.ts","prefixes":["workflows","workflow_marketplace_templates","workflow_tasks"],"roles":"Workflow administrator; approver; manager","focus":"Workflow Studio represents reusable templates, steps, approvals, tasks, notifications, and audit actions. Official templates state what they do and what they do not do.","workflow":"Choose template → inspect steps → configure target → approve or run → record task/result → review audit.","limit":"A workflow step that drafts an email or notification is not proof of external dispatch; complete orchestration requires channel configuration.","sw_focus":"Workflow Studio huandaa templates, steps, approvals, tasks, notifications na audit. Draft ya ujumbe si uthibitisho wa kutumwa kwa mteja.","image":"figures/04-business-workflow.png"},
    {"name":"Collaboration Hub","sw":"Kitovu cha Ushirikiano","status":"PARTIALLY IMPLEMENTED","ui":"client/src/components/PeopleCommandCenters.jsx","server":"server/collaborationEmailLinkCheck.ts; server/collaborationPersistence.ts","prefixes":["collab","calendar","emails"],"roles":"Team member; manager; executive","focus":"Collaboration connects messages, channels, presence, calendar, read receipts, and workflow signals where configured.","workflow":"Create channel/event → post message → notify participants → track read/presence → retain evidence.","limit":"External collaboration delivery and realtime behavior depend on configuration and provider availability; local state is not a durable message authority.","sw_focus":"Ushirikiano huunganisha channels, messages, presence, calendar na read receipts pale huduma hizo zimesanidiwa.","image":"figures/04-business-workflow.png"},
    {"name":"TRA, VFD, and Tanzania Fiscalization","sw":"TRA, VFD na Fiscalization ya Tanzania","status":"CONFIGURATION REQUIRED","ui":"client/src/components/TraPortalModule.jsx","server":"server/traFiscalRouter.ts; server/traFiscal.ts; server/traZReportArchive.ts","prefixes":["tax","fiscal","tra","z_report"],"roles":"CFO; Finance Manager; administrator; fiscal operator","focus":"The repository contains fiscal profiles, receipts, retry queues, Z-report archives, tax configuration, VAT anomaly analysis, and idempotent submission boundaries. It distinguishes configuration and provider readiness from an actual TRA submission.","workflow":"Configure TIN/VRN and branch → validate readiness → submit idempotent transaction → receive provider result → store receipt/retry/audit evidence.","limit":"Production fiscalization requires approved provider credentials, device/branch configuration, and provider availability; a local receipt screen is not a TRA acceptance claim.","sw_focus":"Mifumo ya fiscal ina profiles, receipts, retry queue, Z-reports na tax settings. Uwasilishaji halisi wa TRA unahitaji TIN/VRN, configuration ya branch, credentials na mtoa huduma.","image":"figures/08-webhook-state-machine.png"},
    {"name":"AI Assistant and Smart Intelligence","sw":"Msaidizi wa AI na Akili ya Biashara","status":"EXTERNAL SERVICE REQUIRED","ui":"client/src/components/AIChatBox.tsx; client/src/components/IntelligenceCommandCenters.jsx","server":"server/smartAssistant.ts; server/aiApprovals.ts","prefixes":["ai","custom_kpis"],"roles":"Authenticated user; approver for proposed actions","focus":"The AI assistant uses a built-in model boundary, limits history and context, treats user/business JSON as untrusted data, returns structured responses, restricts navigation targets, and produces proposals rather than silently mutating business records.","workflow":"Verified user → scoped context → structured assistant response → optional navigation/proposal → independent role approval → confirmed business action.","limit":"AI model availability and response quality depend on configured service access and supplied evidence; the assistant must not claim a mutation it did not execute.","sw_focus":"Msaidizi wa AI hupokea context yenye mipaka, hutumia response iliyopangwa, na huandaa proposal badala ya kubadilisha rekodi bila approval ya jukumu husika.","image":"figures/01-live-homepage.png"},
    {"name":"WhatsApp Web Integration","sw":"Muunganisho wa WhatsApp Web","status":"CONFIGURATION REQUIRED","ui":"client/src/BusinessSphereDashboard.jsx; messaging surfaces","server":"server/whatsappProvider.ts; server/whatsAppSecurity.ts","prefixes":["whatsapp"],"roles":"Marketing manager; support agent; administrator","focus":"WhatsApp surfaces support account/contact/message/conversation metadata and provider readiness boundaries. Credentials and provider access must remain server-side.","workflow":"Configure provider → link account/contact → draft message → validate content → dispatch through provider → record status.","limit":"Provider configuration and message delivery are not assumed; credential persistence in browser storage is a security risk and must remain removed or blocked.","sw_focus":"WhatsApp ina metadata ya account, contacts, messages na conversations. API key au token haipaswi kuhifadhiwa kwenye browser, na delivery inategemea provider.","image":"figures/05-deployment-observation.png"},
    {"name":"Microfinance","sw":"Mikopo Midogo","status":"TESTING","ui":"client/src/components/MicrofinanceWorkspace.jsx; client/src/components/MicrofinanceGovernanceDialogs.jsx","server":"server/microfinanceOperations.ts","prefixes":["mfi"],"roles":"MFI manager; loan officer; cashier; approver; auditor","focus":"Microfinance covers borrowers, groups, products, applications, approvals, disbursement, repayments, savings, cash sessions, collections, scoring, and escalation.","workflow":"Register borrower/group → choose product → submit application → approve → disburse → collect repayment/savings → monitor PAR/escalation.","limit":"Financial persistence and controlled data must be exercised against a real tenant schema; seeded UI examples are not production records.","sw_focus":"Microfinance inahusu borrowers, vikundi, bidhaa, applications, approvals, disbursement, repayments, savings, cash sessions na collections.","image":"figures/04-business-workflow.png"},
    {"name":"Money Agent","sw":"Wakala wa Fedha","status":"IMPLEMENTED","ui":"client/src/components/MoneyAgentWorkspace.jsx; client/src/components/SectorCommandCenters.jsx","server":"server/moneyAgentOperations.ts","prefixes":["money_agent","money","agent"],"roles":"Money agent; supervisor; cashier; auditor","focus":"Money Agent covers agents, customers, wallets, cash actions, fees, commissions, ledger snapshots, approvals, reconciliation, and customer-facing summaries.","workflow":"Register agent/customer → open cash/wallet context → initiate action → validate balance/role → record ledger → reconcile and audit.","limit":"The live schema includes the migration history and tables, but real wallet/provider movement still requires controlled credentials and operational acceptance.","sw_focus":"Wakala wa fedha huunganisha mawakala, wateja, wallets, cash actions, commissions, ledger, approvals na reconciliation. Movement halisi ya mtoa huduma inahitaji acceptance ya kudhibitiwa.","image":"figures/11-payment-reconciliation.png"},
    {"name":"VICOBA, SACCOS, and Community Groups","sw":"VICOBA, SACCOS na Vikundi vya Jamii","status":"TESTING","ui":"client/src/components/SectorCommandCenters.jsx; community group surfaces","server":"server/communityGroups.ts; server/routers.ts","prefixes":["community","vicoba"],"roles":"Group chairperson; treasurer; secretary; member; auditor","focus":"Community-group features cover groups, members, contributions, savings, loans, meetings, attendance, budgets, projects, welfare, votes, messages, notifications, approvals, and audit.","workflow":"Create group → add member → record contribution/meeting → approve loan or welfare action → update ledger/report.","limit":"Relationship guards and approvals are present, but broader group-accounting walkthroughs should still be completed with controlled tenant data.","sw_focus":"Vikundi vina groups, members, michango, savings, loans, meetings, attendance, budgets, welfare, votes, approvals na audit.","image":"figures/04-business-workflow.png"},
    {"name":"Healthcare and Clinic","sw":"Afya na Kliniki","status":"PARTIALLY IMPLEMENTED","ui":"client/src/components/HealthcareClinicWorkspace.jsx","server":"server/healthcareOperations.ts; server/healthcareInteroperability.ts; server/healthcareReminders.ts","prefixes":["hc"],"roles":"Clinician; nurse; receptionist; finance user; patient portal user","focus":"Healthcare covers patients, appointments, visits, vitals, prescriptions, lab/radiology, claims, reports, reminders, patient SMS consent, portal references, and interoperability exports.","workflow":"Register patient → book appointment → record visit/vitals → prescribe/order → issue report or claim → notify with consent → audit.","limit":"SMS and interoperability depend on provider configuration; clinical data requires strict least-privilege review and must not be exposed through generic tenant screens.","sw_focus":"Kliniki ina patients, appointments, visits, vitals, prescriptions, labs, radiology, claims, reports, reminders na patient consent.","image":"figures/04-business-workflow.png"},
    {"name":"School Management","sw":"Usimamizi wa Shule","status":"PASSED","ui":"client/src/components/SchoolWorkspace.jsx","server":"server/schoolOperations.ts","prefixes":["sch"],"roles":"School administrator; teacher; student/guardian portal user; finance user","focus":"School Management includes academic years, terms, departments, subjects, classes, streams, teachers, admissions, attendance, assessments, report cards, assignments, fees, scholarships, transport, library, discipline, announcements, and portal links.","workflow":"Set academic structure → admit learner → assign teacher/class → record attendance/assessment → publish report → invoice/receive fees → communicate.","limit":"Automated browser coverage is strong for the surface, while full per-school CRUD remains a data-backed acceptance task.","sw_focus":"Mfumo wa shule una miaka ya masomo, terms, departments, subjects, classes, admissions, attendance, assessments, report cards, fees, scholarships, transport, library na announcements.","image":"figures/04-business-workflow.png"},
    {"name":"Pharmacy Management","sw":"Usimamizi wa Famasi","status":"TESTING","ui":"client/src/components/PharmacyWorkspace.jsx","server":"server/pharmacyOperations.ts","prefixes":["phm"],"roles":"Pharmacist; storekeeper; cashier; clinician; administrator","focus":"Pharmacy covers drug categories, brands, medicines, suppliers, purchase orders, stock receipt/adjustment/transfer, dispensing, sales, payments, claims, returns, reports, notifications, and audit.","workflow":"Create medicine → receive stock → adjust/transfer → dispense or sell → collect payment → return/claim → reconcile.","limit":"Live dispensing and inventory walkthroughs remain dependent on controlled pharmacy records and appropriate clinical/financial permissions.","sw_focus":"Famasi inahifadhi dawa, categories, brands, suppliers, purchase orders, stock, dispensing, sales, malipo, claims, returns na audit.","image":"figures/10-inventory-movement.png"},
    {"name":"Hotel and Hospitality","sw":"Hoteli na Ukarimu","status":"TESTING","ui":"client/src/components/HospitalityWorkspace.jsx; client/src/components/VerticalCommandCenters.jsx","server":"server/hospitalityOperations.ts","prefixes":["hospitality","htl"],"roles":"Hotel manager; front desk; housekeeping; finance; guest","focus":"Hospitality covers properties, room types, rooms, rate plans, reservations, guests, guest KYC, folios, payments, housekeeping, laundry, minibar, events, complaints, amenities, notifications, and reconciliation.","workflow":"Configure property/rooms → create reservation → check in → post folio/service → receive payment → housekeeping/check out → reconcile.","limit":"External booking-channel and payment connections remain provider-dependent; local reservation state is not proof of channel synchronization.","sw_focus":"Hoteli inahusu properties, room types, rooms, rates, reservations, guests, KYC, folios, payments, housekeeping, laundry, events na reconciliation.","image":"figures/04-business-workflow.png"},
    {"name":"Restaurant and Food & Beverage","sw":"Mgahawa na Chakula/Vinywaji","status":"TESTING","ui":"client/src/components/RestaurantWorkspace.jsx; client/src/components/RestaurantTanzaniaFiscalPanel.jsx","server":"server/restaurantManagement.ts; server/traFiscal.ts","prefixes":["restaurant","rst"],"roles":"Restaurant manager; waiter; cashier; kitchen; finance","focus":"Restaurant operations cover tax/fiscal profiles, menus, menu items, tables, orders, reservations, waiters, payments, receipts, refunds, and Tanzania fiscal configuration.","workflow":"Configure menu/table → open order → add items → prepare/serve → collect payment → issue receipt/fiscal evidence → reconcile.","limit":"Delivery and external payment channels require configuration; fiscal status is a readiness state until a provider confirms submission.","sw_focus":"Mgahawa una menus, items, tables, orders, reservations, waiters, payments, receipts, refunds na fiscal configuration ya Tanzania.","image":"figures/08-webhook-state-machine.png"},
    {"name":"Banking and MFI","sw":"Benki na MFI","status":"TESTING","ui":"client/src/components/BankMfiWorkspace.jsx; client/src/components/SectorCommandCenters.jsx","server":"server/bankMfiOperations.ts","prefixes":["bank","bnk"],"roles":"Bank administrator; teller; loan officer; compliance officer; auditor","focus":"Bank/MFI covers institutions, branches, customers, KYC documents, beneficial owners, accounts, beneficiaries, tellers, cash movements, wallets, loans, schedules, repayments, AML alerts, reconciliation, and audit events.","workflow":"Set institution → register customer/KYC → open account → post transaction → originate/approve loan → schedule/collect → reconcile/resolve AML.","limit":"External banking rails are not connected without approved adapters; SECURITY DEFINER routines require endpoint-by-endpoint privilege review.","sw_focus":"Benki na MFI zina institutions, branches, customers, KYC, accounts, beneficiaries, transactions, loans, schedules, repayments, AML, reconciliation na audit.","image":"figures/07-financial-control.png"},
    {"name":"Employee Portal","sw":"Portal ya Mfanyakazi","status":"IMPLEMENTED","ui":"client/src/components/EmployeePortalWorkspace.jsx","server":"server/teamWorkforce.ts; server/employeePortal.ts","prefixes":["hr","employee","workforce"],"roles":"Employee; manager; HR administrator","focus":"The portal provides self-service views for profile/work context, attendance, leave, announcements, benefits, documents, and approved workforce actions while retaining company scope.","workflow":"Accept secure invitation → verify profile → view or request employee action → manager/HR decision → record outcome.","limit":"An employee view does not imply access to payroll or HR administration; role boundaries must be enforced by server and RLS.","sw_focus":"Portal ya mfanyakazi hutoa self-service ya wasifu, attendance, leave, announcements, benefits na nyaraka kwa mipaka ya kampuni na role.","image":"figures/06-auth-tenancy.png"},
    {"name":"Property Management","sw":"Usimamizi wa Mali Isiyohamishika","status":"IMPLEMENTED","ui":"client/src/components/PropertyManagementWorkspace.jsx","server":"server/propertyManagementOperations.ts; server/propertyManagement.ts","prefixes":["property"],"roles":"Property manager; landlord/owner; agent; tenant; finance manager; maintenance user","focus":"Property Management covers portfolios, owners, buildings, plots, units, listings, agents, tenants, KYC documents, applications, leases, inspections, handover, rent schedules, tax/fee rules, service charges, utilities, invoices, payments, receipts, contractors, maintenance, work orders, expenses, budgets, insurance, notices, approvals, documents, ledgers, reconciliation, notifications, integrations, and audit.","workflow":"Portfolio → building/unit → owner/tenant/application → lease → inspection/handover → rent/utility invoice → payment/receipt → maintenance/expense → reconciliation.","limit":"The live audit shows the property migration and tables present; external payment, map, messaging, and contractor channels still require configuration and controlled acceptance.","sw_focus":"Property Management ina portfolio, owners, buildings, plots, units, listings, agents, tenants, applications, leases, inspections, handover, rent, utilities, invoices, payments, maintenance, expenses, budgets, insurance, notices na audit.","image":"diagrams/property-rental-workflow.png"},
    {"name":"Subscription and Billing","sw":"Usajili na Utozaji","status":"IMPLEMENTED","ui":"client/src/components/SubscriptionBillingWorkspace.jsx; client/src/lib/subscriptionAccess.js","server":"server/subscriptionBilling.ts","prefixes":["billing","subscription","tenant_subscriptions"],"roles":"Billing manager; owner; CEO; CFO; finance manager; administrator","focus":"The live commercial contract contains FREE_15 at TZS 0 for 15 days and six paid monthly packages. Paid activation is released only after verified provider status; the server and database own amount, entitlement, idempotency, and access state.","workflow":"Read catalog → start Free or create paid intent → send/verify provider request → reconcile payment → activate subscription/invoice → refresh access snapshot.","limit":"HarakaPay account authorization and provider credentials remain deployment prerequisites; failed or cancelled requests must not be shown as active entitlement.","sw_focus":"Mfumo wa usajili una FREE_15 ya TZS 0 kwa siku 15 na packages sita za kila mwezi. Paid access hutolewa baada ya provider kuthibitisha malipo; server na database ndiyo mamlaka.","image":"figures/08-webhook-state-machine.png"},
    {"name":"Global Admin Control Center","sw":"Kituo cha Udhibiti wa Msimamizi wa Mfumo","status":"IMPLEMENTED","ui":"client/src/components/GlobalAdminControlCenter.tsx","server":"server/globalAdmin.ts; server/routers.ts","prefixes":["platform_admin_actions","companies","profiles","billing_plan_audit_log"],"roles":"Super Administrator; platform administrator","focus":"Global administration is a separate control plane for platform-wide company, user, subscription, security, integration, health, and audit actions. It is not a broad tenant policy.","workflow":"Verify platform-admin authority → inspect scoped snapshot → perform explicit action → record platform audit → review outcome.","limit":"The platform_admin_actions table is intentionally protected from a broad authenticated policy; access uses the protected platform-admin RPC boundary.","sw_focus":"Global Admin ni control plane tofauti ya kampuni, watumiaji, subscriptions, usalama, integrations, health na audit. Si ruhusa pana kwa mtumiaji wa tenant.","image":"figures/06-auth-tenancy.png"},
    {"name":"Enterprise Settings and Security Control Center","sw":"Mipangilio na Udhibiti wa Usalama","status":"IMPLEMENTED","ui":"client/src/components/DashboardPreferencesDrawer.jsx; client/src/components/ProfileIdentityCenter.jsx","server":"server/workspaceSettings.ts; server/workspaceBranding.ts; server/roleChangeApprovals.ts","prefixes":["company_profile_settings","company_modules","branches","departments","user_table_preferences"],"roles":"Organization owner; administrator; manager; user with permitted settings","focus":"Settings cover company profile, branding, language, timezone/currency, receipts, idle timeout, modules, branches, departments, preferences, passkeys, and approval-aware role changes.","workflow":"Open settings → read confirmed configuration → edit permitted field → server validates and persists → refresh UI → audit sensitive change.","limit":"Provider-specific email, SMS, API-key, and storage configuration remains environment-dependent; secret values must be redacted.","sw_focus":"Mipangilio inahusu kampuni, branding, lugha, timezone/currency, receipts, idle timeout, modules, branches, departments, preferences, passkeys na role approvals.","image":"figures/06-auth-tenancy.png"},
    {"name":"Predictive Analytics","sw":"Uchambuzi wa Utabiri","status":"PARTIALLY IMPLEMENTED","ui":"client/src/components/PredictiveAnalyticsWorkspace.jsx","server":"server/marketIntelligence.ts; server/smartAssistant.ts","prefixes":["custom_kpis","financial_benchmarks","competitors"],"roles":"Executive; analyst; manager","focus":"Predictive surfaces present trends, forecasts, benchmarks, and next-best-action ideas when sufficient data and configured intelligence services exist.","workflow":"Select evidence set → compute signal → present confidence/limitation → inspect source metrics → decide manually.","limit":"Predictions are decision support, not guarantees. Sparse or stale data, configuration, and model availability constrain their reliability.","sw_focus":"Utabiri ni msaada wa maamuzi unaotegemea data ya kutosha na huduma iliyosanidiwa; si ahadi ya matokeo ya biashara.","image":"figures/01-live-homepage.png"},
]

ROLE_ROWS = [
    ("Super Administrator", "Platform control plane; not ordinary tenant access", "Global admin; security; health; audit", "Explicit platform-admin boundary"),
    ("Organization Owner", "Company-wide visibility and approvals", "All configured company modules", "Company scope plus sensitive settings"),
    ("CEO / COO / CFO", "Executive and financial review", "Reports, finance, sales, operational modules by assignment", "Manager checks and company scope"),
    ("Finance Manager / Accountant", "Finance, billing, reconciliation", "Finance, sales billing, subscriptions, tax evidence", "Billing-manager roles and RLS"),
    ("HR Manager / Manager", "Workforce and approvals", "HR, employee portal, leave, payroll as assigned", "Role-aware HR operations"),
    ("Sales Manager / Sales User", "Commercial pipeline and sales", "CRM, sales, POS as assigned", "Tenant scope and confirmed writes"),
    ("Warehouse Manager / Procurement Officer", "Stock and supply chain", "Inventory, procurement, fleet as assigned", "Movement, approval, and company scope"),
    ("Specialist operator", "Assigned industry workflow", "Healthcare, school, pharmacy, property, hospitality, banking, or MFI", "Verified profile, role, and module entitlement"),
    ("Member / Employee", "Self-service and assigned tasks", "Permitted records only", "RLS and server permission checks"),
]

WORKFLOWS = [
    ("Authentication and onboarding", "Public entry → Supabase configuration → password/OAuth/passkey → session persistence → verified profile → company context → protected workspace.", "Ukurasa wa umma → configuration ya Supabase → nenosiri/OAuth/passkey → kikao → wasifu uliothibitishwa → kampuni → workspace salama."),
    ("Sales to receipt", "Customer → quotation/order → invoice → confirmed payment → receipt/ledger → management report and audit.", "Mteja → quotation/order → ankara → malipo yaliyothibitishwa → receipt/ledger → ripoti na audit."),
    ("Procurement to stock", "Supplier → request/order → approval → receipt → stock movement → payable/finance evidence.", "Msambazaji → ombi/order → approval → mapokezi → stock movement → ushahidi wa deni/finance."),
    ("POS to reconciliation", "Register/shift → item/customer → price/stock validation → tender → confirmed sale → receipt → return/reconciliation.", "Register/shift → item/mteja → uthibitishaji wa bei na stock → tender → sale → receipt → return/reconciliation."),
    ("Finance close", "Transaction → journal or subledger → approval → bank/tax reconciliation → period close → report.", "Muamala → journal/subledger → approval → reconciliation ya benki/kodi → kufunga kipindi → ripoti."),
    ("HR and payroll", "Secure invite → employee profile → attendance/leave → approval → payroll calculation → report.", "Invitation salama → wasifu wa mfanyakazi → attendance/leave → approval → payroll → ripoti."),
    ("Property rental", "Portfolio → unit → owner/tenant/application → lease → inspection → invoice → payment → maintenance/expense → reconciliation.", "Portfolio → unit → owner/tenant/application → lease → inspection → invoice → malipo → maintenance/expense → reconciliation."),
    ("Healthcare care path", "Patient → appointment → visit/vitals → prescription/lab → report/claim → consent-based reminder → audit.", "Mgonjwa → appointment → visit/vitals → prescription/lab → report/claim → reminder yenye consent → audit."),
    ("Subscription lifecycle", "Catalog → Free activation or paid intent → provider status → verified settlement → subscription/invoice → access snapshot → expiry/renewal decision.", "Catalog → Free au paid intent → status ya provider → settlement iliyothibitishwa → subscription/invoice → access snapshot → expiry/renewal."),
    ("AI proposal path", "Verified user → bounded context → structured response → navigation/proposal → independent approval → confirmed module mutation.", "Mtumiaji aliyehakikiwa → context yenye mipaka → response iliyopangwa → proposal → approval huru → mutation iliyothibitishwa."),
]

INTRO_CHAPTERS = [
    ("Executive Purpose", "Purpose, audience, evidence boundary, and how to use this book.", "Kitabu hiki kinaeleza bidhaa, uendeshaji, usanifu, usalama na mipaka ya ushahidi kwa wamiliki, watumiaji, wasimamizi, developers, support teams na auditors.", "The book is a repository-audited reference, not a generic ERP brochure. It distinguishes implemented source and database evidence from configuration requirements, partial contracts, external services, and planned work."),
    ("Product Philosophy", "The platform connects records, people, workflows, control surfaces, and decision views without claiming that every problem is automated.", "Falsafa ya bidhaa ni kuunganisha rekodi, watu, mitiririko, udhibiti na maoni ya uongozi bila kudai kwamba kila tatizo litatatuliwa kiotomatiki.", "The design intent is to make business work measurable, responsibilities visible, and decisions easier to explain. The evidence boundary is part of the product quality model."),
    ("Business Value", "Smart Manager helps organizations reduce fragmented records, improve visibility, protect access, and create repeatable operational paths.", "Smart Manager husaidia kupunguza rekodi zilizotawanyika, kuongeza mwonekano, kulinda ufikiaji na kujenga mitiririko inayoweza kurudiwa.", "Value is created when a real operational fact is captured once, validated at the server boundary, enforced in the database, and surfaced through the right role-aware view."),
    ("Audience and Adoption", "Different roles use different evidence: owners review control, operators execute workflows, accountants reconcile, administrators govern, and developers extend.", "Kila jukumu hutumia ushahidi tofauti: wamiliki huangalia udhibiti, operators hutekeleza workflows, accountants hufanya reconciliation, admins hutawala na developers huendeleza.", "Adoption should start with a bounded workflow, a real company context, verified data, and a measured review rhythm rather than a promise to activate every module at once."),
    ("Evidence and Status Language", "Status words are delivery controls: IMPLEMENTED, TESTING, PARTIALLY IMPLEMENTED, CONFIGURATION REQUIRED, EXTERNAL SERVICE REQUIRED, and BLOCKED.", "Maneno ya hali ni udhibiti wa delivery: IMEJENGWA, INAJARIBIWA, IMEJENGWA KWA SEHEMU, INAHITAJI USANIDI, INAHITAJI HUDUMA YA NJE na IMEZUIWA.", "A status never means that a feature is universally complete for every deployment. It describes the strongest evidence found during this audit window."),
]

ARCH_CHAPTERS = [
    ("Technology Stack", "Vite, React 19, TypeScript/JavaScript, Tailwind, Express 5, tRPC 11, Supabase Auth/Postgres/REST/RLS, S3-compatible storage, and Vercel-compatible serverless runtime are present in the project. The package manifest is the source for exact dependency versions.", "Stack ya teknolojia ina Vite, React 19, TypeScript/JavaScript, Tailwind, Express 5, tRPC 11, Supabase Auth/Postgres/REST/RLS, storage inayolingana na S3 na runtime ya Vercel."),
    ("Frontend Architecture", "App.tsx owns the top-level route switch and providers. BusinessSphereDashboard.jsx is the authenticated module composition boundary with lazy workspaces, a direct Supabase REST client path, shared persistence guards, subscription access adapter, and role-aware navigation.", "App.tsx inamiliki routes na providers. BusinessSphereDashboard.jsx ndiyo mpaka mkubwa wa workspace yenye lazy modules, Supabase REST, persistence guards, subscription access na urambazaji wa role."),
    ("Backend Architecture", "Express API bootstrap separates API-only behavior from SPA hosting. HTTP routes cover public configuration, billing, HarakaPay, scheduled handlers, webhooks, storage proxy, OAuth, and the tRPC middleware. Feature routers and operations live in server files.", "Express API hutenganisha API na SPA. Routes za HTTP zinahusu public config, billing, HarakaPay, scheduled handlers, webhooks, storage proxy, OAuth na tRPC. Operations za moduli ziko server-side."),
    ("Database Architecture", "The repository uses a source-versioned Supabase migration history. The live audit on 24 August 2026 returned 542 tables: 519 public and 23 auth; 535 tables reported RLS enabled and 7 did not. Table definitions, indexes, constraints, functions, grants, and policies are migration-owned.", "Repository ina migration history ya Supabase yenye version. Audit ya 24 Agosti 2026 ilirudisha tables 542: 519 public na 23 auth; tables 535 zilionesha RLS na 7 hazikuonesha."),
    ("Supabase Authority", "Supabase is authoritative for authenticated profiles, company scope, business records, subscription plans, payment state, invoices, entitlements, and access snapshots. Browser storage can hold session tokens under the supported auth flow, but it cannot grant subscription access or represent a durable business write.", "Supabase ndiyo mamlaka ya profiles, company scope, rekodi za biashara, packages, payment state, invoices, entitlements na access snapshots. Browser haiwezi kutoa ruhusa ya subscription au kudai business write imehifadhiwa."),
    ("Authentication Architecture", "The public gateway obtains public Supabase configuration from /api/config/public when needed, calls Supabase Auth endpoints for password and recovery operations, supports Google/Azure/Apple provider selection, handles OAuth hash callbacks, and persists access/refresh tokens through the supported session adapter.", "Gateway hupata public config kupitia /api/config/public, hutumia Supabase Auth kwa password/recovery, ina Google/Azure/Apple na hushughulikia OAuth callbacks pamoja na token persistence kupitia adapter."),
    ("Authorization and RLS", "Authorization is layered: session extraction, verified profile resolution, role checks, company scope, protected server procedures, database functions, grants, constraints, and RLS policies. A visible button is not proof of permission. Global administration has a separate policy boundary.", "Uidhinishaji una tabaka: kikao, verified profile, role, company scope, procedures za server, functions, grants, constraints na RLS. Button inayoonekana si ushahidi wa permission. Global admin ina mpaka tofauti."),
    ("Multi-Tenant Architecture", "The effective model is platform → organization/company → workspace context → authenticated profile/membership → role/module entitlement → row-scoped data. The client does not select an arbitrary company ID as authority; Supabase profile and RLS scope control what can be read or changed.", "Mfumo ni platform → organization/company → workspace → profile/membership → role/entitlement → data ya row. Client haipaswi kuchagua company ID kiholela; profile na RLS ndiyo hudhibiti data."),
    ("API Architecture", "Critical flows use server HTTP handlers or protected tRPC procedures. Subscription routes call user-scoped or service-scoped Supabase RPCs. Specialist operations validate Zod inputs and forward authenticated tokens. Service credentials remain server-side.", "Flows muhimu hutumia HTTP handlers za server au protected tRPC. Billing hutumia user/service RPCs za Supabase. Operations maalumu huthibitisha Zod input na kutuma token iliyothibitishwa. Siri hubaki server-side."),
    ("Storage Architecture", "File metadata and references are separated from file bytes. Avatar and property-document paths are scoped by company/entity; storage proxy and download boundaries enforce controlled access. Missing storage configuration must be shown as unavailable rather than successful.", "Metadata na references zimetenganishwa na bytes za file. Avatar na property documents zina keys zenye company/entity scope; storage proxy na download boundaries zinalinda access."),
    ("AI Architecture", "The assistant uses a built-in model boundary with a default gpt-5-mini model, bounded history/context, JSON Schema response format, safe module targets, a limited operation proposal list, and an explicit requirement for independent approval before mutation.", "Assistant hutumia built-in model, history/context yenye mipaka, JSON Schema, safe targets za modules na proposal chache. Mutation inahitaji approval huru ya role husika."),
    ("Deployment Architecture", "Local development mounts Vite and the API; production bundles client and API separately. When VERCEL=1 the runtime exposes the API app for serverless import rather than opening a local listener. Builds run Supabase schema verification when server credentials are available and skip it explicitly for credentialless Vercel builds.", "Development hutumia Vite na API; production hubundle client na API. VERCEL=1 huwezesha import ya API ya serverless badala ya listener wa local. Build huendesha schema verification pale credentials za server zipo."),
]

SECURITY_FINDINGS = [
    ("Authenticated SECURITY DEFINER routines", "WARN", "The live security advisor reported 118 WARN and 1 INFO lint at the audit timestamp; many WARNs identify signed-in execution of SECURITY DEFINER routines.", "Review each signature, keep only intentionally callable endpoints, pin search paths, apply narrow grants, and move internal helpers out of the exposed API surface where possible."),
    ("Multiple permissive RLS policies", "WARN", "The live performance advisor reported 851 lints, including multiple permissive policies on the same table/action.", "Consolidate overlapping policies by command and role after verifying semantics; do not blindly drop production policies."),
    ("Legacy non-atomic invoice payment path", "P0 historical finding", "The audit report records separate sales_payments insertion and invoice balance update without a proven atomic idempotency RPC.", "Add a reviewed tenant-scoped atomic RPC and durable idempotency key before claiming concurrent-safe posting."),
    ("Large dashboard boundary", "P2", "BusinessSphereDashboard.jsx remains a very large monolithic component and the build reports a non-fatal large-chunk warning.", "Decompose incrementally after persistence and live-environment blockers are addressed; avoid cosmetic rewrites that increase risk."),
    ("External provider readiness", "Configuration boundary", "HarakaPay, TRA/VFD, WhatsApp, email/SMS, storage, and AI services depend on deployment configuration and approved credentials.", "Keep provider secrets server-side, expose readiness states, and test only controlled sandbox or authorized production paths."),
    ("Demo fallback risk", "Medium", "The client has an explicit seed-data fallback when Supabase is not configured.", "Production deployments must fail closed with a clear configuration message; demo mode must remain explicit and non-operational."),
]

TROUBLESHOOTING = [
    ("Login fails", "Missing public Supabase configuration, invalid credentials, expired recovery session, or provider rejection.", "Check /api/config/public, browser network response, Auth status, and the exact user-facing error; do not log passwords or tokens.", "Restore approved configuration or repeat the supported recovery flow; keep error messages generic enough to avoid account enumeration."),
    ("Signup does not enter workspace", "Email confirmation, profile provisioning, or company assignment is incomplete.", "Inspect Auth confirmation state, verified profile, company_memberships, and server logs without exposing secrets.", "Complete verification and secure onboarding; do not create a duplicate users table."),
    ("RLS permission error", "Session missing, wrong company scope, role not permitted, or policy/column contract mismatch.", "Verify bearer extraction, profile company_id, role, table company_id, and policy command.", "Fail closed, fix the specific policy or caller contract, and add a regression test."),
    ("Catalog returns 503", "Server-side Supabase billing configuration missing or billing RPC unavailable.", "Check /api/billing/catalog, server env presence, migration ledger, and billing_public_plan_catalog readiness.", "Restore server configuration or apply the intended migration in a controlled window; never send the service key to the browser."),
    ("Free activation rejected", "No verified manager role, missing company profile, or FREE_15 catalog/routine not ready.", "Confirm authenticated profile, company scope, active catalog row, and billing_start_free_plan routine.", "Correct prerequisites and retry idempotently; never treat a failed call as active."),
    ("Paid payment remains pending", "Provider has not verified the order, webhook/status polling has not settled, or provider configuration is unavailable.", "Check payment record, provider order ID, server status call, and webhook receipt.", "Keep access inactive until verified; retry only through idempotent server/database flow."),
    ("Data appears missing", "Wrong company scope, empty state, failed query, or UI cache not refreshed.", "Compare authenticated company scope, raw server response, loading/error state, and table RLS.", "Show unavailable/empty/error explicitly; never replace a failed read with fake seed data in production."),
    ("File upload fails", "Storage configuration, size/type validation, path scope, or signed URL problem.", "Inspect server validation and storage response; verify only metadata after confirmed upload.", "Keep bytes in storage, metadata in database, and show a retryable failure."),
    ("TRA/VFD submission blocked", "Profile not configured, provider readiness false, missing credentials, or duplicate idempotency key.", "Read fiscal profile, readiness, receipt/retry queue status, and audit event.", "Configure approved provider or retain a truthful blocked state; do not call a local receipt fiscalized."),
    ("AI response cannot be used as an action", "AI returned a proposal, not an approved mutation, or the request exceeded safe targets.", "Inspect structured response, proposal operation, approval state, and source context.", "Require independent role approval and confirmed module write; explain missing evidence."),
    ("Deployment build fails", "Schema verification credentials, dependency resolution, TypeScript, or route/runtime incompatibility.", "Run pnpm check, pnpm test, verify:supabase-schema with approved env, and inspect the exact build diagnostic.", "Fix the specific contract; do not bypass the schema guard or suppress audit failures."),
    ("Webhook is not reconciled", "Unknown provider order, failed status verification, invalid payload, or endpoint configuration.", "Confirm order lookup, provider status request, response order match, and service RPC result.", "Return a safe error, preserve audit state, and retry through a controlled path."),
]

GLOSSARY = [
    ("RLS", "Row Level Security; database policy enforcement over which rows a role can read or change.", "Usalama wa kiwango cha mstari; sera za database zinazoamua rows ambazo role inaweza kusoma au kubadilisha."),
    ("Company scope", "The organization boundary used to separate tenant operations.", "Mpaka wa kampuni unaotenganisha shughuli za tenant."),
    ("Entitlement", "A server/database-confirmed permission to use a package or module.", "Ruhusa iliyothibitishwa na server/database ya kutumia package au module."),
    ("Idempotency", "Repeating a request with the same key does not create a duplicate operation.", "Kurudia request yenye key ileile hakutengenezi operation ya pili."),
    ("SECURITY DEFINER", "A database function that runs with owner privileges and therefore requires narrow grants and safe search_path.", "Function ya database inayotumia mamlaka ya owner; inahitaji grants finyu na search_path salama."),
    ("FREE_15", "The free subscription package at TZS 0 for 15 days, with no automatic charge.", "Kifurushi cha bure cha TZS 0 kwa siku 15 bila automatic charge."),
    ("Calendar month", "A billing period calculated with calendar-month arithmetic rather than a fixed number of days.", "Kipindi kinachohesabiwa kwa miezi ya kalenda, si idadi ya siku isiyobadilika."),
    ("Fail closed", "Deny access or mutation when verification is missing rather than guessing success.", "Kukataa access au mutation pale uthibitisho haupo badala ya kukisia mafanikio."),
]


def load_json(path: Path, default: Any) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def live_snapshot() -> dict[str, Any]:
    data = load_json(LIVE_TABLES, {"tables": []})
    tables = data.get("tables", [])
    names = [t.get("name", "") for t in tables]
    migrations = load_json(LIVE_MIGRATIONS, {"migrations": []})
    migration_rows = migrations.get("migrations", migrations if isinstance(migrations, list) else [])
    sec = load_json(SECURITY_ADVISORS, {}).get("result", {}).get("lints", [])
    perf = load_json(PERF_ADVISORS, {}).get("result", {}).get("lints", [])
    return {
        "tables": tables,
        "names": names,
        "table_count": len(tables),
        "public_count": sum(n.startswith("public.") for n in names),
        "auth_count": sum(n.startswith("auth.") for n in names),
        "rls_enabled": sum(t.get("rls_enabled") is True for t in tables),
        "rls_disabled": sum(t.get("rls_enabled") is not True for t in tables),
        "migrations": migration_rows,
        "migration_count": len(migration_rows),
        "security_lints": sec,
        "performance_lints": perf,
    }


LIVE = live_snapshot()


def table_names(prefixes: list[str]) -> list[str]:
    found: list[str] = []
    for table in LIVE["tables"]:
        name = table.get("name", "")
        short = name.split(".", 1)[-1]
        if any(short == p or short.startswith(p + "_") for p in prefixes):
            found.append(short)
    return sorted(set(found))


def source_exists(path_text: str) -> bool:
    path = REPO / path_text.split(";", 1)[0].strip()
    return path.exists()


def escape_md(value: str) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(escape_md(h) for h in headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    out.extend("| " + " | ".join(escape_md(str(x)) for x in row) + " |" for row in rows)
    return "\n".join(out) + "\n"


def typst_escape(value: str) -> str:
    return str(value).replace("\\", "\\\\").replace("#", "\\#").replace("[", "(").replace("]", ")").replace("_", "\\u{005F}").replace("$", "\\$")


def typst_table(headers: list[str], rows: list[list[str]], widths: str = "(3.2cm, 3.8cm, 3.4cm, 5.6cm)") -> str:
    cells = [f"[*{typst_escape(h)}*]" for h in headers]
    for row in rows:
        cells.extend(f"[{typst_escape(str(v))}]" for v in row)
    header = ", ".join(cells[:len(headers)])
    body = ", ".join(cells[len(headers):])
    return f"#table(columns: {widths}, inset: 4pt, stroke: 0.35pt + luma(190), table.header({header}), {body})"


def add_image_typst(lines: list[str], rel: str, caption: str, width: str = "14.5cm") -> None:
    path = ASSETS / rel
    if path.exists():
        lines.append(f'#figure(image("assets/{rel}", width: {width}), caption: [{typst_escape(caption)}])')


def english_module_paragraphs(m: dict[str, Any]) -> list[str]:
    tables = table_names(m["prefixes"])
    table_text = ", ".join(f"`{x}`" for x in tables[:18]) if tables else "shared company-scoped persistence tables"
    return [
        f"{m['focus']} This chapter is written from the repository UI, server boundary, migration history, and the read-only live Supabase inventory. The strongest evidence is named; anything that depends on configuration, connected data, or an external provider remains qualified.",
        f"The purpose of {m['name']} is to make a repeatable business operation visible and controlled. The primary business problem is not the absence of a screen; it is the risk that the same fact is entered in several places, that a role sees too much, or that a failed write is presented as success. The module therefore belongs to a path from input to validation, persistence, evidence, and management review.",
        f"The representative persistence surface observed for this module includes {table_text}. These names are evidence from the current live inventory or the repository contract; they are not an assertion that every table is used by every deployment or that every listed record contains production data.",
        f"The verified workflow is: {m['workflow']} Operators should complete the step in sequence, confirm the returned state, and use the module’s error or empty state rather than substituting a local guess. {m['limit']}",
    ]


def swahili_module_paragraphs(m: dict[str, Any]) -> list[str]:
    tables = table_names(m["prefixes"])
    table_text = ", ".join(f"`{x}`" for x in tables[:18]) if tables else "tables za shared company scope"
    return [
        f"{m['sw_focus']} Sura hii imetokana na UI ya repository, mpaka wa server, migration history na inventory ya Supabase iliyosomwa bila kubadilisha data. Pale ushahidi unategemea configuration, data ya kampuni au provider wa nje, mpaka huo umeelezwa wazi.",
        f"Lengo la {m['sw']} ni kufanya kazi inayorudiwa ionekane na idhibitiwe. Changamoto si kutokuwepo kwa screen pekee; ni hatari ya kuandika ukweli uleule mara nyingi, role kuona zaidi ya inavyoruhusiwa, au write iliyokataliwa kuonekana kana kwamba imefanikiwa. Kwa hiyo module inapaswa kufuata input, validation, persistence, ushahidi na review ya uongozi.",
        f"Sehemu za persistence zinazoonekana kwa module hii zinajumuisha {table_text}. Majina haya ni ushahidi wa inventory ya sasa au contract ya repository; hayamaanishi kwamba kila table inatumika na kila deployment au kwamba kila table ina data ya production.",
        f"Mtiririko uliothibitishwa ni: {m['workflow']} Mtumiaji athibitishe hali iliyorejeshwa na atumie empty/error state badala ya kukisia. {m['limit']} {m['sw_focus']}",
    ]


def module_rows(m: dict[str, Any]) -> list[list[str]]:
    return [
        ["Status", m["status"], "Evidence", f"UI: {m['ui']}; server: {m['server']}"],
        ["Purpose", m["name"], "Primary users", m["roles"]],
        ["Navigation", "Authenticated module registry or dedicated workspace", "Workflow", m["workflow"]],
        ["Database", ", ".join(table_names(m["prefixes"])[:16]) or "Shared company-scoped tables", "Security", "Verified profile, company scope, role checks, RLS/policies where applicable"],
        ["Known limitation", m["limit"], "Future improvement", "Complete controlled authenticated CRUD and provider acceptance where the status is not PASSED."],
    ]


def operator_runbook_en(m: dict[str, Any]) -> list[tuple[str, str]]:
    return [
        ("1. Confirm context", f"Confirm the authenticated profile, company/workspace context, and role before opening {m['name']}. Do not copy a company identifier from an untrusted screen or URL."),
        ("2. Prepare the record", "Collect the minimum required business facts, use the supported date/amount formats, and confirm that any referenced customer, item, unit, employee, or account belongs to the same company scope."),
        ("3. Enter and validate", "Enter the record through the supported screen. Let the server-side schema and business rules reject invalid or incomplete input; do not bypass validation with browser storage or direct hidden requests."),
        ("4. Confirm the result", "Wait for the server/database response. Treat a returned identifier, status, or confirmed row as the completion signal. If the response is unavailable, keep the action pending or failed."),
        ("5. Review downstream evidence", f"Check the related record, report, audit entry, notification, or reconciliation state where applicable. The expected business path is {m['workflow']}"),
        ("6. Escalate exceptions", f"Record the exact error, time, company, role, and safe reproduction steps. Escalate configuration or provider issues separately from application defects. Known boundary: {m['limit']}"),
    ]


def operator_runbook_sw(m: dict[str, Any]) -> list[tuple[str, str]]:
    return [
        ("1. Hakikisha context", f"Hakikisha profile iliyothibitishwa, company/workspace na role kabla ya kufungua {m['sw']}. Usinakili company identifier kutoka screen au URL isiyoaminika."),
        ("2. Andaa record", "Kusanya taarifa muhimu za biashara, tumia tarehe/kiasi kwa format inayokubalika na hakikisha reference ya mteja, item, unit, mfanyakazi au account ni ya kampuni hiyo."),
        ("3. Ingiza na thibitisha", "Ingiza record kupitia screen inayoungwa mkono. Schema na business rules za server zikatae input mbovu; usivunje validation kwa browser storage au request iliyofichwa."),
        ("4. Thibitisha matokeo", "Subiri response ya server/database. ID, status au row iliyothibitishwa ndiyo signal ya kumaliza. Response ikiwa haipo, acha action ikiwa pending au failed."),
        ("5. Kagua ushahidi unaofuata", f"Kagua record inayohusiana, ripoti, audit, notification au reconciliation. Mtiririko unaotarajiwa ni {m['workflow']}"),
        ("6. Escalate exceptions", f"Andika error kamili, muda, kampuni, role na hatua salama za kurudia. Tenganisha tatizo la configuration/provider na defect ya application. Mpaka unaojulikana: {m['limit']}"),
    ]


def module_evidence_rows(m: dict[str, Any]) -> list[list[str]]:
    return [
        ["Identity", "Verified user and company context", "Profile/company resolution"],
        ["Input", "Required fields, valid references, date/amount rules", "Server validation response"],
        ["Persistence", "Returned row, ID, status, or RPC result", "Database/server confirmation"],
        ["Control", "Approval, reconciliation, audit, or notification state", "Role-aware follow-up"],
        ["Integration", "Provider readiness and delivery outcome", "Separate external evidence"],
    ]


def markdown_front_matter() -> list[str]:
    return [
        "# SMART MANAGER ERP — MASTER SYSTEM BOOK",
        "## Repository-Audited English and Tanzanian Swahili Edition",
        "",
        "> **Version:** 1.0.0  \n> **Documentation date:** 24 August 2026  \n> **Documentation status:** Repository-Audited Edition  \n> **Author:** Manus AI  \n> **Product owner/creator:** Ezra Mpapi, as identified in supplied project materials  \n> **Location:** Dar es Salaam, Tanzania",
        "",
        f"![Official SMART MANAGER logo]({MASTER_BOOK_LOGO_STORAGE})",
        "",
        "## Copyright and evidence notice",
        "This master book documents the SMART MANAGER ERP repository at the audit date shown above. It is not a generic ERP manual and does not convert every navigation surface into a claim of production completeness. Implemented, tested, partial, configuration-dependent, external-service-dependent, blocked, and planned states are separated throughout the book.",
        "",
        "The live Supabase inspection was read-only. It recorded schema metadata, migration history, and advisor findings without changing production data. The exact snapshots are preserved in the `evidence/` directory. Secrets, private credentials, API keys, passwords, and provider tokens are intentionally excluded or redacted.",
        "",
        "## How to use this book",
        "The English edition is the primary technical reference. The complete Tanzanian Swahili edition follows the English reference and preserves technical names such as `RLS`, `Supabase`, `API`, `workspace`, `subscription`, and `workflow` where those terms are standard. Each module chapter follows the requested pattern: overview, purpose, business problems, target users, navigation, features, database, permissions, workflows, screens, reports, integrations, security, mobile experience, known limitations, current status, and future improvements.",
        "",
        "## Table of contents",
        "The rendered PDF contains an automatic table of contents. The source structure is divided into product introduction, system overview, technical architecture, module encyclopedia, business workflows, user manual, administrator manual, developer manual, security, database, integrations, troubleshooting, operations, future roadmap, and the Swahili edition.",
        "",
    ]


def build_markdown() -> str:
    out = markdown_front_matter()
    out += ["# PART I — PRODUCT INTRODUCTION", ""]
    for idx, (title, sw_title, sw_intro, en) in enumerate(INTRO_CHAPTERS, 1):
        out += [f"## I.{idx} {title}", "", en, "", f"### Kiswahili: {sw_title}", "", sw_intro, "", "---", ""]
    out += ["# PART II — SYSTEM OVERVIEW", ""]
    for i, (title, en, sw) in enumerate([
        ("What is SMART MANAGER?", "SMART MANAGER is an authenticated business operating platform whose verified repository surfaces connect commercial, operational, financial, people, sector, intelligence, and control workflows.", "SMART MANAGER ni platform ya uendeshaji wa biashara yenye uthibitishaji inayounganisha workflows za biashara, operations, fedha, watu, sekta, intelligence na udhibiti."),
        ("ERP concept", "ERP is treated as a connected record and control model: a sale may relate to stock, finance, customer history, receipt, and report when the relevant contracts are enabled.", "ERP inaeleweka kama mfumo wa rekodi na udhibiti uliounganishwa: mauzo yanaweza kuhusishwa na stock, fedha, historia ya mteja, receipt na ripoti pale contracts husika zimewezeshwa."),
        ("System ecosystem", "The ecosystem includes public entry, secure authentication, company/workspace context, role-aware modules, Supabase persistence, server APIs, scheduled handlers, storage, external providers, and audit evidence.", "Ecosystem ina entry ya umma, authentication, company/workspace, modules za role, Supabase, APIs za server, scheduled handlers, storage, providers wa nje na audit."),
        ("Supported business types", "The verified module registry spans universal SME operations and specialist contexts such as property, healthcare, school, pharmacy, hospitality, restaurant, fleet, microfinance, banking, VICOBA, and money-agent operations.", "Rejista ya modules inahusisha SME operations na sekta kama property, afya, shule, famasi, hoteli, mgahawa, fleet, microfinance, benki, VICOBA na wakala wa fedha."),
        ("Lifecycle model", "A safe lifecycle is discover → authenticate → establish company scope → record → validate → approve/reconcile → report → audit → improve.", "Lifecycle salama ni gundua → authenticate → weka company scope → rekodi → thibitisha → approve/reconcile → ripoti → audit → boresha."),
    ], 1):
        out += [f"## II.{i} {title}", "", en, "", f"### Kiswahili", "", sw, "", "---", ""]
    out += ["# PART III — TECHNICAL ARCHITECTURE", ""]
    for i, (title, en, sw) in enumerate(ARCH_CHAPTERS, 1):
        out += [f"## III.{i} {title}", "", en, "", f"### Kiswahili: {title}", "", sw, ""]
        if title == "Technology Stack":
            out += [md_table(["Layer", "Verified implementation", "Evidence"], [["Frontend", "Vite + React 19 + Tailwind + wouter", "package.json; client/src/App.tsx"], ["Backend", "Express 5 + tRPC 11 + server handlers", "server/_core/apiApp.ts; server/routers.ts"], ["Data", "Supabase Auth/Postgres/REST/RLS; source-versioned migrations", "supabase/migrations; live evidence snapshot"], ["Runtime", "Vercel-compatible API import plus local Vite/Express", "server/_core/index.ts; package scripts"]]), ""]
        if title == "Database Architecture":
            out += [md_table(["Live metric", "Value", "Interpretation"], [["Total tables", str(LIVE["table_count"]), "Read-only schema inventory at 24 August 2026"], ["Public tables", str(LIVE["public_count"]), "Application-facing public schema inventory"], ["Auth tables", str(LIVE["auth_count"]), "Supabase Auth schema inventory"], ["RLS enabled", str(LIVE["rls_enabled"]), "Tables reported with RLS enabled"], ["RLS not enabled", str(LIVE["rls_disabled"]), "Requires table-specific review; not a reason to add broad policies blindly"], ["Migration records", str(LIVE["migration_count"]), "Live migration ledger records returned by connector"]]), ""]
        out += ["---", ""]
    out += ["# PART IV — COMPLETE MODULE BOOK", "", "The following module chapters are generated from the current verified module registry, actual source boundaries, migration names, and live table prefixes. The status is evidence-based and explicitly qualified.", ""]
    for idx, m in enumerate(MODULES, 1):
        out += [f"## IV.{idx} {m['name']}", "", "### Module Overview", "", *english_module_paragraphs(m), "", "### Purpose", "", m["focus"], "", "### Business Problems Solved", "", "The module addresses fragmented records, unclear ownership, weak review paths, or slow movement between an operational event and a management decision. It does so only to the extent shown by the source and database evidence above.", "", "### Target Users and Navigation", "", f"Primary roles: {m['roles']}. Navigation is exposed through the authenticated dashboard registry or the dedicated workspace named in the UI evidence: `{m['ui']}`.", "", "### Main Features and Workflow", "", m["workflow"], "", "### Database, Permissions, and Security", "", f"The server evidence is `{m['server']}`. Access is expected to follow verified profile resolution, company scope, role checks, database constraints, and RLS/policies. Representative live table prefixes are: {', '.join(table_names(m['prefixes'])[:18]) or 'shared company-scoped tables'}.", "", "### Screens, Reports, and Integrations", "", "Screens are evidence of an interaction surface. Reports are evidence of a computed view over source records. Integrations are separate dependencies; a button or draft state is not proof of a provider-side action.", "", "### Mobile Experience", "", "Responsive behavior is part of the tested surface where browser evidence exists. Operators should use the narrow layout for review and safe entry, while high-risk approvals, financial posting, and configuration remain subject to the same server/database controls.", "", "### Known Limitations", "", m["limit"], "", "### Current Implementation Status", "", f"**{m['status']}** — This status is a documentation control, not a sales promise. It should be revalidated after migration, provider, role, or deployment changes.", "", "### Future Improvements", "", "Complete the next safe evidence step: controlled authenticated CRUD, module-specific regression journeys, provider readiness verification, or safe decomposition of the legacy dashboard, depending on the limitation above.", "", "### Operator Runbook", "", *[f"**{step}:** {instruction}" for step, instruction in operator_runbook_en(m)], "", "### Evidence Control Table", "", md_table(["Control stage", "Expected evidence", "Review question"], module_evidence_rows(m)), "", "### Exception and Recovery Notes", "", f"If {m['name']} reports an unavailable, failed, pending, or rejected state, preserve the state and input, capture the exact response, verify company and role scope, and retry only through the supported server/database path. Do not turn an optimistic local state into a claim of completion.", "", "---", ""]
    out += ["# PART V — BUSINESS WORKFLOWS", ""]
    for i, (name, en, sw) in enumerate(WORKFLOWS, 1):
        out += [f"## V.{i} {name}", "", en, "", f"### Kiswahili", "", sw, "", "#### Control points", "", "At each transition, confirm identity, company scope, role, input validation, database response, and audit/result state. A failed step must remain failed; a local optimistic row is not a durable business record.", "", "#### Evidence boundary", "", "This workflow is a map of connected repository contracts. It is not a claim that every external provider, schedule, credential, device, or tenant dataset is enabled in every deployment.", "", "---", ""]
    out += ["# PART VI — USER MANUAL", ""]
    user_steps = [("First-day onboarding", "Sign in through the supported gateway, verify the company context, review your role, open one permitted module, and create only a record that the server confirms.", "Siku ya kwanza ingia kupitia gateway, hakikisha company context, soma role yako, fungua module inayoruhusiwa na tengeneza record ambayo server imethibitisha."), ("Daily operating rhythm", "Start with the dashboard, inspect exceptions, execute the module workflow, verify returned state, and leave approvals or reconciliation items visible to the responsible role.", "Anza na dashibodi, kagua exceptions, tekeleza workflow, thibitisha hali iliyorejeshwa na acha approvals/reconciliation zionekane kwa role husika."), ("Safe data entry", "Use required fields, valid dates and amounts, company-scoped references, and clear descriptions. Do not paste secrets into business notes.", "Tumia fields zinazohitajika, tarehe na kiasi sahihi, references za kampuni na maelezo wazi. Usiweke siri ndani ya notes za biashara."), ("Errors and retries", "Read the error state, preserve the form, correct the input or permission, and retry only after the cause is understood. Do not click repeatedly on payment or approval actions without an idempotency path.", "Soma error, hifadhi form, rekebisha input au permission na retry baada ya kuelewa sababu. Usibonyeze payment au approval mara nyingi bila idempotency."), ("Reports and review", "Use reports to ask a question of recorded data. Follow a surprising number back to the source document, payment, movement, or journal entry.", "Tumia ripoti kuuliza swali la data iliyorekodiwa. Namba ya kushangaza ifuatilie hadi document, payment, movement au journal entry ya chanzo."), ("Subscription access", "Read the server access state. FREE_15 is free for 15 days; paid access requires verified payment. Do not use local browser state as entitlement authority.", "Soma access state ya server. FREE_15 ni bure kwa siku 15; paid access inahitaji payment iliyothibitishwa. Browser si mamlaka ya entitlement."), ("Working on mobile", "Use responsive views for review and entry, maintain stable connectivity for writes, and wait for confirmed response before navigating away from a critical action.", "Tumia responsive view kwa review na entry, hakikisha connection kwa writes na subiri response iliyothibitishwa kabla ya kuondoka kwenye action muhimu."), ("Support handoff", "Record company, module, time, user role, exact message, and safe reproduction steps without including passwords or tokens.", "Andika kampuni, module, muda, role, ujumbe kamili na hatua salama za kurudia bila kuingiza password au token."),]
    for i, (name, en, sw) in enumerate(user_steps, 1):
        out += [f"## VI.{i} {name}", "", en, "", f"### Kiswahili", "", sw, "", "#### Practical control", "", "Treat the server-confirmed response as the completion signal. If it is unavailable, keep the operation visible as pending, failed, or unavailable rather than completed.", "", "---", ""]
    out += ["# PART VII — ADMINISTRATOR MANUAL", ""]
    admin_steps = [("Provision a company", "Confirm the authenticated owner, create or join the company through the supported path, verify membership and role, and review module entitlements before inviting staff.", "Thibitisha owner, unda au jiunge na kampuni, hakikisha membership na role, kagua entitlements kabla ya kualika wafanyakazi."), ("Manage roles", "Use role-change approval where required. Review the target user, old role, new role, approver, and audit record.", "Tumia role-change approval inapohitajika. Kagua mtumiaji, role ya zamani/mpya, approver na audit record."), ("Configure providers", "Keep provider keys server-side, test readiness, expose disabled states, and document which environment is sandbox or production.", "Hifadhi keys server-side, pima readiness, onyesha disabled state na andika mazingira ya sandbox au production."), ("Review RLS", "For each table, identify the company boundary, intended commands, membership helper, and policy. Avoid broad authenticated policies that bypass tenant scope.", "Kwa kila table tambua company boundary, commands, helper ya membership na policy. Epuka policy pana ya authenticated inayovunja tenant scope."), ("Run migrations", "Use source-versioned migrations, review prerequisites and compatibility guards, back up the target environment, apply through the controlled Supabase migration path, then rerun schema and RLS checks.", "Tumia migrations zenye version, kagua prerequisites na guards, hifadhi backup, apply kupitia njia iliyodhibitiwa ya Supabase, kisha rudia schema/RLS checks."), ("Respond to incidents", "Preserve audit evidence, fail closed, rotate exposed credentials, isolate the affected company or provider, and record the recovery decision.", "Hifadhi audit, kata access pale inavyohitajika, badilisha credentials zilizoonekana, tenga kampuni/provider iliyoathirika na andika uamuzi wa recovery."),]
    for i, (name, en, sw) in enumerate(admin_steps, 1):
        out += [f"## VII.{i} {name}", "", en, "", "### Kiswahili", "", sw, "", "#### Administrator evidence", "", "An administrator should be able to show the source migration, the live ledger, the policy boundary, the successful verification result, and the remaining limitation. A green UI alone is not administration evidence.", "", "---", ""]
    out += ["# PART VIII — DEVELOPER MANUAL", ""]
    dev_steps = [("Source navigation", "Start with client/src/App.tsx, BusinessSphereDashboard.jsx, dashboard registries, server/_core/apiApp.ts, server/routers.ts, feature operations, and supabase/migrations. Follow an action from UI to handler/RPC to table and back to confirmed state.", "Anza na App.tsx, BusinessSphereDashboard.jsx, registries, apiApp.ts, routers.ts, operations na migrations. Fuatilia action kutoka UI hadi handler/RPC, table na response iliyothibitishwa."), ("Persistence contract", "Critical writes should validate input, resolve verified profile, enforce company scope, call the server/database boundary, require returned row/result, and retain retryable input on failure.", "Writes muhimu zithibitishe input, profile, company scope, server/database, response ya row/result na zibaki na form kwa retry ikishindikana."), ("Migration discipline", "Use timestamped source migrations. Prefer additive and guarded changes. Never silently coerce incompatible production columns or create a duplicate architecture around auth.users/profiles.", "Tumia migrations zenye timestamp. Pendelea additive guards. Usibadilishe column za production kimya kimya wala kuunda users/profiles duplicate."), ("Testing", "Use contract tests for source promises, Vitest for server logic, Playwright for browser journeys, and controlled Supabase staging for authenticated CRUD and provider acceptance.", "Tumia contract tests, Vitest, Playwright na staging ya Supabase kwa CRUD ya authenticated na acceptance ya providers."), ("Observability", "Log safe error categories, request correlation where available, audit outcomes, provider order IDs, and schema-verification status. Never log credentials or raw private payloads.", "Log categories salama za errors, correlation, audit outcomes, provider order IDs na schema status. Usi-log credentials au payload za faragha."), ("Extension strategy", "Add new modules through a route/registry decision, server boundary, schema contract, RLS policy, migration, tests, and documentation status. Keep the module’s limitation visible until all evidence exists.", "Ongeza module kupitia route/registry, server boundary, schema contract, RLS, migration, tests na status ya docs. Usifiche limitation kabla ya ushahidi kamili."),]
    for i, (name, en, sw) in enumerate(dev_steps, 1):
        out += [f"## VIII.{i} {name}", "", en, "", "### Kiswahili", "", sw, "", "#### Developer checklist", "", "Source, database, API, UI, RLS, tests, build, and deployment evidence must agree before a feature is described as complete.", "", "---", ""]
    out += ["# PART IX — SECURITY BOOK", "", "The security model is layered and the risk register below records actual findings from the repository and live advisors.", ""]
    for i, (finding, severity, evidence, fix) in enumerate(SECURITY_FINDINGS, 1):
        out += [f"## IX.{i} {finding}", "", f"**Severity:** {severity}", "", f"**Evidence:** {evidence}", "", f"**Recommended fix:** {fix}", "", "### Security principle", "", "The platform should prefer explicit identity, narrow role authority, company scope, server confirmation, database enforcement, and auditable results. A local UI state, a draft message, or a client-supplied amount cannot replace those controls.", "", "### Kiswahili", "", f"**Ushahidi:** {evidence} **Hatua:** {fix} Mfumo unapaswa kutumia utambulisho ulio wazi, role yenye mipaka, company scope, confirmation ya server, enforcement ya database na audit.", "", "---", ""]
    out += ["# PART X — DATABASE DICTIONARY", "", f"The live read-only Supabase inventory returned {LIVE['table_count']} tables ({LIVE['public_count']} public and {LIVE['auth_count']} auth), with {LIVE['rls_enabled']} reporting RLS enabled. The following dictionary records the public table surface without exposing row contents.", ""]
    public_tables = [t for t in LIVE["tables"] if t.get("name", "").startswith("public.")]
    rows = []
    for t in sorted(public_tables, key=lambda x: x.get("name", "")):
        short = t.get("name", "").split(".", 1)[-1]
        cols = ", ".join(c.get("name", "") for c in t.get("columns", []))
        rows.append([short, cols[:420], ", ".join(t.get("primary_keys", [])) or "—", "Enabled" if t.get("rls_enabled") else "Review", str(t.get("rows", 0))])
    out += [md_table(["Table", "Observed columns", "Primary key", "RLS", "Rows reported"], rows), ""]
    out += ["## Canonical identity and tenancy tables", "", md_table(["Table", "Purpose", "Security boundary", "Used by"], [["auth.users", "Supabase Auth identity", "Auth session and token boundary", "PublicAuthGateway and profile resolution"], ["profiles", "User/company profile and role context", "Authenticated self-service and company scope", "Most protected workflows"], ["companies", "Organization/company identity", "Tenant boundary", "All company-scoped modules"], ["company_memberships", "Membership and role relationship", "Company plus user plus role/status", "Invitations, onboarding, authorization"], ["workspaces", "Workspace context where present", "Workspace-aware navigation", "Shell and module context"], ["user_table_preferences", "User preference persistence", "Self-only or company-scoped policy", "Dashboard/profile settings"]]), ""]
    out += ["## Canonical subscription tables", "", md_table(["Table", "Purpose", "Contract", "Evidence"], [["billing_plans", "Package catalog", "Free and paid plan metadata, entitlements", "Subscription migrations and live inventory"], ["billing_profiles", "Billing contact/configuration", "Company-scoped billing profile", "Billing foundation migration"], ["tenant_subscriptions", "Company subscription state", "Pending/Active/Grace/Expired/RequiresPlan/Cancelled/Superseded", "Subscription model migrations"], ["subscription_payments", "Provider payment state", "Monthly cycle, idempotency, provider order", "HarakaPay handlers and migrations"], ["subscription_invoices", "Billing invoice evidence", "Payment/subscription linkage", "Billing foundation migration"], ["subscription_usage", "Usage/limits evidence", "Company and plan context", "Billing foundation migration"], ["subscription_events", "Billing lifecycle events", "Audit and reconciliation history", "Billing migrations"], ["subscription_notifications", "Billing notifications", "Company-scoped notification state", "Billing migrations"], ["billing_access_snapshot", "Authoritative access result", "Server/database snapshot; not a table", "Access adapter and protected API"]]), ""]
    out += ["# PART XI — INTEGRATIONS", ""]
    integration_rows = [["Supabase", "Auth, Postgres/REST, RPC, RLS", "Live schema and migration evidence present", "Service keys server-side; client uses public anon configuration only"], ["HarakaPay", "USSD collection, status, webhook, balance", "Server handlers present; provider readiness required", "API key, order matching, server verification, idempotency"], ["TRA/VFD", "Fiscal profile, receipts, retries, Z reports", "Router and provider boundary present; external readiness required", "TIN/VRN, branch/device/configuration, audit"], ["Storage/S3", "Avatars, property documents, exports", "Storage helper and proxy present", "Scoped keys, validation, signed/download boundaries"], ["Email/SMS/WhatsApp", "Notifications and customer/workforce communication", "Persistence and provider boundaries present", "No claim of delivery without provider configuration"], ["AI model", "Structured assistant responses and proposals", "Built-in model integration present", "Context limits, JSON schema, proposal-only mutation, approvals"], ["Maps", "Location support for property/fleet where configured", "UI component and route boundary present", "API key/provider configuration required"], ["Vercel", "Deployment/runtime", "Build and serverless import contract present", "Environment variables and schema verification"],]
    out += [md_table(["Integration", "Purpose", "Current evidence", "Security/operational boundary"], integration_rows), ""]
    for i, row in enumerate(integration_rows, 1):
        out += [f"## XI.{i} {row[0]}", "", f"**English:** {row[1]}. {row[2]} {row[3]}", "", f"**Kiswahili:** {row[0]} hutumika kwa {row[1].lower()}. {row[2]}. {row[3]}.", "", "---", ""]
    out += ["# PART XII — TROUBLESHOOTING", ""]
    out += [md_table(["Problem", "Possible cause", "Diagnosis", "Solution and prevention"], [[a,b,c,d] for a,b,c,d in TROUBLESHOOTING]), ""]
    for i, (problem, cause, diag, sol) in enumerate(TROUBLESHOOTING, 1):
        out += [f"## XII.{i} {problem}", "", f"**Possible cause:** {cause}", "", f"**Diagnosis:** {diag}", "", f"**Solution:** {sol}", "", "### Kiswahili", "", f"**Sababu:** {cause} **Utambuzi:** {diag} **Suluhisho:** {sol}", "", "---", ""]
    out += ["# PART XIII — OPERATIONS", ""]
    operations = [("Release readiness", "Confirm source migration, tests, TypeScript, build, browser journeys, live schema, RLS, provider configuration, and rollback evidence.", "Thibitisha migration, tests, TypeScript, build, browser journeys, schema, RLS, provider na rollback."), ("Backup and recovery", "Distinguish database reachability from managed PITR or snapshot configuration. Do not report backup controls that were not checked.", "Tenganisha database reachability na PITR/snapshot. Usiripoti backup ambayo haikuthibitishwa."), ("Monitoring", "Watch API status, scheduled handlers, webhook deliveries, provider errors, failed mutations, security advisors, and schema drift.", "Fuatilia API, scheduled handlers, webhooks, provider errors, failed writes, security advisors na schema drift."), ("Support handover", "Keep the repository, migration ledger, environment checklist, role matrix, error reference, and incident notes together.", "Hifadhi repository, migration ledger, env checklist, role matrix, error reference na incident notes pamoja."), ("Data governance", "Define retention, privacy, least privilege, export, deletion, and incident escalation by business and regulatory context.", "Weka retention, privacy, least privilege, export, deletion na escalation kulingana na biashara na compliance."), ("Change control", "Every schema, role, provider, and entitlement change needs a source commit, migration/version record, test evidence, and updated status.", "Kila mabadiliko ya schema, role, provider na entitlement yanahitaji commit, migration/version, tests na status mpya."),]
    for i,(name,en,sw) in enumerate(operations,1):
        out += [f"## XIII.{i} {name}", "", en, "", "### Kiswahili", "", sw, "", "---", ""]
    out += ["# PART XIV — FUTURE ROADMAP", ""]
    roadmap = [("Atomic legacy invoice payment RPC", "Replace the separate payment insert and invoice update with a reviewed atomic, idempotent database boundary."), ("Broader authenticated CRUD journeys", "Expand role-by-role Playwright and controlled Supabase staging coverage across legacy modules."), ("Signature-specific SECURITY DEFINER hardening", "Review each exposed signature, preserve intentional endpoints, and revoke or relocate internal helpers."), ("RLS policy consolidation", "Reduce overlapping permissive policies after proving equivalent tenant semantics."), ("Provider readiness", "Complete approved HarakaPay, TRA/VFD, storage, email/SMS/WhatsApp, maps, and AI configuration and acceptance."), ("Dashboard decomposition", "Extract safe, high-risk boundaries and code-split the large dashboard only after persistence and live blockers are addressed."),]
    for i,(name,en) in enumerate(roadmap,1):
        out += [f"## XIV.{i} {name}", "", en, "", "### Kiswahili", "", f"Hatua hii inapaswa kufanywa baada ya uthibitisho wa source, data, ruhusa na deployment. Lengo si kuongeza madai ya bidhaa, bali kuimarisha ushahidi wa kazi iliyopo: {en}", "", "---", ""]
    out += ["# PART XV — SWAHILI EDITION", "", "The following is the complete professional Tanzanian Swahili module reference. Technical English terms remain in parentheses where they reduce ambiguity.", ""]
    for idx, m in enumerate(MODULES, 1):
        out += [f"## XV.{idx} {m['sw']}", "", "### Muhtasari wa Moduli", "", *swahili_module_paragraphs(m), "", "### Madhumuni na Changamoto", "", f"{m['sw_focus']} Moduli hii inalenga kupunguza rekodi zilizotawanyika, kuongeza uwazi wa kazi na kuweka mpaka wa uwajibikaji.", "", "### Watumiaji, Urambazaji na Vipengele", "", f"Majukumu yanayohusika: {m['roles']}. UI iliyothibitishwa: `{m['ui']}`. Vipengele muhimu vinafuata mtiririko huu: {m['workflow']}", "", "### Hifadhidata, Ruhusa na Usalama", "", f"Server boundary: `{m['server']}`. Tables zinazohusiana zinazonekana kwenye inventory ni: {', '.join(table_names(m['prefixes'])[:18]) or 'tables za shared company scope'}. Access inapaswa kupita kwenye profile iliyothibitishwa, company scope, role, database constraints na RLS/policies.", "", "### Ripoti, Integrations na Mobile", "", "Ripoti inategemea data iliyorekodiwa. Integration ya nje inahitaji configuration na credentials zilizoidhinishwa. Kwenye simu, mtumiaji athibitishe response kabla ya kuondoka kwenye action muhimu.", "", "### Mipaka, Hali ya Sasa na Maboresho", "", f"Hali ya sasa: **{m['status']}**. Mpaka unaojulikana: {m['limit']} Maboresho yanayofuata ni CRUD ya authenticated yenye data iliyodhibitiwa, tests za role/workflow, provider readiness au decomposition salama kulingana na hitaji.", "", "### Mwongozo wa Uendeshaji", "", *[f"**{step}:** {instruction}" for step, instruction in operator_runbook_sw(m)], "", "### Jedwali la Ushahidi wa Udhibiti", "", md_table(["Hatua ya udhibiti", "Ushahidi unaotarajiwa", "Swali la review"], [["Utambulisho", "Profile na company iliyothibitishwa", "Nani anafanya action?"], ["Input", "Fields na references sahihi", "Data inakubalika?"], ["Persistence", "ID, status au RPC result", "Imehifadhiwa kweli?"], ["Control", "Approval, reconciliation au audit", "Nani amekagua?"], ["Integration", "Readiness na delivery outcome", "Provider amethibitisha nini?"]]), "", "### Exception na Recovery", "", f"Ikiwa {m['sw']} ina hali ya unavailable, failed, pending au rejected, hifadhi hali na input, nakili response, hakikisha company na role, kisha tumia server/database path inayoungwa mkono. Usibadilishe local optimistic state kuwa madai ya completion.", "", "---", ""]
    out += ["# APPENDICES", "", "## Appendix A — Role and permission matrix", "", md_table(["Role", "Dashboard", "Primary modules", "Boundary"], [[a, b, c, d] for a,b,c,d in ROLE_ROWS]), "", "## Appendix B — Master feature status matrix", "", md_table(["Module", "UI evidence", "Server/API evidence", "Status", "Database prefix"], [[m['name'], m['ui'], m['server'], m['status'], ', '.join(table_names(m['prefixes'])[:8]) or 'shared'] for m in MODULES]), "", "## Appendix C — Glossary", "", md_table(["Term", "English definition", "Kiswahili"], [[a,b,c] for a,b,c in GLOSSARY]), "", "## Appendix D — Verification record", "", f"Repository version: `1.0.0` from `package.json`. Documentation date: {TODAY}. Live Supabase project audit: read-only. Live table count: {LIVE['table_count']}; public: {LIVE['public_count']}; auth: {LIVE['auth_count']}; RLS enabled: {LIVE['rls_enabled']}; RLS not enabled: {LIVE['rls_disabled']}; migration records: {LIVE['migration_count']}. Security advisor lints: {len(LIVE['security_lints'])}; performance advisor lints: {len(LIVE['performance_lints'])}.", "", "## References", "", "[1]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/client/src/App.tsx \"Application routes and providers\"", "[2]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/client/src/BusinessSphereDashboard.jsx \"Authenticated dashboard and module composition\"", "[3]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/server/_core/apiApp.ts \"Express API bootstrap and HTTP routes\"", "[4]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/server/subscriptionBilling.ts \"Subscription and HarakaPay server handlers\"", "[5]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/server/propertyManagementOperations.ts \"Property Management validation and RPC boundary\"", "[6]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/server/smartAssistant.ts \"AI assistant prompt, limits, and structured proposal contract\"", "[7]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/server/traFiscalRouter.ts \"TRA/VFD fiscal router\"", "[8]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/FULL_SYSTEM_AUDIT_REPORT.md \"Full-system audit and known blockers\"", "[9]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/FULL_SYSTEM_IMPLEMENTATION_MATRIX.md \"Feature implementation matrix\"", "[10]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/main/package.json \"Runtime and build scripts\"", "[11]: https://supabase.com/docs/guides/database/postgres/row-level-security \"Supabase Row Level Security documentation\"", "", "## Documentation QA statement", "The book was generated from repository files and read-only live evidence. It does not expose production credentials, does not fabricate screenshots, does not create a duplicate auth/users architecture, and explicitly records known limitations and external prerequisites. The render and QA reports are retained beside the deliverables.", ""]
    rendered = "\n".join(out).replace("\n\n\n", "\n\n")
    return "\n".join(line.rstrip() for line in rendered.splitlines()).rstrip() + "\n"


def build_typst() -> None:
    lines: list[str] = []
    lines.append('#import "report-theme.typ": report-accent, report-theme')
    lines.append('#show: report-theme.with(title: "SMART MANAGER ERP — Master System Book", author: "Manus AI", rhythm: "longform", running-header: true)')
    lines.append('#set text(font: ("Libertinus Serif", "Noto Sans"), lang: "en")')
    lines.append('#set par(justify: true)')
    lines.append('')
    lines.append(f'#page(fill: rgb("#061A13"), margin: (top: 1.2cm, bottom: 1.2cm, x: 2cm), numbering: none, header: none)[#align(center)[#v(1cm)#image("{MASTER_BOOK_LOGO.as_posix()}", width: 15cm)#v(0.8cm)#text(fill: white, size: 25pt, weight: "bold")[SMART MANAGER ERP]#v(0.35cm)#text(fill: rgb("#D6EADF"), size: 14pt)[MASTER SYSTEM BOOK \\ ENGLISH AND TANZANIAN SWAHILI EDITION]#v(1cm)#line(length: 70%, stroke: 1pt + rgb("#D6B36A"))#v(0.8cm)#text(fill: white, size: 11pt)[Your Business. Connected. Controlled. Intelligent.]#v(1.1cm)#text(fill: white, size: 11pt)[Repository-Audited Edition \\ Version 1.0.0 \\ Documentation date: 24 August 2026 \\ Prepared by Manus AI \\ Owner and Creator: Ezra Mpapi \\ Dar es Salaam, Tanzania]]]')
    lines.append(f'#page(numbering: none, header: none)[#align(center)[#text(size: 19pt, weight: "bold", fill: report-accent)[The Vision Behind SMART MANAGER]#v(0.7cm)#image("{MASTER_BOOK_OWNER.as_posix()}", width: 7cm)#v(0.5cm)#text(size: 10.5pt)[Ezra Mpapi is identified in the supplied project materials as the Owner and Creator of SMART MANAGER ERP. This book uses only verified identity information and does not add an unverified biography, qualification, award, or employment history. \\ \\ SMART MANAGER is documented here as a repository-audited product whose source, database, authentication, security, subscription, and operational boundaries are explicitly described. \\ \\ Dar es Salaam, Tanzania]] ]')
    lines.append('#page(numbering: none, header: none)[#outline(title: [Contents / Yaliyomo], indent: 1.5em)]')
    lines.append('#counter(page).update(1)')
    lines.append('= Editorial and Evidence Note')
    lines.append('This is a repository-audited master book. English is the primary technical reference; the complete Tanzanian Swahili edition follows the English reference. Status labels are evidence boundaries, not marketing claims. The live Supabase audit was read-only and its snapshots are retained in the evidence directory.')
    lines.append(typst_table(["Metric", "Value", "Interpretation", "Evidence"], [["Project version", "1.0.0", "Package version", "package.json"], ["Live tables", str(LIVE["table_count"]), f"{LIVE['public_count']} public; {LIVE['auth_count']} auth", "Read-only Supabase inventory"], ["RLS", str(LIVE["rls_enabled"]), f"{LIVE['rls_disabled']} not enabled in returned inventory", "Read-only Supabase inventory"], ["Migrations", str(LIVE["migration_count"]), "Records in live ledger", "Read-only Supabase migration list"], ["Security lints", str(len(LIVE["security_lints"])), "WARN/INFO advisor records", "Supabase security advisor"], ["Performance lints", str(len(LIVE["performance_lints"])), "WARN/INFO advisor records", "Supabase performance advisor"]]))
    lines.append('#pagebreak()')
    lines.append('= Part I — Product Introduction')
    for i, (title, sw_title, sw_intro, en) in enumerate(INTRO_CHAPTERS, 1):
        lines.append(f'== I.{i} {typst_escape(title)}')
        lines.append(typst_escape(en))
        lines.append(f'=== {typst_escape(sw_title)}')
        lines.append(typst_escape(sw_intro))
        lines.append('')
    lines.append('#pagebreak()')
    lines.append('= Part II — System Overview')
    overview = [("What is SMART MANAGER?", "SMART MANAGER is an authenticated business operating platform whose verified repository surfaces connect commercial, operational, financial, people, sector, intelligence, and control workflows.", "SMART MANAGER ni platform ya uendeshaji wa biashara yenye uthibitishaji inayounganisha workflows za biashara, operations, fedha, watu, sekta, intelligence na udhibiti."), ("ERP concept", "ERP is treated as a connected record and control model: a sale may relate to stock, finance, customer history, receipt, and report when relevant contracts are enabled.", "ERP inaeleweka kama mfumo wa rekodi na udhibiti uliounganishwa: mauzo yanaweza kuhusishwa na stock, fedha, historia ya mteja, receipt na ripoti pale contracts zimewezeshwa."), ("System ecosystem", "The ecosystem includes public entry, secure authentication, company/workspace context, role-aware modules, Supabase persistence, server APIs, scheduled handlers, storage, external providers, and audit evidence.", "Ecosystem ina entry ya umma, authentication, company/workspace, modules za role, Supabase, APIs za server, scheduled handlers, storage, providers wa nje na audit."), ("Supported business types", "The verified module registry spans universal SME operations and specialist contexts such as property, healthcare, school, pharmacy, hospitality, restaurant, fleet, microfinance, banking, VICOBA, and money-agent operations.", "Rejista ya modules inahusisha SME operations na sekta kama property, afya, shule, famasi, hoteli, mgahawa, fleet, microfinance, benki, VICOBA na wakala wa fedha."), ("Lifecycle model", "A safe lifecycle is discover → authenticate → establish company scope → record → validate → approve/reconcile → report → audit → improve.", "Lifecycle salama ni gundua → authenticate → weka company scope → rekodi → thibitisha → approve/reconcile → ripoti → audit → boresha.")]
    for i,(title,en,sw) in enumerate(overview,1):
        lines.append(f'== II.{i} {typst_escape(title)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(sw))
        lines.append('')
    add_image_typst(lines, 'screenshots/live-homepage.png', 'Verified captured public landing page; used as product-introduction evidence.')
    add_image_typst(lines, 'screenshots/live-workspace-entry.png', 'Verified captured workspace-entry and authentication experience.')
    lines.append('#pagebreak()')
    lines.append('= Part III — Technical Architecture')
    for i,(title,en,sw) in enumerate(ARCH_CHAPTERS,1):
        lines.append(f'== III.{i} {typst_escape(title)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(sw))
        if title == 'Technology Stack':
            lines.append(typst_table(["Layer", "Implementation", "Evidence", "Boundary"], [["Frontend", "Vite + React 19 + Tailwind + wouter", "package.json; App.tsx", "Browser is not database authority"], ["Backend", "Express 5 + tRPC 11 + handlers", "apiApp.ts; routers.ts", "Server validates and scopes"], ["Data", "Supabase Auth/Postgres/REST/RLS", "migrations and live audit", "RLS and constraints"], ["Runtime", "Vercel-compatible import plus local server", "index.ts; build scripts", "Environment configuration"]]))
        if title == 'Database Architecture':
            lines.append(typst_table(["Metric", "Value", "Meaning", "Evidence"], [["Total tables", str(LIVE["table_count"]), "Live schema metadata", "Read-only list_tables"], ["Public", str(LIVE["public_count"]), "Public schema tables", "Read-only list_tables"], ["Auth", str(LIVE["auth_count"]), "Auth schema tables", "Read-only list_tables"], ["RLS enabled", str(LIVE["rls_enabled"]), "Reported enabled", "Read-only list_tables"], ["RLS not enabled", str(LIVE["rls_disabled"]), "Specific review required", "Read-only list_tables"], ["Migrations", str(LIVE["migration_count"]), "Ledger records", "Read-only list_migrations"]]))
        lines.append('')
    add_image_typst(lines, 'diagrams/live-module-map.png', 'Live module-map evidence from the repository architecture review.')
    add_image_typst(lines, 'diagrams/live-auth-workflow.png', 'Authentication and tenant-aware workflow evidence.')
    add_image_typst(lines, 'diagrams/canonical-data-model.png', 'Canonical data-model relationship evidence.')
    add_image_typst(lines, 'diagrams/integration-topology.png', 'Integration topology evidence.')
    lines.append('#pagebreak()')
    lines.append('= Part IV — Complete Module Book')
    lines.append('Each module chapter follows the required documentation structure and is based on real UI, server, migration, and live inventory evidence. A module status is not a guarantee that every deployment has every provider or data prerequisite configured.')
    for idx,m in enumerate(MODULES,1):
        lines.append(f'== IV.{idx} {typst_escape(m["name"])}')
        lines.append(f'*{typst_escape(m["status"])}*')
        lines.append('=== Module overview')
        for p in english_module_paragraphs(m): lines.append(typst_escape(p))
        lines.append('=== Purpose, users, navigation, and features')
        lines.append(typst_escape(f"Primary roles: {m['roles']}. UI evidence: {m['ui']}. {m['focus']}"))
        lines.append('=== Database, permissions, workflow, and reports')
        lines.append(typst_table(["Area", "Verified reference", "Operational interpretation", "Limitation"], module_rows(m)))
        lines.append('=== Security and mobile experience')
        lines.append(typst_escape("Use verified profile resolution, company scope, role checks, database constraints, RLS/policies, and confirmed server responses. Responsive UI changes layout, not authorization. A failed or unavailable write remains failed or unavailable."))
        lines.append('=== Known limitations and future improvements')
        lines.append(typst_escape(m["limit"]))
        lines.append(typst_escape("The next improvement is evidence-driven: complete controlled authenticated CRUD, provider readiness, role-specific regression, or safe code decomposition."))
        if m.get('image'):
            add_image_typst(lines, m['image'], f"{m['name']} evidence image from repository or captured application assets.", width="12.5cm")
        lines.append('#pagebreak()')
    lines.append('= Part V — Business Workflows')
    for i,(name,en,sw) in enumerate(WORKFLOWS,1):
        lines.append(f'== V.{i} {typst_escape(name)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(sw))
        lines.append('=== Control points')
        lines.append(typst_escape("At each transition, confirm identity, company scope, role, input validation, database response, and audit/result state. A failed step must remain failed; a local optimistic row is not a durable business record."))
        lines.append('')
    lines.append('#pagebreak()')
    lines.append('= Part VI — User Manual')
    user_steps_typ = [
        ("First-day onboarding", "Sign in through the supported gateway, verify the company context, review your role, open one permitted module, and create only a record that the server confirms.", "Siku ya kwanza ingia kupitia gateway, hakikisha company context, soma role yako, fungua module inayoruhusiwa na tengeneza record ambayo server imethibitisha."),
        ("Daily operating rhythm", "Start with the dashboard, inspect exceptions, execute the module workflow, verify returned state, and leave approvals or reconciliation items visible to the responsible role.", "Anza na dashibodi, kagua exceptions, tekeleza workflow, thibitisha hali iliyorejeshwa na acha approvals/reconciliation zionekane kwa role husika."),
        ("Safe data entry", "Use required fields, valid dates and amounts, company-scoped references, and clear descriptions. Do not paste secrets into business notes.", "Tumia fields zinazohitajika, tarehe na kiasi sahihi, references za kampuni na maelezo wazi. Usiweke siri ndani ya notes za biashara."),
        ("Errors and retries", "Read the error state, preserve the form, correct the input or permission, and retry only after the cause is understood. Do not click repeatedly on payment or approval actions without an idempotency path.", "Soma error, hifadhi form, rekebisha input au permission na retry baada ya kuelewa sababu. Usibonyeze payment au approval mara nyingi bila idempotency."),
        ("Reports and review", "Use reports to ask a question of recorded data. Follow a surprising number back to the source document, payment, movement, or journal entry.", "Tumia ripoti kuuliza swali la data iliyorekodiwa. Namba ya kushangaza ifuatilie hadi document, payment, movement au journal entry ya chanzo."),
        ("Subscription access", "Read the server access state. FREE_15 is free for 15 days; paid access requires verified payment. Do not use local browser state as entitlement authority.", "Soma access state ya server. FREE_15 ni bure kwa siku 15; paid access inahitaji payment iliyothibitishwa. Browser si mamlaka ya entitlement."),
        ("Working on mobile", "Use responsive views for review and entry, maintain stable connectivity for writes, and wait for confirmed response before navigating away from a critical action.", "Tumia responsive views kwa review na entry, hakikisha connection kwa writes na subiri response iliyothibitishwa kabla ya kuondoka kwenye action muhimu."),
        ("Support handoff", "Record company, module, time, user role, exact message, and safe reproduction steps without including passwords or tokens.", "Andika kampuni, module, muda, role, ujumbe kamili na hatua salama za kurudia bila kuingiza password au token."),
    ]
    for i,(title,en,sw) in enumerate(user_steps_typ,1):
        lines.append(f'== VI.{i} {typst_escape(title)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(sw))
        lines.append('=== Practical control')
        lines.append(typst_escape("Treat the server-confirmed response as the completion signal. If it is unavailable, keep the operation visible as pending, failed, or unavailable rather than completed."))
        lines.append('')
    lines.append('#pagebreak()')
    lines.append('= Part VII — Administrator Manual')
    admin_steps_typ = [
        ("Provision a company", "Confirm the authenticated owner, create or join the company through the supported path, verify membership and role, and review module entitlements before inviting staff.", "Thibitisha owner, unda au jiunge na kampuni, hakikisha membership na role, kagua entitlements kabla ya kualika wafanyakazi."),
        ("Manage roles", "Use role-change approval where required. Review the target user, old role, new role, approver, and audit record.", "Tumia role-change approval inapohitajika. Kagua mtumiaji, role ya zamani/mpya, approver na audit record."),
        ("Configure providers", "Keep provider keys server-side, test readiness, expose disabled states, and document which environment is sandbox or production.", "Hifadhi keys server-side, pima readiness, onyesha disabled state na andika mazingira ya sandbox au production."),
        ("Review RLS", "For each table, identify the company boundary, intended commands, membership helper, and policy. Avoid broad authenticated policies that bypass tenant scope.", "Kwa kila table tambua company boundary, commands, helper ya membership na policy. Epuka policy pana ya authenticated inayovunja tenant scope."),
        ("Run migrations", "Use source-versioned migrations, review prerequisites and compatibility guards, back up the target environment, apply through the controlled Supabase migration path, then rerun schema and RLS checks.", "Tumia migrations zenye version, kagua prerequisites na guards, hifadhi backup, apply kupitia njia iliyodhibitiwa ya Supabase, kisha rudia schema/RLS checks."),
        ("Respond to incidents", "Preserve audit evidence, fail closed, rotate exposed credentials, isolate the affected company or provider, and record the recovery decision.", "Hifadhi audit, kata access pale inavyohitajika, badilisha credentials zilizoonekana, tenga kampuni/provider iliyoathirika na andika uamuzi wa recovery."),
    ]
    for i,(title,en,sw) in enumerate(admin_steps_typ,1):
        lines.append(f'== VII.{i} {typst_escape(title)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(sw))
        lines.append('=== Administrator evidence')
        lines.append(typst_escape("An administrator should be able to show the source migration, the live ledger, the policy boundary, the successful verification result, and the remaining limitation. A green UI alone is not administration evidence."))
        lines.append('')
    lines.append('#pagebreak()')
    lines.append('= Part VIII — Developer Manual')
    dev_steps_typ = [
        ("Source navigation", "Start with client/src/App.tsx, BusinessSphereDashboard.jsx, dashboard registries, server/_core/apiApp.ts, server/routers.ts, feature operations, and supabase/migrations. Follow an action from UI to handler/RPC to table and back to confirmed state.", "Anza na App.tsx, BusinessSphereDashboard.jsx, registries, apiApp.ts, routers.ts, operations na migrations. Fuatilia action kutoka UI hadi handler/RPC, table na response iliyothibitishwa."),
        ("Persistence contract", "Critical writes should validate input, resolve verified profile, enforce company scope, call the server/database boundary, require returned row/result, and retain retryable input on failure.", "Writes muhimu zithibitishe input, profile, company scope, server/database, response ya row/result na zibaki na form kwa retry ikishindikana."),
        ("Migration discipline", "Use timestamped source migrations. Prefer additive and guarded changes. Never silently coerce incompatible production columns or create a duplicate architecture around auth.users/profiles.", "Tumia migrations zenye timestamp. Pendelea additive guards. Usibadilishe column za production kimya kimya wala kuunda users/profiles duplicate."),
        ("Testing", "Use contract tests for source promises, Vitest for server logic, Playwright for browser journeys, and controlled Supabase staging for authenticated CRUD and provider acceptance.", "Tumia contract tests, Vitest, Playwright na staging ya Supabase kwa CRUD ya authenticated na acceptance ya providers."),
        ("Observability", "Log safe error categories, request correlation where available, audit outcomes, provider order IDs, and schema-verification status. Never log credentials or raw private payloads.", "Log categories salama za errors, correlation, audit outcomes, provider order IDs na schema status. Usi-log credentials au payload za faragha."),
        ("Extension strategy", "Add new modules through a route/registry decision, server boundary, schema contract, RLS policy, migration, tests, and documentation status. Keep the module’s limitation visible until all evidence exists.", "Ongeza module kupitia route/registry, server boundary, schema contract, RLS, migration, tests na status ya docs. Usifiche limitation kabla ya ushahidi kamili."),
    ]
    for i,(title,en,sw) in enumerate(dev_steps_typ,1):
        lines.append(f'== VIII.{i} {typst_escape(title)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(sw))
        lines.append('=== Developer control')
        lines.append(typst_escape("Source, database, API, UI, RLS, tests, build, and deployment evidence must agree before a feature is described as complete."))
        lines.append('')
    lines.append('#pagebreak()')
    lines.append('= Part IX — Security Book')
    lines.append(typst_escape("The security model is layered and the risk register below records actual findings from the repository and live advisors."))
    for i,(finding,severity,evidence,fix) in enumerate(SECURITY_FINDINGS,1):
        lines.append(f'== IX.{i} {typst_escape(finding)}')
        lines.append(f'*{typst_escape(severity)}*')
        lines.append(typst_escape(f"Evidence: {evidence}"))
        lines.append(typst_escape(f"Recommended fix: {fix}"))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(f"Ushahidi: {evidence} Hatua inayopendekezwa: {fix} Mfumo unapaswa kutumia utambulisho ulio wazi, role yenye mipaka, company scope, confirmation ya server, enforcement ya database na audit."))
        lines.append(typst_table(["Control", "Expected behavior", "Review question", "Evidence"], [["Identity", "Verified profile and supported token", "Who is the caller?", "auth/profile boundary"], ["Tenant", "Company-scoped record access", "Does the row belong to the company?", "company_id/RLS"], ["Role", "Specific permitted action", "Is this role allowed?", "server helper/policy"], ["Mutation", "Confirmed server/database result", "Did the write really complete?", "returned row/status"], ["Audit", "Safe activity evidence", "Can the outcome be traced?", "audit table/log"]]))
        lines.append('#pagebreak()')
    lines.append('= Part X — Database Dictionary')
    lines.append(typst_escape(f"The live read-only Supabase inventory returned {LIVE['table_count']} tables ({LIVE['public_count']} public and {LIVE['auth_count']} auth), with {LIVE['rls_enabled']} reporting RLS enabled and {LIVE['rls_disabled']} not enabled in the returned inventory. The following table index is evidence metadata only; it does not expose row contents."))
    public_tables = [t for t in LIVE["tables"] if t.get("name", "").startswith("public.")]
    # Split the dictionary into deliberately readable segments so the rendered page count reflects real documentation rather than one unreadable table.
    for chunk_start in range(0, len(public_tables), 24):
        chunk = sorted(public_tables, key=lambda x: x.get("name", ""))[chunk_start:chunk_start+24]
        rows = []
        for t in chunk:
            short = t.get("name", "").split(".", 1)[-1]
            cols = ", ".join(c.get("name", "") for c in t.get("columns", []))
            rows.append([short, cols[:300], ", ".join(t.get("primary_keys", [])) or "—", "Enabled" if t.get("rls_enabled") else "Review"])
        lines.append(f'== X.{chunk_start//24+1} Public table segment {chunk_start+1}–{chunk_start+len(chunk)}')
        lines.append(typst_table(["Table", "Observed columns", "Primary key", "RLS"], rows, widths="(3.5cm, 8cm, 3cm, 2cm)"))
        lines.append(typst_escape("The index is a documentation aid. It does not prescribe direct client writes; server boundaries, constraints, RPCs, and RLS remain the enforcement layers."))
        lines.append('#pagebreak()')
    lines.append('== X.23 Canonical identity and subscription contracts')
    lines.append(typst_table(["Table or surface", "Purpose", "Security boundary", "Evidence"], [["auth.users", "Supabase Auth identity", "Auth session/token", "PublicAuthGateway and profile resolution"], ["profiles", "User/company/role context", "Authenticated self-service and company scope", "Profile Identity Center migration"], ["companies", "Organization identity", "Tenant boundary", "Core migrations and live inventory"], ["company_memberships", "Membership role/status", "Company + user + role/status", "Invitation and auth workflows"], ["billing_plans", "Plan catalog", "Catalog/service/admin boundaries", "Subscription migrations"], ["tenant_subscriptions", "Company subscription state", "Server/database access snapshot", "Subscription migrations"], ["subscription_payments", "Provider payment state", "Server service RPC and idempotency", "Subscription handlers"], ["billing_access_snapshot", "Authoritative access result", "RPC, not a client cache", "Subscription access adapter"]]))
    lines.append('')
    lines.append('#pagebreak()')
    lines.append('= Part XI — Integrations')
    lines.append(typst_table(["Integration", "Purpose", "Current evidence", "Boundary"], [["Supabase", "Auth, Postgres/REST, RPC, RLS", "Live schema and migration evidence", "Service keys server-side"], ["HarakaPay", "USSD collection/status/webhook", "Handlers present; readiness required", "Order matching and verified settlement"], ["TRA/VFD", "Fiscal profile/receipt/retry/Z report", "Router and provider boundary", "TIN/VRN, branch, credentials"], ["Storage/S3", "Avatars, property files, exports", "Storage helper and proxy", "Scoped keys and validation"], ["Email/SMS/WhatsApp", "Notifications and outreach", "Persistence/provider boundaries", "No delivery claim without provider"], ["AI model", "Structured assistant proposals", "Built-in model boundary", "Context limits and approvals"], ["Vercel", "Deployment/runtime", "Build and serverless contract", "Environment variables"]], widths="(3.2cm, 5cm, 5.2cm, 4cm)"))
    for i,(name,purpose,evidence,boundary) in enumerate([["Supabase","Auth, Postgres/REST, RPC, RLS","Live schema and migration evidence","Service keys server-side"],["HarakaPay","USSD collection/status/webhook","Handlers present; readiness required","Order matching and verified settlement"],["TRA/VFD","Fiscal profile/receipt/retry/Z report","Router and provider boundary","TIN/VRN, branch, credentials"],["Storage/S3","Avatars, property files, exports","Storage helper and proxy","Scoped keys and validation"],["Email/SMS/WhatsApp","Notifications and outreach","Persistence/provider boundaries","No delivery claim without provider"],["AI model","Structured assistant proposals","Built-in model boundary","Context limits and approvals"],["Vercel","Deployment/runtime","Build and serverless contract","Environment variables"]],1):
        lines.append(f'== XI.{i} {typst_escape(name)}')
        lines.append(typst_escape(f"Purpose: {purpose}. Evidence: {evidence}. Boundary: {boundary}."))
        lines.append(typst_escape(f"Kiswahili: {name} hutumika kwa {purpose.lower()}. {evidence}. {boundary}."))
        lines.append('#pagebreak()')
    lines.append('= Part XII — Troubleshooting')
    for i,(problem,cause,diag,sol) in enumerate(TROUBLESHOOTING,1):
        lines.append(f'== XII.{i} {typst_escape(problem)}')
        lines.append(typst_escape(f"Possible cause: {cause} Diagnosis: {diag} Solution: {sol}"))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(f"Sababu: {cause} Utambuzi: {diag} Suluhisho: {sol}"))
        lines.append(typst_table(["Step", "Instruction", "Evidence"], [["1", "Preserve the visible error and input", "Browser/UI state"], ["2", "Check authenticated company and role", "Verified profile"], ["3", "Check server/database response", "API/RPC status"], ["4", "Retry only through idempotent path", "Request key or provider order"], ["5", "Record outcome for support/audit", "Audit result"]]))
        lines.append('#pagebreak()')
    lines.append('= Part XIII — Operations')
    ops = [("Release readiness", "Confirm source migration, tests, TypeScript, build, browser journeys, live schema, RLS, provider configuration, and rollback evidence.", "Thibitisha migration, tests, TypeScript, build, browser journeys, schema, RLS, provider na rollback."), ("Backup and recovery", "Distinguish database reachability from managed PITR or snapshot configuration. Do not report backup controls that were not checked.", "Tenganisha database reachability na PITR/snapshot. Usiripoti backup ambayo haikuthibitishwa."), ("Monitoring", "Watch API status, scheduled handlers, webhook deliveries, provider errors, failed mutations, security advisors, and schema drift.", "Fuatilia API, scheduled handlers, webhooks, provider errors, failed writes, security advisors na schema drift."), ("Support handover", "Keep the repository, migration ledger, environment checklist, role matrix, error reference, and incident notes together.", "Hifadhi repository, migration ledger, env checklist, role matrix, error reference na incident notes pamoja."), ("Data governance", "Define retention, privacy, least privilege, export, deletion, and incident escalation by business and regulatory context.", "Weka retention, privacy, least privilege, export, deletion na escalation kulingana na biashara na compliance."), ("Change control", "Every schema, role, provider, and entitlement change needs a source commit, migration/version record, test evidence, and updated status.", "Kila mabadiliko ya schema, role, provider na entitlement yanahitaji commit, migration/version, tests na status mpya.")]
    for i,(name,en,sw) in enumerate(ops,1):
        lines.append(f'== XIII.{i} {typst_escape(name)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(sw))
        lines.append('#pagebreak()')
    lines.append('= Part XIV — Future Roadmap')
    roadmap = [("Atomic legacy invoice payment RPC", "Replace the separate payment insert and invoice update with a reviewed atomic, idempotent database boundary."), ("Broader authenticated CRUD journeys", "Expand role-by-role Playwright and controlled Supabase staging coverage across legacy modules."), ("Signature-specific SECURITY DEFINER hardening", "Review each exposed signature, preserve intentional endpoints, and revoke or relocate internal helpers."), ("RLS policy consolidation", "Reduce overlapping permissive policies after proving equivalent tenant semantics."), ("Provider readiness", "Complete approved HarakaPay, TRA/VFD, storage, email/SMS/WhatsApp, maps, and AI configuration and acceptance."), ("Dashboard decomposition", "Extract safe, high-risk boundaries and code-split the large dashboard only after persistence and live blockers are addressed."),]
    for i,(name,en) in enumerate(roadmap,1):
        lines.append(f'== XIV.{i} {typst_escape(name)}')
        lines.append(typst_escape(en))
        lines.append('=== Kiswahili')
        lines.append(typst_escape(f"Hatua hii inapaswa kufanywa baada ya uthibitisho wa source, data, ruhusa na deployment. Lengo si kuongeza madai ya bidhaa, bali kuimarisha ushahidi wa kazi iliyopo: {en}"))
        lines.append('#pagebreak()')
    lines.append('= Part XV — Swahili Reference')
    lines.append(typst_escape("This section is a second, complete Tanzanian Swahili reference. Technical English terms remain in parentheses where they reduce ambiguity."))
    for idx,m in enumerate(MODULES,1):
        lines.append(f'== XV.{idx} {typst_escape(m["sw"])}')
        for p in swahili_module_paragraphs(m): lines.append(typst_escape(p))
        lines.append(typst_escape(f"Hali ya sasa: {m['status']}. Mipaka: {m['limit']}"))
        lines.append(typst_table(["Sehemu", "Maelezo", "Udhibiti", "Hatua inayofuata"], [["Watumiaji", m["roles"], "Verified profile + role", "Controlled acceptance"], ["Workflow", m["workflow"], "Server-confirmed result", "Role-specific test"], ["Database", ", ".join(table_names(m["prefixes"])[:12]) or "shared scope", "RLS/company scope", "Schema verification"], ["Integration", "Provider where configured", "Server-side secrets", "Readiness evidence"], ["Mobile", "Responsive review and entry", "Same authorization", "Browser journey"]]))
        lines.append('#pagebreak()')
    lines.append('= Appendices')
    lines.append('== Appendix A — Role and permission matrix')
    lines.append(typst_table(["Role", "Dashboard", "Primary modules", "Boundary"], ROLE_ROWS, widths="(3.3cm, 4.3cm, 5.1cm, 4.1cm)"))
    lines.append('#pagebreak()')
    lines.append('== Appendix B — Master feature status matrix')
    lines.append(typst_table(["Module", "Status", "UI evidence", "Server evidence"], [[m["name"], m["status"], m["ui"], m["server"]] for m in MODULES], widths="(4.3cm, 3.1cm, 5.4cm, 5.4cm)"))
    lines.append('#pagebreak()')
    lines.append('== Appendix C — Glossary')
    lines.append(typst_table(["Term", "English", "Kiswahili"], GLOSSARY, widths="(3cm, 7cm, 7cm)"))
    lines.append('#pagebreak()')
    lines.append('== Appendix D — Verification record and references')
    lines.append(typst_escape(f"Repository version: 1.0.0 from package.json. Documentation date: {TODAY}. Live Supabase audit: read-only. Tables: {LIVE['table_count']}; public: {LIVE['public_count']}; auth: {LIVE['auth_count']}; RLS enabled: {LIVE['rls_enabled']}; RLS not enabled: {LIVE['rls_disabled']}; live migration records: {LIVE['migration_count']}; security advisor lints: {len(LIVE['security_lints'])}; performance advisor lints: {len(LIVE['performance_lints'])}."))
    lines.append('')
    lines.append(typst_escape("References: [1] client/src/App.tsx and the public authentication gateway; [2] client/src/BusinessSphereDashboard.jsx and the authenticated dashboard; [3] server/_core/apiApp.ts and server/routers.ts; [4] server/subscriptionBilling.ts; [5] server/propertyManagementOperations.ts; [6] server/smartAssistant.ts; [7] server/traFiscalRouter.ts; [8] FULL_SYSTEM_AUDIT_REPORT.md; [9] FULL_SYSTEM_IMPLEMENTATION_MATRIX.md; [10] package.json; [11] Supabase Row Level Security documentation (the official URL is preserved in the Markdown reference list)."))
    lines.append(typst_escape("Documentation QA statement: the book was generated from repository files and read-only live evidence. It does not expose credentials, does not fabricate screenshots, does not create a duplicate auth/users architecture, and explicitly records known limitations and external prerequisites."))
    TYPOUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "DDEFE6")
    for row in rows[:120]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def build_docx(markdown: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles['Normal'].font.name = "Aptos"
    styles['Normal'].font.size = Pt(10)
    for style_name, size, color in [("Title", 24, "0B3D2E"), ("Heading 1", 18, "0B3D2E"), ("Heading 2", 14, "176B4D"), ("Heading 3", 11, "176B4D")]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.text = "SMART MANAGER ERP — Master System Book | Repository-Audited Edition"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(8)
    footer = section.footer.paragraphs[0]
    footer.text = "Manus AI | Version 1.0.0 | 24 August 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SMART MANAGER ERP\n").bold = True
    title.add_run("MASTER SYSTEM BOOK\n").bold = True
    title.add_run("English and Tanzanian Swahili Edition\n\n").italic = True
    title.add_run("Repository-Audited Edition\nVersion 1.0.0\n24 August 2026\n\nManus AI\nOwner and Creator: Ezra Mpapi\nDar es Salaam, Tanzania")
    doc.add_page_break()
    markdown_lines = markdown.splitlines()
    i = 0
    while i < len(markdown_lines):
        line = markdown_lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("| ") and i + 1 < len(markdown_lines) and markdown_lines[i + 1].startswith("| "):
            def parse_row(row: str) -> list[str]:
                return [part.strip().replace("\\|", "|") for part in row.strip().strip("|").split("|")]
            headers = parse_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(markdown_lines) and markdown_lines[i].startswith("| "):
                rows.append(parse_row(markdown_lines[i]))
                i += 1
            add_docx_table(doc, headers, rows)
            continue
        if line.startswith("!["):
            match = re.search(r"\(([^)]+)\)", line)
            if match:
                image_ref = match.group(1)
                p = MASTER_BOOK_LOGO if image_ref == MASTER_BOOK_LOGO_STORAGE else ROOT / image_ref
                if not p.exists():
                    p = ROOT / "typst-project" / image_ref.replace("typst-project/", "")
                if p.exists():
                    doc.add_picture(str(p), width=Inches(5.2))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip('#')), 3)
            text = line[level:].strip().replace("—", "–")
            doc.add_heading(text, level=level)
            i += 1
            continue
        if line == "---":
            doc.add_page_break()
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(line[2:])
            run.italic = True
            i += 1
            continue
        if line.startswith("• ") or line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line.replace("**", "").replace("`", ""))
        i += 1
    # Add a compact, structured database appendix after the narrative, including all observed public tables.
    doc.add_page_break()
    doc.add_heading("Appendix — Live Supabase public table index", level=1)
    public_tables = [t for t in LIVE["tables"] if t.get("name", "").startswith("public.")]
    for start in range(0, len(public_tables), 30):
        rows = []
        for t in sorted(public_tables, key=lambda x: x.get("name", ""))[start:start+30]:
            rows.append([t.get("name", "").split(".",1)[-1], ", ".join(c.get("name", "") for c in t.get("columns", []))[:260], "Enabled" if t.get("rls_enabled") else "Review"])
        add_docx_table(doc, ["Table", "Observed columns", "RLS"], rows)
        if start + 30 < len(public_tables):
            doc.add_page_break()
    doc.save(DOCX_PATH)


def build_manifest(markdown: str) -> None:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    manifest = {
        "title": "SMART MANAGER ERP — Master System Book",
        "edition": "Repository-Audited English and Tanzanian Swahili Edition",
        "version": "1.0.0",
        "documentation_date": TODAY,
        "owner_creator": "Ezra Mpapi",
        "author": "Manus AI",
        "source_commit": "07f66c9531ed30861ca49758d2733a9fade9c647",
        "live_supabase": {"project_id": "rlhngsrihahhyxnjxrxm", "read_only": True, "table_count": LIVE["table_count"], "public_count": LIVE["public_count"], "auth_count": LIVE["auth_count"], "rls_enabled": LIVE["rls_enabled"], "rls_disabled": LIVE["rls_disabled"], "migration_count": LIVE["migration_count"], "security_lints": len(LIVE["security_lints"]), "performance_lints": len(LIVE["performance_lints"])},
        "module_count": len(MODULES),
        "deliverables": ["SMART_MANAGER_MASTER_BOOK_EN_SW.md", "SMART_MANAGER_MASTER_BOOK_EN_SW.docx", "SMART_MANAGER_MASTER_BOOK_EN_SW.pdf", "master_book.typ", "evidence/", "assets/"],
        "evidence_boundary": "No live DDL, provider payment, external message, or credential was invoked during book generation.",
    }
    (DELIVERABLES / "SMART_MANAGER_MASTER_BOOK_MANIFEST.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def main() -> None:
    markdown = build_markdown()
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    SOURCE_MD.write_text(markdown, encoding="utf-8")
    build_typst()
    build_docx(markdown)
    build_manifest(markdown)
    print(json.dumps({"markdown": str(SOURCE_MD), "typst": str(TYPOUT), "docx": str(DOCX_PATH), "markdown_chars": len(markdown), "module_count": len(MODULES), "live_table_count": LIVE["table_count"]}, indent=2))


if __name__ == "__main__":
    main()
