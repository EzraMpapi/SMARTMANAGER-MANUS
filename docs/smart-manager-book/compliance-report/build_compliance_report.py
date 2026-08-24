from __future__ import annotations

import json
import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[3]
BOOK = ROOT / "docs" / "smart-manager-book" / "master-book"
OUT = ROOT / "docs" / "smart-manager-book" / "compliance-report"
DELIVERABLES = OUT / "deliverables"
EVIDENCE = BOOK / "evidence"
SOURCE_MD = BOOK / "SMART_MANAGER_MASTER_BOOK_EN_SW.md"
LIVE_TABLES = EVIDENCE / "live_supabase_tables_2026-08-24.json"
LIVE_METRICS = EVIDENCE / "live_audit_metrics_2026-08-24.txt"
ADVISOR_COUNTS = EVIDENCE / "advisor_counts_2026-08-24.txt"

MD_PATH = OUT / "SMART_MANAGER_COMPLIANCE_REPORT.md"
TYPST_PATH = OUT / "main.typ"
DOCX_PATH = DELIVERABLES / "SMART_MANAGER_COMPLIANCE_REPORT.docx"
PDF_PATH = DELIVERABLES / "SMART_MANAGER_COMPLIANCE_REPORT.pdf"
MANIFEST_PATH = DELIVERABLES / "SMART_MANAGER_COMPLIANCE_REPORT_MANIFEST.json"

AUDIT_DATE = "24 August 2026"
PROJECT_ID = "rlhngsrihahhyxnjxrxm"
COMMIT = "d20a9b922e8596d54f3c7538b6389f71f4aef869"

RISK_REGISTER = [
    {
        "id": "R-01",
        "title": "Authenticated SECURITY DEFINER routines",
        "severity": "WARN",
        "evidence": "The live security advisor reported 118 WARN and 1 INFO lint at the audit timestamp; many WARNs identify signed-in execution of SECURITY DEFINER routines.",
        "action": "Review each signature, keep only intentionally callable endpoints, pin search paths, apply narrow grants, and move internal helpers out of the exposed API surface where possible.",
        "owner": "Database/security engineering",
        "priority": "P1",
    },
    {
        "id": "R-02",
        "title": "Multiple permissive RLS policies",
        "severity": "WARN",
        "evidence": "The live performance advisor reported 851 lints, including multiple permissive policies on the same table/action.",
        "action": "Consolidate overlapping policies by command and role after verifying semantics; do not blindly drop production policies.",
        "owner": "Database/security engineering",
        "priority": "P1",
    },
    {
        "id": "R-03",
        "title": "Legacy non-atomic invoice payment path",
        "severity": "P0 historical finding",
        "evidence": "The audit report records separate sales_payments insertion and invoice balance update without a proven atomic idempotency RPC.",
        "action": "Add a reviewed tenant-scoped atomic RPC and durable idempotency key before claiming concurrent-safe posting.",
        "owner": "Finance platform engineering",
        "priority": "P0",
    },
    {
        "id": "R-04",
        "title": "Large dashboard boundary",
        "severity": "P2",
        "evidence": "BusinessSphereDashboard.jsx remains a very large monolithic component and the build reports a non-fatal large-chunk warning.",
        "action": "Decompose incrementally after persistence and live-environment blockers are addressed; avoid cosmetic rewrites that increase risk.",
        "owner": "Frontend platform engineering",
        "priority": "P2",
    },
    {
        "id": "R-05",
        "title": "External provider readiness",
        "severity": "Configuration boundary",
        "evidence": "HarakaPay, TRA/VFD, WhatsApp, email/SMS, storage, and AI services depend on deployment configuration and approved credentials.",
        "action": "Keep provider secrets server-side, expose readiness states, and test only controlled sandbox or authorized production paths.",
        "owner": "Platform operations",
        "priority": "P1",
    },
    {
        "id": "R-06",
        "title": "Demo fallback risk",
        "severity": "Medium",
        "evidence": "The client has an explicit seed-data fallback when Supabase is not configured.",
        "action": "Production deployments must fail closed with a clear configuration message; demo mode must remain explicit and non-operational.",
        "owner": "Application/platform engineering",
        "priority": "P1",
    },
]

CANONICAL_IDENTITY = [
    ["auth.users", "Supabase Auth identity", "Auth session and token boundary", "PublicAuthGateway and profile resolution"],
    ["profiles", "User/company profile and role context", "Authenticated self-service and company scope", "Most protected workflows"],
    ["companies", "Organization/company identity", "Tenant boundary", "All company-scoped modules"],
    ["company_memberships", "Membership and role relationship", "Company plus user plus role/status", "Invitations, onboarding, authorization"],
    ["workspaces", "Workspace context where present", "Workspace-aware navigation", "Shell and module context"],
    ["user_table_preferences", "User preference persistence", "Self-only or company-scoped policy", "Dashboard/profile settings"],
]

CANONICAL_SUBSCRIPTION = [
    ["billing_plans", "Package catalog", "Free and paid plan metadata, entitlements", "Subscription migrations and live inventory"],
    ["billing_profiles", "Billing contact/configuration", "Company-scoped billing profile", "Billing foundation migration"],
    ["tenant_subscriptions", "Company subscription state", "Pending/Active/Grace/Expired/RequiresPlan/Cancelled/Superseded", "Subscription model migrations"],
    ["subscription_payments", "Provider payment state", "Monthly cycle, idempotency, provider order", "HarakaPay handlers and migrations"],
    ["subscription_invoices", "Billing invoice evidence", "Payment/subscription linkage", "Billing foundation migration"],
    ["subscription_usage", "Usage/limits evidence", "Company and plan context", "Billing foundation migration"],
    ["subscription_events", "Billing lifecycle events", "Audit and reconciliation history", "Billing migrations"],
    ["subscription_notifications", "Billing notifications", "Company-scoped notification state", "Billing migrations"],
    ["billing_access_snapshot", "Authoritative access result", "Server/database snapshot; not a table", "Access adapter and protected API"],
]


def read_kv(path: Path) -> dict[str, str]:
    result: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if "=" in line:
            key, value = line.split("=", 1)
            result[key.strip()] = value.strip()
    return result


def load_tables() -> list[dict[str, Any]]:
    raw = json.loads(LIVE_TABLES.read_text(encoding="utf-8"))
    tables = raw.get("tables", raw if isinstance(raw, list) else [])
    return [t for t in tables if isinstance(t, dict) and str(t.get("name", "")).startswith("public.")]


def live_metrics() -> dict[str, str]:
    metrics = read_kv(LIVE_METRICS)
    advisor_lines = ADVISOR_COUNTS.read_text(encoding="utf-8").splitlines()
    section = "security"
    for line in advisor_lines:
        if line.strip().lower().startswith("=== performance"):
            section = "performance"
            continue
        if "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip()
        metrics[f"{section}_{key}"] = value
    metrics["table_count"] = str(int(metrics.get("live_public_table_count", "0")) + int(metrics.get("live_auth_table_count", "0")))
    return metrics


def md_escape(value: Any) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    lines = ["| " + " | ".join(md_escape(h) for h in headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    lines.extend("| " + " | ".join(md_escape(v) for v in row) + " |" for row in rows)
    return "\n".join(lines)


def table_rows(tables: list[dict[str, Any]]) -> list[list[str]]:
    rows = []
    for table in sorted(tables, key=lambda x: str(x.get("name", ""))):
        name = str(table.get("name", "")).split(".", 1)[-1]
        columns = ", ".join(str(c.get("name", "")) for c in table.get("columns", []))
        pks = ", ".join(str(pk) for pk in table.get("primary_keys", [])) or "—"
        rls = "Enabled" if table.get("rls_enabled") else "Review"
        rows.append([name, columns, pks, rls, str(table.get("rows", 0))])
    return rows


def build_markdown(tables: list[dict[str, Any]], metrics: dict[str, str]) -> str:
    public_rows = table_rows(tables)
    out: list[str] = [
        "# SMART MANAGER ERP — SECURITY AND DATABASE COMPLIANCE REPORT",
        "## Focused extraction from the repository-audited master system book",
        "",
        f"> **Audit date:** {AUDIT_DATE}; **Supabase project:** `{PROJECT_ID}`; **Evidence mode:** Read-only; **Source commit:** `{COMMIT}`; **Prepared by:** Manus AI",
        "",
        "## Executive scope",
        "",
        "This focused report extracts the **security risk register** and **live Supabase database schema dictionary** from the Smart Manager ERP master book. It is a compliance evidence report, not a statement that every module, provider, policy, or external integration is production-complete. The live database evidence was collected read-only and is time-bound to the audit date. [1] [2]",
        "",
        "The report separates observed evidence from remediation recommendations. It does not authorize destructive policy changes, broad RLS rewrites, credential rotation, provider payments, or production migration application. Any remediation must remain source-versioned, reviewed, tenant-scoped, and validated against the live environment before release.",
        "",
        "## Evidence snapshot",
        "",
        md_table(
            ["Metric", "Observed value", "Interpretation"],
            [
                ["Total tables", metrics.get("table_count", "—"), "Combined public and auth tables returned by the read-only inventory"],
                ["Public tables", metrics.get("live_public_table_count", "—"), "Application-facing public-schema inventory"],
                ["Auth tables", metrics.get("live_auth_table_count", "—"), "Supabase Auth schema inventory"],
                ["RLS enabled", metrics.get("live_rls_enabled_count", "—"), "Tables reported with RLS enabled"],
                ["RLS not enabled", metrics.get("live_rls_disabled_count", "—"), "Requires table-specific review; not a reason to add broad policies blindly"],
                ["Migration records", metrics.get("migration_count", "—"), "Records returned by the live migration ledger"],
                ["Security advisor", metrics.get("security_total", "—"), f"{metrics.get('security_WARN', '—')} WARN and {metrics.get('security_INFO', '—')} INFO"],
                ["Performance advisor", metrics.get("performance_total", "—"), f"{metrics.get('performance_WARN', '—')} WARN and {metrics.get('performance_INFO', '—')} INFO in the saved advisor-count file"],
            ],
        ),
        "",
        "> **Interpretation note.** Advisor counts are findings, not proof that every finding has the same severity or exploitability. The database dictionary reports metadata and row estimates; it does not expose row contents. [3] [4]",
        "",
        "# 1. Security risk register",
        "",
        "The risk register below preserves the six risks extracted from the master book. The owner and priority columns are compliance-management additions for triage; they do not change the underlying evidence or severity wording.",
        "",
        md_table(["ID", "Risk", "Severity", "Priority", "Primary owner"], [[r["id"], r["title"], r["severity"], r["priority"], r["owner"]] for r in RISK_REGISTER]),
        "",
    ]
    for r in RISK_REGISTER:
        out += [
            f"## {r['id']} — {r['title']}",
            "",
            f"**Severity:** {r['severity']}; **Priority:** {r['priority']}; **Primary owner:** {r['owner']}",
            "",
            f"**Evidence.** {r['evidence']}",
            "",
            f"**Recommended remediation.** {r['action']}",
            "",
            "**Control principle.** The platform should prefer explicit identity, narrow role authority, company scope, server confirmation, database enforcement, and auditable results. A local UI state, a draft message, or a client-supplied amount cannot replace those controls.",
            "",
        ]
    out += [
        "## 1.1 Remediation sequencing",
        "",
        "The historical P0 invoice-payment finding should be handled as a finance-safety gate before concurrent-safe posting is claimed. The SECURITY DEFINER and RLS-policy findings should be addressed through signature-specific and command-specific review rather than bulk revocation or blanket policy replacement. Provider readiness and demo fallback are release-boundary controls, while dashboard decomposition is a P2 engineering improvement that should not precede data-integrity and access-control work.",
        "",
        md_table(["Sequence", "Gate", "Required evidence before closure"], [
            ["1", "Invoice payment atomicity", "Tenant-scoped atomic RPC, durable idempotency, concurrent test evidence, reconciliation evidence"],
            ["2", "Security-definer exposure", "Per-signature inventory, pinned search paths, least-privilege grants, intentional endpoint record"],
            ["3", "RLS policy consolidation", "Command/role semantics reviewed, tenant isolation tests, no broad policy regression"],
            ["4", "Provider readiness", "Server-side secret boundary, sandbox/approved production acceptance, visible readiness state"],
            ["5", "Demo fallback", "Fail-closed production configuration and explicit non-operational demo mode"],
            ["6", "Dashboard decomposition", "Incremental module boundaries, bundle evidence, regression coverage"],
        ]),
        "",
        "# 2. Database schema dictionary",
        "",
        f"The live read-only Supabase inventory returned **{metrics.get('table_count', '—')} tables**: {metrics.get('live_public_table_count', '—')} public and {metrics.get('live_auth_table_count', '—')} auth. This dictionary lists the **{len(public_rows)} public tables** observed in the snapshot, their observed columns, reported primary keys, RLS state, and reported row estimates. It is metadata evidence, not a data export. [3]",
        "",
        "## 2.1 Canonical identity and tenancy contract",
        "",
        "The Smart Manager architecture uses Supabase Auth plus the existing profile, company, membership, workspace, and preference surfaces. A separate application `users` table must not be introduced merely to restate `auth.users` or `profiles`.",
        "",
        md_table(["Table", "Purpose", "Security boundary", "Used by"], CANONICAL_IDENTITY),
        "",
        "## 2.2 Canonical subscription contract",
        "",
        "Subscription authority remains in the existing billing and subscription tables plus the server/database access snapshot. Entitlements are represented through the catalog and access snapshot contract rather than a parallel subscription-items or duplicate-entitlements architecture.",
        "",
        md_table(["Table or contract", "Purpose", "Contract", "Evidence"], CANONICAL_SUBSCRIPTION),
        "",
        "## 2.3 Public table dictionary",
        "",
        md_table(["Table", "Observed columns", "Primary key", "RLS", "Rows reported"], public_rows),
        "",
        "## 2.4 Dictionary interpretation rules",
        "",
        "The reported row estimates are a point-in-time audit snapshot and should not be treated as a completeness assertion. A table with zero reported rows can still be a valid persistence contract, while a non-zero estimate does not prove that every workflow is fully tested. RLS state must be evaluated together with policies, helper functions, grants, foreign keys, triggers, and server authorization. Changes should use source-versioned migrations and should fail safely when existing objects are incompatible.",
        "",
        "## 3. Compliance conclusion",
        "",
        "The evidence supports a substantial existing Supabase architecture with canonical identity, tenancy, billing, subscription, and module persistence surfaces already present. The most important open compliance work is targeted: protect the payment posting boundary with atomic idempotency, review exposed SECURITY DEFINER signatures, consolidate overlapping RLS policies only after semantic proof, maintain provider secret boundaries, and keep demo behavior explicit and non-operational. The report does not support recreating the schema, adding duplicate identity tables, or applying blanket production DDL.",
        "",
        "## References",
        "",
        f"[1]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/docs/smart-manager-book/master-book/SMART_MANAGER_MASTER_BOOK_EN_SW.md \"Repository-audited master book source\"",
        f"[2]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/FULL_SYSTEM_AUDIT_REPORT.md \"Full-system audit and historical findings\"",
        f"[3]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/docs/smart-manager-book/master-book/evidence/live_supabase_tables_2026-08-24.json \"Read-only Supabase table inventory snapshot\"",
        f"[4]: https://github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/docs/smart-manager-book/master-book/evidence/advisor_counts_2026-08-24.txt \"Saved Supabase advisor count summary\"",
        "[5]: https://supabase.com/docs/guides/database/postgres/row-level-security \"Supabase Row Level Security documentation\"",
        "",
        "## Documentation and security notice",
        "",
        "This report contains schema metadata, security findings, and remediation guidance. It intentionally excludes credentials, service-role keys, provider tokens, passwords, private customer records, and raw business payloads. It should be distributed together with its evidence date and source commit.",
        "",
    ]
    return "\n".join(out).rstrip() + "\n"


def typst_escape(value: Any) -> str:
    return str(value).replace("\\", "\\\\").replace("#", "\\#").replace("[", "(").replace("]", ")").replace("_", "\\u{005F}").replace("$", "\\$")


def typst_table(headers: list[str], rows: list[list[Any]], widths: str) -> str:
    cells = [f"[*{typst_escape(h)}*]" for h in headers]
    for row in rows:
        cells.extend(f"[{typst_escape(v)}]" for v in row)
    return f"#table(columns: {widths}, inset: 3pt, stroke: 0.3pt + luma(190), table.header({', '.join(cells[:len(headers)])}), {', '.join(cells[len(headers):])})"


def build_typst(tables: list[dict[str, Any]], metrics: dict[str, str]) -> None:
    theme_source = BOOK / "typst-project" / "report-theme.typ"
    OUT.mkdir(parents=True, exist_ok=True)
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    shutil.copy2(theme_source, OUT / "report-theme.typ")
    rows = table_rows(tables)
    lines = [
        '#import "report-theme.typ": report-accent, report-theme',
        '#show: report-theme.with(title: "SMART MANAGER ERP — Security and Database Compliance Report", author: "Manus AI", rhythm: "longform", running-header: true)',
        '#set text(font: ("Libertinus Serif", "Noto Sans"), lang: "en")',
        '#set par(justify: true)',
        '#page(fill: rgb("#061A13"), margin: (top: 1.4cm, bottom: 1.4cm, x: 2cm), numbering: none, header: none)[#align(center)[#v(2.2cm)#text(fill: white, size: 25pt, weight: "bold")[SMART MANAGER ERP]#v(0.5cm)#text(fill: rgb("#D6EADF"), size: 17pt)[SECURITY AND DATABASE COMPLIANCE REPORT]#v(1cm)#line(length: 70%, stroke: 1pt + rgb("#D6B36A"))#v(0.9cm)#text(fill: white, size: 11pt)[Focused extraction from the repository-audited master system book]#v(1.4cm)#text(fill: white, size: 11pt)[Read-only evidence \\ Audit date: 24 August 2026 \\ Supabase project: rlhngsrihahhyxnjxrxm \\ Prepared by Manus AI]]]',
        '#page(numbering: none, header: none)[#align(center)[#text(size: 20pt, weight: "bold", fill: report-accent)[Compliance scope]#v(0.8cm)#text(size: 11pt)[This report extracts the security risk register and live Supabase schema dictionary. It distinguishes observed evidence from recommended remediation and does not authorize destructive production changes. The live database inspection was read-only and time-bound to the audit date.]]]',
        '#page(numbering: none, header: none)[#outline(title: [Contents], indent: 1.5em)]',
        '#counter(page).update(1)',
        '= Executive evidence snapshot',
        typst_table(["Metric", "Value", "Interpretation"], [["Total tables", metrics.get("table_count", "—"), "Public plus auth metadata returned by read-only inventory"], ["Public tables", metrics.get("live_public_table_count", "—"), "Application-facing public schema"], ["Auth tables", metrics.get("live_auth_table_count", "—"), "Supabase Auth schema"], ["RLS enabled", metrics.get("live_rls_enabled_count", "—"), "Reported enabled in returned metadata"], ["RLS not enabled", metrics.get("live_rls_disabled_count", "—"), "Requires table-specific review"], ["Migration records", metrics.get("migration_count", "—"), "Live migration ledger records"], ["Security lints", metrics.get("security_total", "—"), f"{metrics.get('security_WARN', '—')} WARN; {metrics.get('security_INFO', '—')} INFO"], ["Performance lints", metrics.get("performance_total", "—"), "Saved performance advisor count"]], "(3.1cm, 2.1cm, 11.0cm)"),
        "The dictionary reports metadata and row estimates; it does not expose row contents. Advisor counts are findings rather than a claim that every finding has equal severity or exploitability.",
        '#pagebreak()',
        '= 1. Security risk register',
        'The six extracted risks are shown below. The owner and priority fields support triage and do not replace the underlying severity wording.',
        typst_table(["ID", "Risk", "Severity", "Priority", "Owner"], [[r["id"], r["title"], r["severity"], r["priority"], r["owner"]] for r in RISK_REGISTER], "(1.1cm, 5.0cm, 3.0cm, 1.6cm, 6.0cm)"),
    ]
    for r in RISK_REGISTER:
        lines += [
            f"== {typst_escape(r['id'])} — {typst_escape(r['title'])}",
            f"*Severity:* {typst_escape(r['severity'])} \\ *Priority:* {typst_escape(r['priority'])} \\ *Primary owner:* {typst_escape(r['owner'])}",
            f"*Evidence.* {typst_escape(r['evidence'])}",
            f"*Recommended remediation.* {typst_escape(r['action'])}",
            "*Control principle.* The platform should prefer explicit identity, narrow role authority, company scope, server confirmation, database enforcement, and auditable results. A local UI state, a draft message, or a client-supplied amount cannot replace those controls.",
        ]
    lines += [
        '== Remediation sequencing',
        'The historical P0 invoice-payment finding should be handled as a finance-safety gate before concurrent-safe posting is claimed. SECURITY DEFINER and RLS findings require signature-specific and command-specific review rather than bulk revocation or blanket replacement.',
        typst_table(["Sequence", "Gate", "Closure evidence"], [["1", "Invoice payment atomicity", "Atomic RPC, durable idempotency, concurrent and reconciliation tests"], ["2", "SECURITY DEFINER exposure", "Per-signature inventory, search_path, grants, intentional endpoint record"], ["3", "RLS policy consolidation", "Command/role semantics and tenant-isolation tests"], ["4", "Provider readiness", "Server-side secret boundary and controlled acceptance"], ["5", "Demo fallback", "Fail-closed production configuration"], ["6", "Dashboard decomposition", "Incremental boundaries and regression coverage"]], "(1.4cm, 5.1cm, 10.2cm)"),
        '#pagebreak()',
        '= 2. Database schema dictionary',
        f"The read-only inventory returned {typst_escape(metrics.get('table_count', '—'))} tables: {typst_escape(metrics.get('live_public_table_count', '—'))} public and {typst_escape(metrics.get('live_auth_table_count', '—'))} auth. The public dictionary contains {len(rows)} observed tables.",
        '== Canonical identity and tenancy contract',
        'The architecture uses Supabase Auth plus profiles, companies, memberships, workspaces, and preferences. Do not introduce a duplicate application users table merely to restate auth.users or profiles.',
        typst_table(["Table", "Purpose", "Security boundary", "Used by"], CANONICAL_IDENTITY, "(3.2cm, 4.1cm, 4.8cm, 4.2cm)"),
        '== Canonical subscription contract',
        'Subscription authority remains in the existing billing and subscription tables plus the server/database access snapshot. Entitlements are not recreated as a parallel subscription-items architecture.',
        typst_table(["Table or contract", "Purpose", "Contract", "Evidence"], CANONICAL_SUBSCRIPTION, "(3.4cm, 4.0cm, 4.8cm, 4.1cm)"),
        '== Public table dictionary',
    ]
    for start in range(0, len(rows), 40):
        chunk = rows[start:start + 40]
        lines.append(typst_table(["Table", "Observed columns", "Primary key", "RLS", "Rows"], chunk, "(3.0cm, 9.3cm, 2.0cm, 1.2cm, 1.3cm)"))
        if start + 40 < len(rows):
            lines.append('#pagebreak()')
    lines += [
        '#pagebreak()',
        '= 3. Compliance conclusion',
        'The evidence supports a substantial existing Supabase architecture with canonical identity, tenancy, billing, subscription, and module persistence surfaces already present. The important open work is targeted: protect payment posting with atomic idempotency, review exposed SECURITY DEFINER signatures, consolidate overlapping RLS policies only after semantic proof, maintain provider secret boundaries, and keep demo behavior explicit and non-operational.',
        'This report does not support recreating the schema, adding duplicate identity tables, or applying blanket production DDL.',
        '== References',
        f"[1] Repository-audited master book source at github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/docs/smart-manager-book/master-book/SMART_MANAGER_MASTER_BOOK_EN_SW.md.",
        f"[2] Full-system audit at github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/FULL_SYSTEM_AUDIT_REPORT.md.",
        f"[3] Read-only Supabase table inventory at github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/docs/smart-manager-book/master-book/evidence/live_supabase_tables_2026-08-24.json.",
        f"[4] Saved advisor counts at github.com/EzraMpapi/SMARTMANAGER-MANUS/blob/{COMMIT}/docs/smart-manager-book/master-book/evidence/advisor_counts_2026-08-24.txt.",
        'Supabase Row Level Security documentation: supabase.com/docs/guides/database/postgres/row-level-security.',
    ]
    TYPST_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(7.5)


def add_docx_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell(cell, value)
    doc.add_paragraph()


def add_docx_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=min(level, 3))
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 94, 150)


def build_docx(tables: list[dict[str, Any]], metrics: dict[str, str]) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    header = sec.header.paragraphs[0]
    header.text = "SMART MANAGER ERP — Security and Database Compliance Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    footer = sec.footer.paragraphs[0]
    footer.text = "Manus AI | Read-only evidence | 24 August 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SMART MANAGER ERP\n")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(6, 26, 19)
    run = title.add_run("SECURITY AND DATABASE COMPLIANCE REPORT\n")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(31, 94, 150)
    run = title.add_run("Focused extraction from the repository-audited master system book\n\n")
    run.italic = True
    run.font.size = Pt(11)
    title.add_run(f"Read-only evidence\nAudit date: {AUDIT_DATE}\nSupabase project: {PROJECT_ID}\nPrepared by Manus AI")
    doc.add_page_break()
    add_docx_heading(doc, "Executive scope", 1)
    doc.add_paragraph("This focused report extracts the security risk register and live Supabase database schema dictionary from the Smart Manager ERP master book. It separates observed evidence from remediation recommendations and does not authorize destructive production changes. The live database evidence was collected read-only and is time-bound to the audit date.")
    add_docx_heading(doc, "Evidence snapshot", 1)
    add_docx_table(doc, ["Metric", "Observed value", "Interpretation"], [["Total tables", metrics.get("table_count", "—"), "Combined public and auth tables"], ["Public tables", metrics.get("live_public_table_count", "—"), "Application-facing public schema"], ["Auth tables", metrics.get("live_auth_table_count", "—"), "Supabase Auth schema"], ["RLS enabled", metrics.get("live_rls_enabled_count", "—"), "Reported enabled"], ["RLS not enabled", metrics.get("live_rls_disabled_count", "—"), "Table-specific review required"], ["Migration records", metrics.get("migration_count", "—"), "Live ledger records"], ["Security advisor", metrics.get("total", "—"), f"{metrics.get('WARN', '—')} WARN; {metrics.get('INFO', '—')} INFO"], ["Performance advisor", metrics.get("total", "—"), "Saved performance count"]])
    doc.add_paragraph("Advisor counts are findings, not proof that every finding has equal severity or exploitability. The dictionary reports metadata and row estimates; it does not expose row contents.")
    doc.add_page_break()
    add_docx_heading(doc, "1. Security risk register", 1)
    doc.add_paragraph("The six extracted risks are shown below. Owner and priority fields support triage and do not replace the underlying severity wording.")
    add_docx_table(doc, ["ID", "Risk", "Severity", "Priority", "Owner"], [[r["id"], r["title"], r["severity"], r["priority"], r["owner"]] for r in RISK_REGISTER])
    for r in RISK_REGISTER:
        add_docx_heading(doc, f"{r['id']} — {r['title']}", 2)
        doc.add_paragraph(f"Severity: {r['severity']} | Priority: {r['priority']} | Primary owner: {r['owner']}")
        doc.add_paragraph(f"Evidence. {r['evidence']}")
        doc.add_paragraph(f"Recommended remediation. {r['action']}")
        doc.add_paragraph("Control principle. The platform should prefer explicit identity, narrow role authority, company scope, server confirmation, database enforcement, and auditable results. A local UI state, a draft message, or a client-supplied amount cannot replace those controls.")
    add_docx_heading(doc, "Remediation sequencing", 2)
    doc.add_paragraph("The historical P0 invoice-payment finding should be handled as a finance-safety gate before concurrent-safe posting is claimed. SECURITY DEFINER and RLS findings require signature-specific and command-specific review rather than bulk revocation or blanket replacement.")
    add_docx_table(doc, ["Sequence", "Gate", "Closure evidence"], [["1", "Invoice payment atomicity", "Atomic RPC, durable idempotency, concurrent and reconciliation tests"], ["2", "SECURITY DEFINER exposure", "Per-signature inventory, search_path, grants, endpoint record"], ["3", "RLS policy consolidation", "Command/role semantics and tenant-isolation tests"], ["4", "Provider readiness", "Server-side secret boundary and controlled acceptance"], ["5", "Demo fallback", "Fail-closed production configuration"], ["6", "Dashboard decomposition", "Incremental boundaries and regression coverage"]])
    doc.add_page_break()
    add_docx_heading(doc, "2. Database schema dictionary", 1)
    doc.add_paragraph(f"The read-only inventory returned {metrics.get('table_count', '—')} tables: {metrics.get('live_public_table_count', '—')} public and {metrics.get('live_auth_table_count', '—')} auth. This dictionary lists the {len(tables)} public tables observed in the snapshot, their observed columns, reported primary keys, RLS state, and reported row estimates.")
    add_docx_heading(doc, "Canonical identity and tenancy contract", 2)
    doc.add_paragraph("The architecture uses Supabase Auth plus profiles, companies, memberships, workspaces, and preferences. Do not introduce a duplicate application users table merely to restate auth.users or profiles.")
    add_docx_table(doc, ["Table", "Purpose", "Security boundary", "Used by"], CANONICAL_IDENTITY)
    add_docx_heading(doc, "Canonical subscription contract", 2)
    doc.add_paragraph("Subscription authority remains in the existing billing and subscription tables plus the server/database access snapshot. Entitlements are not recreated as a parallel subscription-items architecture.")
    add_docx_table(doc, ["Table or contract", "Purpose", "Contract", "Evidence"], CANONICAL_SUBSCRIPTION)
    add_docx_heading(doc, "Public table dictionary", 2)
    rows = table_rows(tables)
    for start in range(0, len(rows), 35):
        add_docx_table(doc, ["Table", "Observed columns", "Primary key", "RLS", "Rows reported"], rows[start:start + 35])
        if start + 35 < len(rows):
            doc.add_page_break()
    add_docx_heading(doc, "Dictionary interpretation rules", 2)
    doc.add_paragraph("The reported row estimates are a point-in-time audit snapshot and are not a completeness assertion. RLS state must be evaluated together with policies, helper functions, grants, foreign keys, triggers, and server authorization. Changes should use source-versioned migrations and fail safely when existing objects are incompatible.")
    add_docx_heading(doc, "3. Compliance conclusion", 1)
    doc.add_paragraph("The evidence supports a substantial existing Supabase architecture with canonical identity, tenancy, billing, subscription, and module persistence surfaces already present. The important open work is targeted: protect payment posting with atomic idempotency, review exposed SECURITY DEFINER signatures, consolidate overlapping RLS policies only after semantic proof, maintain provider secret boundaries, and keep demo behavior explicit and non-operational. This report does not support recreating the schema, adding duplicate identity tables, or applying blanket production DDL.")
    add_docx_heading(doc, "References", 1)
    for ref in [
        f"[1] Repository-audited master book source: GitHub commit {COMMIT}.",
        f"[2] Full-system audit: GitHub commit {COMMIT}.",
        f"[3] Read-only Supabase table inventory snapshot: GitHub commit {COMMIT}.",
        f"[4] Saved Supabase advisor count summary: GitHub commit {COMMIT}.",
        "[5] Supabase Row Level Security documentation.",
    ]:
        doc.add_paragraph(ref, style="List Bullet")
    doc.add_paragraph("Security notice: this report excludes credentials, service-role keys, provider tokens, passwords, private customer records, and raw business payloads.")
    doc.save(DOCX_PATH)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    tables = load_tables()
    metrics = live_metrics()
    markdown = build_markdown(tables, metrics)
    MD_PATH.write_text(markdown, encoding="utf-8")
    build_typst(tables, metrics)
    build_docx(tables, metrics)
    manifest = {
        "title": "SMART MANAGER ERP — Security and Database Compliance Report",
        "audit_date": AUDIT_DATE,
        "source_commit": COMMIT,
        "supabase_project_id": PROJECT_ID,
        "read_only_evidence": True,
        "public_table_rows": len(tables),
        "risk_count": len(RISK_REGISTER),
        "deliverables": [str(MD_PATH), str(TYPST_PATH), str(DOCX_PATH), str(PDF_PATH)],
        "sources": [str(SOURCE_MD), str(LIVE_TABLES), str(LIVE_METRICS), str(ADVISOR_COUNTS)],
        "secrets_included": False,
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps({"markdown": str(MD_PATH), "typst": str(TYPST_PATH), "docx": str(DOCX_PATH), "tables": len(tables), "risks": len(RISK_REGISTER)}, indent=2))


if __name__ == "__main__":
    main()
